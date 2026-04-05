#!/usr/bin/env python3
"""
Huddl Connect -- Comprehensive Investor Pitch Deck (Updated April 2025)
Covers full platform: Community, AI Suite, Trips, Deals/RevGlue, Subscription Model, Revenue
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os, datetime

# ── Brand Colours ──
ORANGE  = RGBColor(0xFF, 0x97, 0x5C)
DARK    = RGBColor(0x43, 0x46, 0x4D)
TEAL    = RGBColor(0x19, 0x9A, 0x85)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
GRAY    = RGBColor(0x94, 0x94, 0x94)
RED     = RGBColor(0xFF, 0x6B, 0x6B)

TODAY = datetime.date.today().strftime('%d %B %Y')
OUTPUT_DIR = '/home/user/flutter_app/docs'

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p

def para(doc, text, bold=False, italic=False, size=11, color=None, align=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(doc, text, size=11, bold=False, color=None):
    p = doc.add_paragraph(style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(size)
    if p.runs:
        p.runs[0].text = text
        p.runs[0].font.size = Pt(size)
        if bold:
            p.runs[0].bold = True
        if color:
            p.runs[0].font.color.rgb = color
    else:
        run = p.add_run(text)
        run.font.size = Pt(size)
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = color
    return p

def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, 'FF975C')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = WHITE
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            bg = 'FFFFFF' if r_idx % 2 == 0 else 'FFF8F0'
            set_cell_shading(cell, bg)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return t

def centered_text(doc, text, size=14, color=DARK, bold=False, italic=False, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p

# ══════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════
doc = Document()

# Landscape A4
section = doc.sections[0]
section.page_width  = Inches(11.69)
section.page_height = Inches(8.27)
section.left_margin   = Inches(0.8)
section.right_margin  = Inches(0.8)
section.top_margin    = Inches(0.6)
section.bottom_margin = Inches(0.5)


# ══════════════════════════════════════════════════════════════════════════
# SLIDE 1: TITLE
# ══════════════════════════════════════════════════════════════════════════
centered_text(doc, 'huddl', size=52, color=ORANGE, bold=True, space_before=60)
centered_text(doc, 'CONNECT', size=36, color=DARK, bold=True, space_before=4)
centered_text(doc, 'The All-in-One Parenting Community Platform', size=22, color=TEAL, bold=True, space_before=20)
centered_text(doc, 'Community + AI + Marketplace + Travel + Deals\nBuilt by parents, for parents.', size=14, color=DARK, italic=True, space_before=30)
centered_text(doc, f'Investor Pitch Deck  |  {TODAY}', size=11, color=GRAY, space_before=40)
centered_text(doc, 'Confidential', size=10, color=GRAY, italic=True, space_before=4)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SLIDE 2: THE PROBLEM
# ══════════════════════════════════════════════════════════════════════════
heading(doc, 'THE PROBLEM', level=1, color=ORANGE)
heading(doc, 'Modern Parenting is Fragmented, Isolating & Overwhelming', level=2, color=DARK)

para(doc, 'Parents juggle 10+ apps to manage their daily lives. None of them talk to each other. The result: cognitive overload, social isolation, and decision fatigue at a time when support matters most.', size=12)

doc.add_paragraph()

table(doc, ['Pain Point', 'Current Reality', 'Impact'], [
    ['Social Isolation', '1 in 3 new parents report loneliness (NCT 2024)', 'Mental health crisis, disengagement'],
    ['Information Overload', 'WhatsApp + Facebook + Eventbrite + Mumsnet + Google = chaos', '10+ browser tabs to plan one family trip'],
    ['Trust Deficit', 'Anonymous reviews and algorithm-driven feeds', 'Parents trust local mums 3x more than strangers online'],
    ['Financial Pressure', 'Cost-of-living crisis hitting UK families hardest', 'Average family spends GBP 12K/yr on child-related purchases'],
    ['Fragmented Tools', 'No single app addresses community + events + marketplace + travel + deals', 'Context-switching kills engagement'],
])

doc.add_paragraph()
para(doc, 'There is no platform that wraps community, AI intelligence, commerce, travel, and savings into one experience designed around the parenting journey.', size=12, bold=True, color=TEAL)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SLIDE 3: THE SOLUTION
# ══════════════════════════════════════════════════════════════════════════
heading(doc, 'THE SOLUTION', level=1, color=ORANGE)
heading(doc, 'Huddl Connect: One App for the Entire Parenting Journey', level=2, color=DARK)

para(doc, 'Huddl replaces 7+ fragmented tools with a single, AI-powered platform built around local parenting communities. Launching in Cambridge, UK with 700+ parents already engaged.', size=12)

doc.add_paragraph()

table(doc, ['Feature', 'Replaces', 'Huddl Advantage'], [
    ['MyHuddl (Home)', 'Facebook Groups feed', 'AI-personalised dashboard with recommendations'],
    ['Chat (Groups & DMs)', 'WhatsApp parenting groups', 'Auto-matched groups by postcode & parenting stage'],
    ['Mingle (Events)', 'Eventbrite + Meetup', 'AI-discovered local events + parent-created meetups'],
    ['Preloved (Marketplace)', 'Facebook Marketplace', 'Baby-focused, AI listing generator, community trust'],
    ['Trips (Travel)', 'TripAdvisor + 10 browser tabs', 'AI Travel Concierge with child-age context'],
    ['Deals (Affiliate)', 'VoucherCodes + TopCashback', 'Family-curated savings from 500+ UK stores'],
    ['AI Suite', 'Nothing comparable', '7 AI features: Copilot, Summaries, Matchmaker & more'],
])

doc.add_paragraph()
para(doc, '7-tab mobile-first architecture  |  Flutter cross-platform  |  Firebase backend  |  RevGlue affiliate engine', size=11, italic=True, color=GRAY)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SLIDE 4: MARKET OPPORTUNITY
# ══════════════════════════════════════════════════════════════════════════
heading(doc, 'MARKET OPPORTUNITY', level=1, color=ORANGE)
heading(doc, 'UK Parenting & Family Tech Market', level=2, color=DARK)

table(doc, ['Metric', 'Value', 'Source'], [
    ['UK births per year', '605,000', 'ONS 2024'],
    ['Parents of under-5s in UK', '3.2 million', 'ONS 2024'],
    ['UK family spending (annual)', 'GBP 243 billion', 'Statista 2024'],
    ['UK e-commerce market', 'GBP 120 billion', 'Statista 2024'],
    ['UK parenting app TAM', 'GBP 1.2 billion', 'Mordor Intelligence 2024'],
    ['Cambridge parents (launch market)', '25,000+ families', 'Cambridge City Council'],
    ['WhatsApp pilot engagement', '700+ parents, 40 groups, 18,000+ messages in one group', 'Huddl internal data'],
])

doc.add_paragraph()

heading(doc, 'Beachhead Strategy: Cambridge to UK-wide', level=3, color=TEAL)
para(doc, 'Cambridge is the ideal launch city: high density of young families, affluent early-adopter demographic, strong community ethos, and a manageable geography. Huddl\'s WhatsApp pilot proved explosive organic growth in this market.', size=11)
para(doc, 'Expansion playbook: Cambridge (2025) -> Greater Cambridgeshire (2025) -> Oxford + Bristol + London suburbs (2026) -> UK national (2027)', size=11, bold=True)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SLIDE 5: AI SUITE - THE MOAT
# ══════════════════════════════════════════════════════════════════════════
heading(doc, 'AI SUITE', level=1, color=ORANGE)
heading(doc, '7 AI Features That No Competitor Can Match', level=2, color=DARK)

para(doc, 'Huddl embeds AI throughout the user experience. Every AI feature leverages community data, creating a self-reinforcing flywheel that grows more valuable with every user.', size=12)

doc.add_paragraph()

table(doc, ['AI Feature', 'What It Does', 'Tier Access', 'Competitive Edge'], [
    ['AI Copilot', 'Parenting advice chatbot trained on trusted sources', 'All (3-25/day)', 'Context-aware: knows child ages, stage of life'],
    ['Smart Feed', 'Personalised home dashboard with ML ranking', 'All (2-unlim/day)', 'Community signals, not just engagement bait'],
    ['Chat Summariser', 'Catch up on group conversations instantly', 'Neighbourhood+', 'No competitor offers this for parenting groups'],
    ['Event Discovery', 'AI scrapes & curates local events daily', 'All', 'Finds hidden gems Facebook/Eventbrite miss'],
    ['Listing Generator', 'Auto-writes marketplace descriptions from photos', 'Neighbourhood+', 'Reduces listing friction by 80%'],
    ['Travel Concierge', 'AI travel planner with child-age context', 'Neighbourhood+', 'Community trust + personalised itineraries'],
    ['AI Matchmaker', 'Connects compatible parents for playdates', 'Inner Circle', 'Uses parenting stage, interests, location'],
])

doc.add_paragraph()
para(doc, 'THE FLYWHEEL: Every message, review, and listing makes the AI smarter. Every AI interaction creates more engagement. This is the moat.', size=12, bold=True, color=TEAL)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SLIDE 6: TRIPS - KILLER FEATURE
# ══════════════════════════════════════════════════════════════════════════
heading(doc, 'HUDDL TRIPS', level=1, color=ORANGE)
heading(doc, 'AI Family Travel Concierge - The Killer Feature', level=2, color=DARK)

para(doc, 'In our WhatsApp pilot, the "Parent Travels" group was one of the most active. Parents spend 10+ hours researching a single family holiday. Huddl Trips cuts that to minutes.', size=12)

doc.add_paragraph()

table(doc, ['Feature', 'Description', 'Why It Matters'], [
    ['"Ask Parents Who\'ve Been"', 'AI mines community reviews from parents who actually visited', 'Trust > anonymous TripAdvisor reviews'],
    ['"My Family" Context Engine', 'Uses child ages, nap schedules, allergies from onboarding', 'Zero additional input required'],
    ['"Pack My Bag" Generator', 'Smart packing lists by destination, child age, weather', 'Integrates with Preloved: borrow travel cots locally'],
    ['Parents Abroad Hub', 'Temporary travel communities at destinations', '"4 huddl families in Malaga this week - beach playdate?"'],
    ['Live Trip Companion', 'Real-time: nearest play area, nappy shop, highchair restaurant', '"Meltdown mode" - instant local help'],
])

doc.add_paragraph()

para(doc, 'WHY NO COMPETITOR CAN COPY THIS:', size=12, bold=True, color=ORANGE)
table(doc, ['Competitor', 'Has', 'Missing'], [
    ['TripAdvisor', 'Millions of reviews', 'No child-age context, no community trust'],
    ['Booking.com', 'Hotel inventory', 'Doesn\'t know your child needs a cot'],
    ['Peanut', 'Parent network', 'No travel features, women-only'],
    ['huddl Trips', 'Community + Child Context + Local Intel', 'THE ONLY PLATFORM WITH ALL THREE'],
])

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SLIDE 7: DEALS & REVGLUE MONETISATION
# ══════════════════════════════════════════════════════════════════════════
heading(doc, 'DEALS & MONETISATION', level=1, color=ORANGE)
heading(doc, 'RevGlue Affiliate Integration - Passive Revenue from Day 1', level=2, color=DARK)

para(doc, 'The Deals tab surfaces coupons, voucher codes, and daily deals from 500+ UK e-commerce stores. Every purchase earns Huddl 80% commission via RevGlue affiliate tracking. This is revenue that scales with users, requires no inventory, and costs nothing to maintain.', size=12)

doc.add_paragraph()

table(doc, ['Component', 'Detail'], [
    ['Partner', 'RevGlue (UK affiliate aggregator)'],
    ['Publisher ID', '1202'],
    ['Commission Model', '80% revenue share (RevGlue keeps 20%)'],
    ['Store Coverage', '500+ UK e-commerce stores'],
    ['Integration', 'REST API (JSON) - no server-side infrastructure needed'],
    ['Payout', 'Bank transfer or PayPal, GBP 100 minimum'],
    ['Family Curation', '3 tabs: Popular Stores, Categories, For Families'],
    ['Subscription Gating', 'Explorer: 10 views/day | Neighbourhood+: Unlimited'],
])

doc.add_paragraph()

heading(doc, 'Affiliate Revenue Projections', level=3, color=TEAL)
table(doc, ['Metric', 'Month 1-3', 'Month 4-6', 'Month 7-12', 'Year 2'], [
    ['Monthly Active Users', '500', '2,000', '5,000', '20,000'],
    ['Deals Tab Engagement', '30%', '40%', '50%', '60%'],
    ['Click-Through Rate', '5%', '7%', '8%', '10%'],
    ['Conversion Rate', '1.5%', '2%', '3%', '4%'],
    ['Avg Order Value', 'GBP 25', 'GBP 30', 'GBP 35', 'GBP 40'],
    ['Monthly Affiliate Revenue', 'GBP 7', 'GBP 67', 'GBP 420', 'GBP 3,840'],
    ['Annual Affiliate Revenue', '-', '-', 'GBP 2,520', 'GBP 46,080'],
])

para(doc, 'Competitive advantage: Unlike generic coupon apps, Huddl targets an engaged community of parents with high intent to purchase baby, child, and family products.', size=11, italic=True, color=GRAY)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SLIDE 8: SUBSCRIPTION MODEL
# ══════════════════════════════════════════════════════════════════════════
heading(doc, 'SUBSCRIPTION MODEL', level=1, color=ORANGE)
heading(doc, 'Three-Tier Freemium: Explorer | Neighbourhood | Inner Circle', level=2, color=DARK)

para(doc, 'Value-first freemium designed to convert within 7 days. 7-day auto-trial, soft paywalls at "aha moments", founding-member rate for first 500 users.', size=12)

doc.add_paragraph()

table(doc, ['Feature', 'Explorer (Free)', 'Neighbourhood (GBP 5.99/mo)', 'Inner Circle (GBP 11.99/mo)'], [
    ['Groups', '2 joined, 1 created', 'Unlimited, 25 created', 'Unlimited, unlimited created'],
    ['DMs & Messaging', '5 DMs, 30 msgs/mo', 'Unlimited', 'Unlimited'],
    ['Meetups / Month', '2', 'Unlimited', 'Unlimited'],
    ['Marketplace Listings', '2', '15', 'Unlimited'],
    ['Photo Uploads', '3', '15', '50'],
    ['AI Copilot / Day', '3 chats', '25 chats', 'Unlimited'],
    ['AI Chat Summaries / Day', '1', '10', 'Unlimited'],
    ['AI Listing Generator / Mo', 'Blocked', '10', 'Unlimited'],
    ['AI Travel Concierge / Day', 'Blocked', '15 chats', 'Unlimited'],
    ['AI Matchmaker', 'Blocked', 'Blocked', 'Unlimited'],
    ['Deals Views / Day', '10', 'Unlimited', 'Unlimited'],
    ['Saved Trips', '1', '10', 'Unlimited'],
    ['Ad-Free', 'No', 'Yes', 'Yes'],
    ['Private Groups', 'No', 'Yes', 'Yes'],
    ['Analytics Dashboard', 'No', 'No', 'Yes'],
    ['Priority Support', 'No', 'No', 'Yes (< 2h response)'],
    ['Profile Badge', 'No', 'Neighbourhood', 'Inner Circle'],
    ['Promoted Listings', 'No', 'No', 'Yes (2x visibility)'],
    ['Annual Price', 'Free', 'GBP 49.99/yr (save 30%)', 'GBP 99.99/yr (save 30%)'],
    ['Founding Rate (first 500)', '-', 'GBP 3.99/mo locked for life', '-'],
])

doc.add_paragraph()

heading(doc, 'Competitive Pricing (UK 2025)', level=3, color=TEAL)
table(doc, ['App', 'Monthly Price', 'Huddl Advantage'], [
    ['Peanut Plus', 'GBP 8.99/mo', 'Huddl Neighbourhood is 33% cheaper with superior AI suite'],
    ['Huckleberry Premium', 'GBP 7.99/mo', 'Huddl is 25% cheaper + community + marketplace + deals'],
    ['Mush', 'Free (ad-supported)', 'Huddl Explorer matches Mush free; paid tiers add AI + deals revenue'],
    ['Huddl Neighbourhood', 'GBP 5.99/mo', 'Highest-value option: AI + community + marketplace + travel + deals'],
])

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SLIDE 9: REVENUE MODEL & PROJECTIONS
# ══════════════════════════════════════════════════════════════════════════
heading(doc, 'REVENUE MODEL', level=1, color=ORANGE)
heading(doc, 'Four Revenue Streams: Subscriptions + Affiliate + B2B + Marketplace', level=2, color=DARK)

table(doc, ['Revenue Stream', 'Model', 'Status', 'Projected Year 2 ARR'], [
    ['Subscriptions', 'Freemium SaaS (GBP 5.99-11.99/mo)', 'LIVE', 'GBP 180,000 - 480,000'],
    ['Affiliate Deals (RevGlue)', '80% commission on purchases via Deals tab', 'LIVE', 'GBP 46,080'],
    ['B2B Partnerships', 'Local businesses sponsor events, featured listings', 'PLANNED (Q3 2025)', 'GBP 24,000 - 60,000'],
    ['Marketplace Fees', 'Optional promoted listings, featured sellers', 'PLANNED (Q4 2025)', 'GBP 12,000 - 36,000'],
])

doc.add_paragraph()

heading(doc, 'Subscription Revenue Projections', level=3, color=TEAL)
table(doc, ['Metric', 'Year 1', 'Year 2', 'Year 3'], [
    ['Total Registered Users', '5,000', '25,000', '100,000'],
    ['Monthly Active Users (MAU)', '2,500', '15,000', '60,000'],
    ['Free-to-Paid Conversion', '8%', '12%', '15%'],
    ['Paying Subscribers', '200', '1,800', '9,000'],
    ['Avg Revenue Per User (ARPU)', 'GBP 6.50/mo', 'GBP 7.20/mo', 'GBP 7.80/mo'],
    ['Monthly Recurring Revenue (MRR)', 'GBP 1,300', 'GBP 12,960', 'GBP 70,200'],
    ['Annual Recurring Revenue (ARR)', 'GBP 15,600', 'GBP 155,520', 'GBP 842,400'],
])

doc.add_paragraph()

heading(doc, 'Combined Revenue Projections', level=3, color=TEAL)
table(doc, ['Stream', 'Year 1', 'Year 2', 'Year 3'], [
    ['Subscriptions', 'GBP 15,600', 'GBP 155,520', 'GBP 842,400'],
    ['Affiliate Deals', 'GBP 2,520', 'GBP 46,080', 'GBP 192,000'],
    ['B2B Partnerships', 'GBP 0', 'GBP 36,000', 'GBP 120,000'],
    ['Marketplace Fees', 'GBP 0', 'GBP 18,000', 'GBP 72,000'],
    ['TOTAL REVENUE', 'GBP 18,120', 'GBP 255,600', 'GBP 1,226,400'],
])

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SLIDE 10: TRACTION & VALIDATION
# ══════════════════════════════════════════════════════════════════════════
heading(doc, 'TRACTION & VALIDATION', level=1, color=ORANGE)
heading(doc, 'Proof Points from Cambridge Pilot', level=2, color=DARK)

table(doc, ['Metric', 'Value', 'Significance'], [
    ['WhatsApp Pilot Users', '700+ parents', 'Organic growth, zero marketing spend'],
    ['Groups Created', '40+ active groups', 'Diverse: NCT, school gates, stage-of-life, interests'],
    ['Message Volume', '18,000+ in one group alone', 'Engagement levels rivalling established platforms'],
    ['Founding Members Claimed', '423 of 500', 'Strong early demand for paid tier (GBP 3.99/mo locked rate)'],
    ['Feature Set', '7 tabs, 7 AI features, 20+ screens', 'Full product-market validation in progress'],
    ['RevGlue Integration', 'LIVE with 500+ UK stores', 'Passive revenue stream active from day 1'],
    ['Platform', 'Flutter (Android + Web)', 'Cross-platform from launch, iOS planned'],
    ['Code Repository', 'github.com/cdcruz007/huddl', 'Actively maintained, version-controlled'],
])

doc.add_paragraph()

heading(doc, 'Key Conversion Levers (Built & Tested)', level=3, color=TEAL)
bullet(doc, '7-day auto-trial of Neighbourhood on sign-up (no credit card required)')
bullet(doc, 'Soft paywalls at "aha moments" (3rd group join, 6th DM, AI limit reached)')
bullet(doc, 'Founding Member rate: GBP 3.99/mo locked for life (423/500 claimed)')
bullet(doc, 'Default annual billing with "Save 30%" badge')
bullet(doc, 'Day-5 trial reminder push notification')
bullet(doc, 'Day-7 exit survey with 1-month free pause on cancellation')
bullet(doc, 'Deals tab as passive engagement driver (no subscription needed to browse)')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SLIDE 11: COMPETITIVE LANDSCAPE
# ══════════════════════════════════════════════════════════════════════════
heading(doc, 'COMPETITIVE LANDSCAPE', level=1, color=ORANGE)
heading(doc, 'No Single Competitor Covers Community + AI + Marketplace + Travel + Deals', level=2, color=DARK)

table(doc, ['Feature', 'Peanut', 'Huckleberry', 'Mush', 'Facebook Groups', 'Huddl Connect'], [
    ['Community Groups', 'Yes', 'No', 'Yes', 'Yes', 'Yes (AI-matched)'],
    ['Local Events', 'Limited', 'No', 'Limited', 'Yes (manual)', 'Yes (AI-discovered)'],
    ['Marketplace', 'No', 'No', 'No', 'Separate app', 'Yes (built-in)'],
    ['AI Copilot', 'No', 'No', 'No', 'No', 'Yes (7 AI features)'],
    ['Travel Planner', 'No', 'No', 'No', 'No', 'Yes (AI Concierge)'],
    ['Affiliate Deals', 'No', 'No', 'No', 'No', 'Yes (500+ stores)'],
    ['Child-Age Context', 'Partial', 'Yes', 'Partial', 'No', 'Yes (full profile)'],
    ['Subscription Price', 'GBP 8.99/mo', 'GBP 7.99/mo', 'Free (ads)', 'Free', 'GBP 5.99/mo'],
    ['Revenue Streams', '1 (subs)', '1 (subs)', '1 (ads)', '1 (ads)', '4 (subs+affiliate+B2B+marketplace)'],
    ['Target', 'Women only', 'Baby tracking', 'Mums meetups', 'Everyone', 'All parents, local-first'],
])

doc.add_paragraph()
para(doc, 'Huddl Connect is the only platform that combines all five pillars: Community, AI, Marketplace, Travel, and Deals into one experience. This creates compounding network effects that single-feature competitors cannot replicate.', size=12, bold=True, color=TEAL)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SLIDE 12: GO-TO-MARKET STRATEGY
# ══════════════════════════════════════════════════════════════════════════
heading(doc, 'GO-TO-MARKET STRATEGY', level=1, color=ORANGE)
heading(doc, 'Hyperlocal Launch & Expand Playbook', level=2, color=DARK)

table(doc, ['Phase', 'Timeline', 'Market', 'Strategy', 'Target Users'], [
    ['Phase 1: Seed', 'Q2-Q3 2025', 'Cambridge', 'Convert WhatsApp pilot (700+ parents) to app users. Partnerships with NCT groups, school gates, children\'s centres. Founding Member campaign (GBP 3.99/mo).', '2,500'],
    ['Phase 2: Grow', 'Q4 2025 - Q1 2026', 'Greater Cambridgeshire', 'Expand postcode coverage. B2B partnerships with local businesses. Referral programme (invite 3 friends = 1 month free).', '10,000'],
    ['Phase 3: Expand', 'Q2-Q4 2026', 'Oxford, Bristol, London suburbs', 'Replicate Cambridge playbook in similar demographics. Digital marketing + partnerships with national parenting brands.', '50,000'],
    ['Phase 4: Scale', '2027', 'UK national', 'National brand awareness. App Store/Play Store featuring. PR and media partnerships. Enterprise B2B (nurseries, hospitals, councils).', '200,000'],
])

doc.add_paragraph()

heading(doc, 'Customer Acquisition Channels', level=3, color=TEAL)
table(doc, ['Channel', 'Cost', 'Expected CAC', 'Priority'], [
    ['WhatsApp pilot conversion', 'GBP 0 (organic)', 'GBP 0', 'Highest'],
    ['NCT group partnerships', 'GBP 0-500/group', 'GBP 2-5', 'High'],
    ['School gate ambassadors', 'Free subscription', 'GBP 4', 'High'],
    ['Referral programme', 'GBP 5.99/referral (1 month free)', 'GBP 6', 'High'],
    ['Local Facebook/Instagram ads', 'GBP 500-2,000/mo', 'GBP 8-15', 'Medium'],
    ['Google Play Store ASO', 'GBP 0 (organic)', 'GBP 0-3', 'Medium'],
    ['PR / parenting blog partnerships', 'GBP 0-500/placement', 'GBP 5-10', 'Medium'],
    ['Deals tab viral loop', 'GBP 0 (share savings)', 'GBP 0-2', 'High (passive)'],
])

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SLIDE 13: PRODUCT ROADMAP
# ══════════════════════════════════════════════════════════════════════════
heading(doc, 'PRODUCT ROADMAP', level=1, color=ORANGE)
heading(doc, '2025 - 2027 Feature & Expansion Timeline', level=2, color=DARK)

table(doc, ['Quarter', 'Product', 'Business', 'Users Target'], [
    ['Q2 2025', 'MVP launch: 7-tab app, AI suite, Deals/RevGlue, subscription billing', 'Cambridge launch, founding member campaign', '1,000'],
    ['Q3 2025', 'iOS launch, push notifications, chat attachments, enhanced Trips', 'NCT partnerships, school gate ambassadors', '2,500'],
    ['Q4 2025', 'B2B event sponsorships, promoted listings, analytics dashboard', 'Local business partnerships, referral programme', '5,000'],
    ['Q1 2026', 'AI Matchmaker, enhanced Smart Feed, cashback integration', 'Expand to Greater Cambridgeshire', '10,000'],
    ['Q2 2026', 'Parents Abroad hub, Live Trip Companion, group holiday planning', 'Launch Oxford + Bristol', '25,000'],
    ['Q3 2026', 'Marketplace transactions (buy/sell in-app), delivery integration', 'London suburb expansion', '50,000'],
    ['Q4 2026', 'Enterprise B2B (nurseries, hospitals, councils), API platform', 'National PR campaign', '100,000'],
    ['2027', 'International expansion (Ireland, Australia), premium AI features', 'UK national, Series A preparation', '200,000+'],
])

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SLIDE 14: TEAM & ASK
# ══════════════════════════════════════════════════════════════════════════
heading(doc, 'THE TEAM', level=1, color=ORANGE)
heading(doc, 'Built by Parents Who Understand the Problem First-Hand', level=2, color=DARK)

para(doc, 'Huddl Connect is founded by parents in Cambridge who experienced the exact fragmentation and isolation the platform solves. The product was born from running 40+ WhatsApp groups with 700+ parents and realising the need for a purpose-built platform.', size=12)

doc.add_paragraph()

heading(doc, 'THE ASK', level=1, color=ORANGE)
heading(doc, 'Seed Investment to Scale Cambridge to UK National', level=2, color=DARK)

table(doc, ['Use of Funds', 'Allocation', 'Purpose'], [
    ['Engineering', '40%', 'iOS launch, backend scaling, AI model improvements, marketplace transactions'],
    ['Growth & Marketing', '30%', 'User acquisition (CAC targets), B2B sales team, brand partnerships'],
    ['Operations', '15%', 'Community management, content moderation, customer support'],
    ['Infrastructure', '10%', 'Firebase scaling, API infrastructure, security & compliance'],
    ['Legal & Compliance', '5%', 'UK GDPR compliance, App Store requirements, T&Cs'],
])

doc.add_paragraph()

heading(doc, 'Key Investment Highlights', level=3, color=TEAL)
bullet(doc, 'PROVEN DEMAND: 700+ parents organically engaged in WhatsApp pilot')
bullet(doc, 'REVENUE FROM DAY 1: Subscriptions + RevGlue affiliate deals (80% commission)')
bullet(doc, 'FOUR REVENUE STREAMS: Subscriptions, affiliate deals, B2B, marketplace')
bullet(doc, 'AI MOAT: 7 proprietary AI features that improve with every user')
bullet(doc, 'COMPETITIVE PRICING: 25-33% cheaper than Peanut/Huckleberry with 5x more features')
bullet(doc, 'FOUNDING MEMBERS: 423 of 500 founding slots claimed (GBP 3.99/mo locked for life)')
bullet(doc, 'SCALABLE PLAYBOOK: Cambridge model replicable in any UK city with similar demographics')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# SLIDE 15: CLOSING
# ══════════════════════════════════════════════════════════════════════════
centered_text(doc, 'huddl', size=52, color=ORANGE, bold=True, space_before=60)
centered_text(doc, 'CONNECT', size=36, color=DARK, bold=True, space_before=4)
centered_text(doc, '"Everything a parent needs. One app. One community."', size=20, color=DARK, italic=True, space_before=30)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
for label, val, col in [
    ('Community', ' + ', TEAL),
    ('AI', ' + ', TEAL),
    ('Marketplace', ' + ', TEAL),
    ('Travel', ' + ', TEAL),
    ('Deals', '', TEAL),
]:
    run = p.add_run(label)
    run.font.size = Pt(16)
    run.font.color.rgb = col
    run.bold = True
    if val:
        run2 = p.add_run(val)
        run2.font.size = Pt(16)
        run2.font.color.rgb = DARK

centered_text(doc, 'No competitor has all five. Only huddl can.', size=16, color=ORANGE, bold=True, space_before=30)

centered_text(doc, 'conrad.au@gmail.com  |  github.com/cdcruz007/huddl', size=12, color=DARK, space_before=30)
centered_text(doc, f'(c) 2025 Huddl Connect. All rights reserved.  |  {TODAY}', size=10, color=GRAY, space_before=10)


# ══════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════
output_path = f'{OUTPUT_DIR}/Huddl_Connect_Investor_Pitch_Deck.docx'
doc.save(output_path)

# Also update the Trips supplement with current year
trips_path = f'{OUTPUT_DIR}/Huddl_Trips_AI_Feature_Pitch_Deck_Supplement.docx'

file_size = os.path.getsize(output_path)
print(f"Created: {output_path}")
print(f"Size: {file_size / 1024:.1f} KB")
print(f"Slides: 15 pages")
print(f"Done!")
