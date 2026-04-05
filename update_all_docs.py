#!/usr/bin/env python3
"""
Update all Huddl Connect design documents with the Deals/RevGlue integration
and regenerate downloadable DOCX files.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, datetime

BRAND_ORANGE = RGBColor(0xFF, 0x97, 0x5C)
BRAND_DARK   = RGBColor(0x26, 0x2A, 0x35)
BRAND_TEAL   = RGBColor(0x19, 0x9A, 0x85)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
GRAY         = RGBColor(0x6C, 0x6C, 0x6C)

OUTPUT_DIR = '/home/user/flutter_app/docs'
os.makedirs(OUTPUT_DIR, exist_ok=True)

TODAY = datetime.date.today().strftime('%d %B %Y')

# ═══════════════════════════════════════════════════════════════════
# Helper functions
# ═══════════════════════════════════════════════════════════════════
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_DARK
    return h

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    s = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading.append(s)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, 'FF975C')
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


# ═══════════════════════════════════════════════════════════════════
# 1. FUNCTIONAL DESIGN DOCUMENT (Updated)
# ═══════════════════════════════════════════════════════════════════
def create_functional_design():
    doc = Document()
    
    # Title
    title = doc.add_heading('Huddl Connect', level=0)
    for run in title.runs:
        run.font.color.rgb = BRAND_ORANGE
        run.font.size = Pt(28)
    add_para(doc, 'Functional Design Document (FDD)', bold=True, size=16)
    add_para(doc, f'Version 2.1 | {TODAY}', italic=True, size=11)
    add_para(doc, 'Classification: Confidential', italic=True, size=10)
    doc.add_page_break()

    # Document Control
    add_heading(doc, 'Document Control')
    add_table(doc, ['Version', 'Date', 'Author', 'Changes'], [
        ['1.0', '01 Mar 2025', 'Product Team', 'Initial FDD'],
        ['2.0', '03 Apr 2025', 'Product Team', 'Added AI features, Trips, subscription model'],
        ['2.1', TODAY, 'Product Team', 'Added Deals tab with RevGlue affiliate integration, 7-tab navigation'],
    ])
    doc.add_page_break()

    # ToC placeholder
    add_heading(doc, 'Table of Contents')
    add_para(doc, '1. Introduction & Purpose')
    add_para(doc, '2. Product Overview')
    add_para(doc, '3. User Personas')
    add_para(doc, '4. Feature Specifications')
    add_para(doc, '5. Subscription Model')
    add_para(doc, '6. Deals & Monetisation (RevGlue)')
    add_para(doc, '7. Data Architecture')
    add_para(doc, '8. Non-Functional Requirements')
    add_para(doc, '9. Appendices')
    doc.add_page_break()

    # 1. Introduction
    add_heading(doc, '1. Introduction & Purpose')
    add_para(doc, 'Huddl Connect is a mobile-first community platform for parents in Cambridge, UK. It combines group messaging, local event discovery, a preloved marketplace, AI-powered features, family travel planning, and affiliate-powered deals into a single, cohesive experience designed around the parenting journey.')
    
    add_heading(doc, 'Intended Audience', level=2)
    add_bullet(doc, 'Development team (Flutter/Dart)')
    add_bullet(doc, 'Product owners and stakeholders')
    add_bullet(doc, 'QA engineers')
    add_bullet(doc, 'UX/UI designers')
    add_bullet(doc, 'Investor and business development teams')

    add_heading(doc, 'Scope', level=2)
    add_para(doc, 'This document covers all user-facing features, business rules, subscription tiers, and the RevGlue affiliate integration for the Deals monetisation channel.')

    # 2. Product Overview
    add_heading(doc, '2. Product Overview')
    
    add_heading(doc, 'Core Value Proposition', level=2)
    add_para(doc, 'A single app that replaces fragmented parenting tools: WhatsApp groups, Facebook Marketplace, Eventbrite, Mumsnet, and travel forums. Huddl adds AI intelligence and an affiliate deals engine to make the platform self-sustaining.')

    add_heading(doc, 'Primary Navigation - 7-Tab Architecture', level=2)
    add_table(doc, ['Tab Index', 'Name', 'Icon', 'Purpose'], [
        ['0', 'MyHuddl', 'home', 'Personalised dashboard with AI recommendations'],
        ['1', 'Chat', 'people', 'Group messaging and direct messages'],
        ['2', 'Mingle', 'groups', 'Events, meetups and activities'],
        ['3', 'Preloved', 'storefront', 'Buy, sell and give away baby items'],
        ['4', 'Trips', 'flight', 'Family travel planning with AI concierge'],
        ['5', 'Deals', 'local_offer', 'Affiliate coupons, vouchers and daily deals (RevGlue)'],
        ['6', 'Profile', 'person', 'Account settings, subscription management'],
    ])

    add_heading(doc, 'Platform Targets', level=2)
    add_bullet(doc, 'Android (primary) - Google Play Store')
    add_bullet(doc, 'iOS (planned) - Apple App Store')
    add_bullet(doc, 'Web (preview/testing)')

    add_heading(doc, 'Technology Stack', level=2)
    add_table(doc, ['Layer', 'Technology'], [
        ['Framework', 'Flutter 3.35.4 / Dart 3.9.2'],
        ['Backend', 'Firebase (Firestore, Auth, Storage, Messaging)'],
        ['State Management', 'Provider'],
        ['Affiliate Engine', 'RevGlue RevEmbed API (Publisher ID 1202)'],
        ['In-App Purchases', 'StoreKit (iOS) / Google Play Billing (Android)'],
        ['Local Storage', 'SharedPreferences, Hive'],
    ])

    # 3. User Personas
    add_heading(doc, '3. User Personas')
    
    personas = [
        ('Aspiring Parent - "Alex"', 'Researching parenthood, joining communities for advice and support.'),
        ('Expecting Parent - "Sarah"', 'Preparing for baby arrival, finding local groups and buying essentials. Deals tab helps find discounted baby gear.'),
        ('New Parent - "Marcus"', 'Sleep-deprived but eager to connect. Uses Chat and Mingle heavily. Discovers savings through Deals.'),
        ('Experienced Parent - "Priya"', 'Multiple children, sells outgrown items on Preloved, organises meetups, uses Trips for holidays. Power user of Deals for family shopping.'),
        ('Local Service Provider - "Dr. Chen"', 'Offers parenting classes, uses Events to promote. May sponsor deals through the platform.'),
    ]
    for name, desc in personas:
        add_heading(doc, name, level=2)
        add_para(doc, desc)

    # 4. Feature Specifications
    add_heading(doc, '4. Feature Specifications')
    
    features = [
        ('4.1 MyHuddl (Home Dashboard)', [
            'AI-recommended events and groups',
            'Quick actions: create meetup, list item, browse deals',
            'Community activity feed',
            'Personalised content based on postcode and parenting stage',
        ]),
        ('4.2 Chat (Groups & DMs)', [
            'Auto-assigned groups based on postcode and stage of life',
            'User-created public and private groups',
            'Direct messaging with online status',
            'Message threading and saved messages',
            'AI Chat Summariser for catching up on conversations',
        ]),
        ('4.3 Mingle (Events & Meetups)', [
            'AI-discovered local events (daily web scraping)',
            'B2B partner events from local businesses',
            'Parent-created meetups with RSVP',
            'Filters: area, date, age group, category',
            'Map view and list view',
        ]),
        ('4.4 Preloved (Marketplace)', [
            'Buy, sell and give away baby items',
            'AI Listing Generator for product descriptions',
            'Category browsing and search',
            'Direct messaging to sellers',
            'Photo uploads (up to 5 per listing)',
        ]),
        ('4.5 Trips (Family Travel)', [
            'AI Travel Concierge chatbot',
            'Destination browser with family-friendly ratings',
            'AI-generated packing lists',
            'Parents Abroad tips and advice',
            'Saved trips and travel planning',
        ]),
        ('4.6 Deals (RevGlue Affiliate Integration) [NEW]', [
            'Live coupons and voucher codes from 500+ UK stores',
            'Three sub-tabs: Popular Stores, Categories, For Families',
            'Store search with real-time filtering',
            'Store detail view with all active coupons/offers',
            'Copy-to-clipboard for coupon codes',
            'Affiliate exit-click tracking (80% commission)',
            'Family-curated picks (baby, toys, kids clothing)',
            'Banner carousel for featured promotions',
            'Subscription-gated daily view limits (Explorer: 10/day)',
            '25+ categories with subcategories',
        ]),
        ('4.7 Profile & Settings', [
            'Edit profile (name, bio, photo, children info)',
            'Subscription management and upgrade flow',
            'Notification preferences',
            'Privacy settings',
            'Run Tutorial option',
            'Help & Support, About, Legal links',
        ]),
    ]
    for title, bullets in features:
        add_heading(doc, title, level=2)
        for b in bullets:
            add_bullet(doc, b)

    # 5. Subscription Model
    add_heading(doc, '5. Subscription Model')
    add_para(doc, 'Three-tier subscription model with freemium conversion strategy:')
    
    add_table(doc, ['Feature', 'Explorer (Free)', 'Neighbourhood (GBP5.99/mo)', 'Inner Circle (GBP11.99/mo)'], [
        ['Groups Joined', '2', 'Unlimited', 'Unlimited'],
        ['Direct Messages', '5', 'Unlimited', 'Unlimited'],
        ['Meetups/Month', '2', 'Unlimited', 'Unlimited'],
        ['Marketplace Listings', '2', '15', 'Unlimited'],
        ['AI Copilot Chats/Day', '3', '25', 'Unlimited'],
        ['AI Chat Summaries/Day', '1', '10', 'Unlimited'],
        ['AI Listing Generator/Mo', '0', '10', 'Unlimited'],
        ['AI Travel Concierge/Day', '0', '15', 'Unlimited'],
        ['AI Matchmaker', 'No', 'No', 'Yes'],
        ['Deals Views/Day', '10', 'Unlimited', 'Unlimited'],
        ['Deals Cashback', 'No', 'Yes', 'Yes'],
        ['Exclusive Deals', 'No', 'No', 'Yes (early access)'],
        ['Ad-Free', 'No', 'Yes', 'Yes'],
        ['Profile Badge', 'No', 'Yes', 'Yes'],
        ['Priority Support', 'No', 'No', 'Yes'],
        ['Saved Trips', '1', '10', 'Unlimited'],
    ])

    # 6. Deals & Monetisation
    add_heading(doc, '6. Deals & Monetisation (RevGlue)')
    
    add_heading(doc, '6.1 Overview', level=2)
    add_para(doc, 'The Deals tab integrates RevGlue\'s RevEmbed Open API to surface affiliate coupons, voucher codes, and daily deals from 500+ UK e-commerce stores. This creates a passive revenue stream where Huddl earns 80% commission on every purchase made through the app\'s affiliate links.')
    
    add_heading(doc, '6.2 Technical Integration', level=2)
    add_table(doc, ['Component', 'Detail'], [
        ['Publisher ID', '1202'],
        ['Base API', 'https://www.revglue.com/iframeapi/'],
        ['Commission Model', '80% revenue share (RevGlue keeps 20%)'],
        ['Payout Threshold', 'GBP100 minimum (bank transfer or PayPal)'],
        ['Data Format', 'JSON REST API'],
        ['Cache Duration', '30 minutes (in-app)'],
        ['Exit-Click Tracking', 'https://www.revglue.com/revembed/coupon_exitclick/1202/{storeId}'],
    ])

    add_heading(doc, '6.3 API Endpoints Used', level=2)
    add_table(doc, ['Endpoint', 'Purpose', 'Returns'], [
        ['top_stores/1202', 'Popular stores with offer counts', 'List of store objects (id, title, icon, offer count)'],
        ['coupon_allcategories/1202', 'All coupon categories', 'Nested categories with subcategories'],
        ['coupon_menu_categories/1202', 'Navigation menu categories', 'Top-level categories with children'],
        ['store_detail/1202/{storeId}/{page}', 'Coupons for a specific store', 'Coupon objects (code, title, expiry, type)'],
        ['homepage_placement_banners/1202', 'Featured promotional banners', 'Banner objects (image, store link)'],
    ])

    add_heading(doc, '6.4 Revenue Projections', level=2)
    add_table(doc, ['Scenario', 'MAU', 'CTR', 'Conversion', 'Avg Order', 'Commission', 'Monthly Revenue'], [
        ['Conservative', '1,000', '5%', '2%', 'GBP30', '5%', 'GBP24'],
        ['Moderate', '5,000', '8%', '3%', 'GBP40', '7%', 'GBP672'],
        ['Optimistic', '20,000', '12%', '5%', 'GBP50', '10%', 'GBP9,600'],
    ])

    add_heading(doc, '6.5 User Experience Flow', level=2)
    add_para(doc, '1. User taps Deals tab in bottom navigation')
    add_para(doc, '2. Sees Popular Stores grid, Categories list, or Family Picks')
    add_para(doc, '3. Searches or browses for a store')
    add_para(doc, '4. Taps a store to see active coupons and offers')
    add_para(doc, '5. Taps a coupon to view details and copy code')
    add_para(doc, '6. Taps "Use Code & Shop" or "Get Offer" to open retailer site')
    add_para(doc, '7. Purchase tracked via RevGlue affiliate link')
    add_para(doc, '8. Commission credited to Huddl\'s RevGlue account')

    # 7. Data Architecture
    add_heading(doc, '7. Data Architecture')
    add_para(doc, 'Firebase Firestore collections: users, groups, messages, conversations, events, marketplace_items, trips, bookmarks.')
    add_para(doc, 'Local storage: SharedPreferences (settings, subscription state, usage counters), Hive (offline cache).')
    add_para(doc, 'External API: RevGlue iframeapi (deals data, cached 30 minutes in-app).')

    # 8. Non-Functional Requirements
    add_heading(doc, '8. Non-Functional Requirements')
    add_table(doc, ['Requirement', 'Target'], [
        ['App Launch Time', '< 3 seconds cold start'],
        ['API Response Time', '< 500ms for cached data, < 2s for fresh'],
        ['Offline Support', 'Core features available offline (cached data)'],
        ['Accessibility', 'WCAG 2.1 AA compliance'],
        ['Security', 'Firebase Auth, encrypted local storage'],
        ['Performance', '60fps smooth scrolling on mid-range devices'],
        ['Deals Availability', '500+ UK stores, updated daily by RevGlue'],
    ])

    doc.save(f'{OUTPUT_DIR}/Huddl_Connect_Functional_Design_Document.docx')
    print('  Created: Functional Design Document')


# ═══════════════════════════════════════════════════════════════════
# 2. SOFTWARE DESIGN DOCUMENT (Updated)
# ═══════════════════════════════════════════════════════════════════
def create_software_design():
    doc = Document()
    
    title = doc.add_heading('Huddl Connect', level=0)
    for run in title.runs:
        run.font.color.rgb = BRAND_ORANGE
        run.font.size = Pt(28)
    add_para(doc, 'Software Design Document (SDD)', bold=True, size=16)
    add_para(doc, f'Version 2.1 | {TODAY}', italic=True, size=11)
    doc.add_page_break()

    add_heading(doc, 'Document Control')
    add_table(doc, ['Version', 'Date', 'Author', 'Changes'], [
        ['1.0', '01 Mar 2025', 'Engineering', 'Initial SDD'],
        ['2.0', '03 Apr 2025', 'Engineering', 'Added AI services, Trips module, subscription service'],
        ['2.1', TODAY, 'Engineering', 'Added RevGlue service, Deals screen, 7-tab navigation'],
    ])
    doc.add_page_break()

    # 1. Introduction
    add_heading(doc, '1. Introduction')
    add_heading(doc, '1.1 Purpose', level=2)
    add_para(doc, 'This document describes the software architecture, component design, data models, and service integrations for Huddl Connect, a Flutter-based parenting community app.')
    
    add_heading(doc, '1.2 Scope', level=2)
    add_para(doc, 'Covers all application layers: UI (screens/widgets), business logic (services), data models, navigation, state management, external API integrations (Firebase, RevGlue), and the subscription/monetisation system.')

    add_heading(doc, '1.3 Technology Stack', level=2)
    add_table(doc, ['Component', 'Technology', 'Version'], [
        ['Framework', 'Flutter', '3.35.4'],
        ['Language', 'Dart', '3.9.2'],
        ['Backend', 'Firebase (Firestore, Auth, Storage)', 'Latest compatible'],
        ['State Management', 'Provider', '6.1.5+1'],
        ['HTTP Client', 'http package', '1.5.0'],
        ['Affiliate API', 'RevGlue iframeapi', 'Publisher 1202'],
        ['Local Storage', 'SharedPreferences / Hive', '2.5.3 / 2.2.3'],
        ['In-App Purchases', 'in_app_purchase', '3.2.0'],
        ['URL Launcher', 'url_launcher', '6.2.5'],
    ])

    # 2. System Architecture
    add_heading(doc, '2. System Architecture')
    add_heading(doc, '2.1 High-Level Architecture', level=2)
    add_para(doc, 'The app follows a layered architecture:')
    add_bullet(doc, 'Presentation Layer: Flutter screens, widgets, theme')
    add_bullet(doc, 'Business Logic Layer: Services (singleton pattern)')
    add_bullet(doc, 'Data Layer: Firebase Firestore, RevGlue API, local storage')
    add_bullet(doc, 'Navigation Layer: Named routes via AppRouter')

    add_heading(doc, '2.2 Navigation Architecture', level=2)
    add_para(doc, 'MainShell uses an IndexedStack with 7 child screens and a custom bottom navigation bar. The AppRouter handles all named route transitions with custom page animations (FadePageRoute, SlidePageRoute, ScalePageRoute).')
    
    add_table(doc, ['Index', 'Screen', 'File', 'Description'], [
        ['0', 'HomeScreen', 'lib/screens/home/home_screen.dart', 'AI-powered dashboard'],
        ['1', 'GroupsScreen', 'lib/screens/groups/groups_screen.dart', 'Groups and DMs'],
        ['2', 'EventsScreen', 'lib/screens/events/events_screen.dart', 'Events and meetups'],
        ['3', 'MarketplaceScreen', 'lib/screens/marketplace/marketplace_screen.dart', 'Preloved marketplace'],
        ['4', 'TripsScreen', 'lib/screens/trips/trips_screen.dart', 'Family travel'],
        ['5', 'DealsScreen', 'lib/screens/deals/deals_screen.dart', 'RevGlue affiliate deals'],
        ['6', 'ProfileScreen', 'lib/screens/profile/profile_screen.dart', 'User profile and settings'],
    ])

    # 3. Service Layer
    add_heading(doc, '3. Service Layer')
    add_para(doc, 'All services use the singleton pattern for shared state:')
    
    add_table(doc, ['Service', 'File', 'Responsibility'], [
        ['SubscriptionService', 'subscription_service.dart', 'Tier management, feature gating, usage tracking'],
        ['RevGlueService', 'revglue_service.dart', 'Affiliate deals API, store/coupon data, caching'],
        ['TutorialService', 'tutorial_service.dart', 'Onboarding tutorial state and step definitions'],
        ['OTPService', 'otp_service.dart', 'Phone verification OTP generation and validation'],
        ['TestAccountService', 'test_account_service.dart', 'Demo account profiles and data'],
        ['FirebaseAuthService', 'firebase_auth_service.dart', 'Authentication flows'],
        ['EventService', 'event_service.dart', 'Event data management'],
        ['RehomeService', 'rehome_service.dart', 'Marketplace listing management'],
        ['TravelService', 'travel_service.dart', 'Travel destination and trip data'],
        ['AiCopilotService', 'ai_copilot_service.dart', 'AI chatbot conversations'],
        ['AiChatSummariserService', 'ai_chat_summariser_service.dart', 'Chat summary generation'],
        ['AiListingGeneratorService', 'ai_listing_generator_service.dart', 'AI marketplace listings'],
        ['AiMatchmakerService', 'ai_matchmaker_service.dart', 'AI parent matching'],
        ['AiSmartFeedService', 'ai_smart_feed_service.dart', 'Personalised feed ranking'],
    ])

    # 4. RevGlue Integration (NEW)
    add_heading(doc, '4. RevGlue Integration (Deals Service)')
    
    add_heading(doc, '4.1 Service Architecture', level=2)
    add_para(doc, 'RevGlueService is a singleton that wraps the RevGlue iframeapi REST endpoints. It provides cached access to store, category, and coupon data with a 30-minute TTL.')
    
    add_heading(doc, '4.2 Data Models', level=2)
    add_table(doc, ['Model', 'Key Fields', 'Source'], [
        ['RevGlueStore', 'id, title, titleUrl, offerCouponStr, storeIcon, storeLargeIcon', 'top_stores endpoint'],
        ['RevGlueCategory', 'id, title, catUrl, offerCouponStr, smallIcon, subCategories', 'coupon_allcategories endpoint'],
        ['RevGlueCoupon', 'id, storeId, storeTitle, voucherCode, voucherTitle, expiryDate, offerCoupon', 'store_detail endpoint'],
        ['RevGlueBanner', 'src, storeId, titleUrl, title', 'homepage_placement_banners endpoint'],
    ])

    add_heading(doc, '4.3 Affiliate Tracking', level=2)
    add_para(doc, 'All outbound clicks route through RevGlue\'s exit-click URLs which embed the Publisher ID (1202) for commission attribution:')
    add_bullet(doc, 'Coupon clicks: https://www.revglue.com/revembed/coupon_exitclick/1202/{storeId}')
    add_bullet(doc, 'Daily deal clicks: https://www.revglue.com/revembed/daily_deal_exitclick/1202/{storeId}/{dealId}')
    add_para(doc, 'The url_launcher package opens these URLs in the device\'s default browser, ensuring proper cookie tracking for commission attribution.')

    add_heading(doc, '4.4 Caching Strategy', level=2)
    add_table(doc, ['Data', 'Cache Duration', 'Refresh Trigger'], [
        ['Top Stores', '30 minutes', 'Pull-to-refresh or cache expiry'],
        ['Categories', '30 minutes', 'Pull-to-refresh or cache expiry'],
        ['Store Coupons', 'No cache (fresh per visit)', 'Each store tap'],
        ['Banners', 'Per session', 'App restart'],
    ])

    # 5. Subscription & Feature Gating
    add_heading(doc, '5. Subscription & Feature Gating')
    add_para(doc, 'SubscriptionService manages three tiers (Explorer, Neighbourhood, Inner Circle) with feature limits tracked via usage counters in local storage. The Deals screen checks canViewMoreDeals before allowing store detail views, showing an upgrade dialog at the limit.')
    
    add_table(doc, ['Feature Gate', 'Explorer', 'Neighbourhood', 'Inner Circle'], [
        ['canJoinGroup', '2 max', 'Unlimited', 'Unlimited'],
        ['canCreateMeetup', 'No', 'Yes', 'Yes'],
        ['canUseAiCopilot', '3/day', '25/day', 'Unlimited'],
        ['canViewDeals', '10/day', 'Unlimited', 'Unlimited'],
        ['canAccessCashback', 'No', 'Yes', 'Yes'],
    ])

    # 6. Screen Component Design
    add_heading(doc, '6. Screen Component Design')
    add_heading(doc, '6.1 DealsScreen Architecture', level=2)
    add_para(doc, 'The DealsScreen uses a NestedScrollView with a SliverPersistentHeader for the pinned TabBar. It manages two view states: main view (3 tabs) and store detail view (back navigation within the screen).')
    
    add_table(doc, ['Component', 'Type', 'Responsibility'], [
        ['DealsScreen', 'StatefulWidget', 'Main container with TabController and state management'],
        ['_StoreCard', 'StatelessWidget', 'Grid tile showing store logo, name, offer count'],
        ['_StoreListTile', 'StatelessWidget', 'Horizontal list tile for family picks'],
        ['_CategoryTile', 'StatelessWidget', 'Category row with icon, name, subcategory count'],
        ['_CouponCard', 'StatelessWidget', 'Coupon row with type badge (Code/Offer), title, expiry'],
        ['_TabBarDelegate', 'SliverPersistentHeaderDelegate', 'Pinned tab bar for Popular/Categories/Families'],
    ])

    # 7. File Structure
    add_heading(doc, '7. Project File Structure')
    add_para(doc, '''lib/
  main.dart
  config/
    router.dart
  models/
    subscription.dart
  screens/
    main_shell.dart
    home/ home_screen.dart
    groups/ groups_screen.dart, group_chat_screen.dart, ...
    events/ events_screen.dart, create_meetup_screen.dart, ...
    marketplace/ marketplace_screen.dart, item_detail_screen.dart
    trips/ trips_screen.dart, travel_concierge_screen.dart, ...
    deals/ deals_screen.dart
    ai/ ai_copilot_screen.dart, ...
    profile/ profile_screen.dart
    subscription/ subscription_plans_screen.dart, ...
    onboarding/ splash_screen.dart, ...
    auth/ login_screen.dart, login_otp_screen.dart
  services/
    revglue_service.dart
    subscription_service.dart
    tutorial_service.dart
    firebase_auth_service.dart
    event_service.dart
    rehome_service.dart
    travel_service.dart
    ai_copilot_service.dart
    ...
  theme/
    huddl_colors.dart
    huddl_theme.dart
  widgets/
    tutorial/ tutorial_overlay.dart
  utils/
    page_transitions.dart''', size=9)

    doc.save(f'{OUTPUT_DIR}/Huddl_Connect_Software_Design_Document.docx')
    print('  Created: Software Design Document')


# ═══════════════════════════════════════════════════════════════════
# 3. UX DESIGN DOCUMENTATION (Updated)
# ═══════════════════════════════════════════════════════════════════
def create_ux_design():
    doc = Document()
    
    title = doc.add_heading('Huddl Connect', level=0)
    for run in title.runs:
        run.font.color.rgb = BRAND_ORANGE
        run.font.size = Pt(28)
    add_para(doc, 'UX Design Documentation', bold=True, size=16)
    add_para(doc, f'Version 2.1 | {TODAY}', italic=True, size=11)
    doc.add_page_break()

    add_heading(doc, '1. Executive Summary')
    add_para(doc, 'Huddl Connect is designed as a warm, approachable mobile experience for parents. The UX prioritises simplicity, trust, and delight. Version 2.1 introduces the Deals tab, bringing the bottom navigation to 7 tabs and adding a new monetisation-driven feature that enhances the user value proposition.')

    add_heading(doc, '2. Design Principles')
    principles = [
        ('Trust First', 'Every interaction reinforces safety and community trust. Verified accounts, moderated content, and transparent data practices.'),
        ('Reduce Cognitive Load', 'Parents are busy. Surface the right content at the right time. AI-powered recommendations replace manual searching.'),
        ('Celebrate Milestones', 'The parenting journey is full of milestones. The app acknowledges and celebrates them.'),
        ('Save Money, Save Time', 'The Deals tab directly addresses parents\' financial pressures with curated savings opportunities.'),
    ]
    for name, desc in principles:
        add_heading(doc, name, level=2)
        add_para(doc, desc)

    add_heading(doc, '3. Information Architecture')
    add_heading(doc, '3.1 Navigation Model - 7-Tab Bottom Bar', level=2)
    add_para(doc, 'The bottom navigation uses a custom floating pill-shaped bar with 7 equally-spaced icons. Each tab has an outlined (inactive) and filled (active) icon variant with the Huddl orange accent colour.')
    
    add_table(doc, ['Tab', 'Primary Action', 'Secondary Actions'], [
        ['MyHuddl', 'View dashboard', 'Tap cards to explore sections'],
        ['Chat', 'Read messages', 'Create group, start DM, search'],
        ['Mingle', 'Browse events', 'Create meetup (FAB), filter, map view'],
        ['Preloved', 'Browse items', 'Create listing, search, filter by category'],
        ['Trips', 'Browse destinations', 'Travel concierge, packing lists, parents abroad'],
        ['Deals', 'Browse stores/coupons', 'Search stores, view categories, copy codes'],
        ['Profile', 'View profile', 'Edit, subscription, settings, tutorial'],
    ])

    add_heading(doc, '3.2 Complete Screen Inventory', level=2)
    add_table(doc, ['Screen', 'Route', 'Access'], [
        ['Splash', '/splash', 'App launch'],
        ['Onboarding Carousel', '/onboarding', 'First launch'],
        ['Name Input', '/name_input', 'Onboarding flow'],
        ['Parent Type', '/parent_type', 'Onboarding flow'],
        ['Stage of Life', '/stage_of_life', 'Onboarding flow'],
        ['Postcode', '/postcode', 'Onboarding flow'],
        ['Phone Verification', '/verification', 'Auth flow'],
        ['Login', '/login', 'Returning users'],
        ['Home (MainShell)', '/home', 'Post-auth'],
        ['Group Chat', '/group_chat', 'Chat tab'],
        ['DM Chat', '/dm_chat', 'Chat tab'],
        ['Event Detail', '/event_detail', 'Mingle tab'],
        ['Create Meetup', 'Push route', 'Mingle FAB'],
        ['Item Detail', '/item_detail', 'Preloved tab'],
        ['Create Listing', '/create_listing', 'Preloved tab'],
        ['Destination Detail', '/destination_detail', 'Trips tab'],
        ['Travel Concierge', '/travel_concierge', 'Trips tab'],
        ['Deals (in MainShell)', 'Tab index 5', 'Deals tab'],
        ['Subscription Plans', '/subscription_plans', 'Profile or gate'],
        ['Manage Subscription', '/manage_subscription', 'Profile tab'],
    ])

    # 4. Deals Tab UX Design
    add_heading(doc, '4. Deals Tab - UX Design')
    
    add_heading(doc, '4.1 Screen Layout', level=2)
    add_para(doc, 'The Deals screen uses a NestedScrollView with a collapsible header containing:')
    add_bullet(doc, 'Brand header with gradient tag icon, title "Deals", subtitle "Save money on top UK brands"')
    add_bullet(doc, 'Explorer limit badge (shows remaining daily views)')
    add_bullet(doc, 'Search bar with real-time store filtering')
    add_bullet(doc, 'Promotional banner carousel (horizontal scroll)')
    add_bullet(doc, 'Pinned TabBar with 3 tabs: Popular Stores | Categories | For Families')

    add_heading(doc, '4.2 Interaction Patterns', level=2)
    add_table(doc, ['Interaction', 'Behaviour'], [
        ['Tap store card', 'Navigate to in-screen store detail with back button'],
        ['Tap coupon card', 'Show bottom sheet with code, copy button, and shop link'],
        ['Copy code button', 'Copy to clipboard + green snackbar confirmation'],
        ['Use Code & Shop', 'Opens retailer website in default browser (affiliate tracked)'],
        ['Pull-to-refresh', 'Refreshes store and category data from RevGlue API'],
        ['Search bar typing', 'Real-time filter of store grid'],
        ['Limit reached (Explorer)', 'Shows upgrade dialog with Neighbourhood CTA'],
        ['Category tap', 'Shows bottom sheet with subcategories'],
    ])

    add_heading(doc, '4.3 Visual Design', level=2)
    add_para(doc, 'The Deals screen follows the Huddl design system:')
    add_bullet(doc, 'Store cards: White containers with 16px border radius, subtle shadow, store logo, and green "X Codes" badge')
    add_bullet(doc, 'Coupon cards: Horizontal layout with type badge (orange for Code, green for Offer), title, expiry date')
    add_bullet(doc, 'Category tiles: Left icon with contextual colour, category name, subcategory count')
    add_bullet(doc, 'Family picks banner: Peach gradient background with family icon')
    add_bullet(doc, 'Code modal: Centered code in source-code font with copy icon, prominent CTA button')

    add_heading(doc, '4.4 Colour Mapping', level=2)
    add_table(doc, ['Element', 'Colour', 'Hex'], [
        ['Header gradient', 'Red to Orange', '#FF6B6B to #FF975C'],
        ['Code badge', 'Huddl Orange', '#FF975C'],
        ['Offer badge', 'Huddl Teal', '#199A85'],
        ['Family banner', 'Peach gradient', '#FFF3ED to #FFF8F0'],
        ['Explorer limit badge', 'Yellow', '#FFF7C9'],
        ['Success snackbar', 'Teal', '#199A85'],
    ])

    # 5. Tutorial Flow
    add_heading(doc, '5. Onboarding Tutorial')
    add_para(doc, 'The tutorial overlay walks users through all 7 tabs with a swipeable card interface:')
    add_table(doc, ['Step', 'Tab', 'Headline', 'CTA'], [
        ['1', 'MyHuddl', 'Your personalised dashboard', 'Tap cards to explore'],
        ['2', 'Chat', 'Your groups & conversations', 'Tap a group to start chatting'],
        ['3', 'Mingle', 'Events, meetups & activities', 'Tap + to create a meetup'],
        ['4', 'Preloved', 'Buy, sell & give away baby items', 'Tap + to list an item'],
        ['5', 'Trips', 'Family travel made easy', 'Tap Explore to browse destinations'],
        ['6', 'Deals', 'Save money on top UK brands', 'Tap a store to see offers'],
        ['7', 'Profile', 'Your account & settings', 'Scroll down for settings'],
    ])

    # 6. Accessibility
    add_heading(doc, '6. Accessibility')
    add_bullet(doc, 'Minimum touch target: 44x44px for all interactive elements')
    add_bullet(doc, 'Colour contrast: WCAG 2.1 AA minimum (4.5:1 for text)')
    add_bullet(doc, 'Screen reader support: Semantic labels on all icons and buttons')
    add_bullet(doc, 'Font scaling: Respects system text size preferences')
    add_bullet(doc, 'Reduced motion: Animations respect system accessibility settings')

    doc.save(f'{OUTPUT_DIR}/Huddl_Connect_UX_Design_Documentation.docx')
    print('  Created: UX Design Documentation')


# ═══════════════════════════════════════════════════════════════════
# 4. AI SYSTEM DESIGN DOCUMENT (Updated)
# ═══════════════════════════════════════════════════════════════════
def create_ai_system_design():
    doc = Document()
    
    title = doc.add_heading('Huddl Connect', level=0)
    for run in title.runs:
        run.font.color.rgb = BRAND_ORANGE
        run.font.size = Pt(28)
    add_para(doc, 'AI System Design Document', bold=True, size=16)
    add_para(doc, f'Version 2.1 | {TODAY}', italic=True, size=11)
    doc.add_page_break()

    add_heading(doc, 'Document Control')
    add_table(doc, ['Version', 'Date', 'Author', 'Changes'], [
        ['1.0', '01 Mar 2025', 'AI Team', 'Initial AI system design'],
        ['2.0', '03 Apr 2025', 'AI Team', 'Added Travel Concierge, Matchmaker, Smart Feed'],
        ['2.1', TODAY, 'AI Team', 'Added Deals recommendation engine, family curation logic'],
    ])
    doc.add_page_break()

    add_heading(doc, '1. Executive Summary')
    add_para(doc, 'Huddl Connect embeds AI throughout the user experience. This document covers all AI-powered features, model architectures, data pipelines, and the intelligent content curation system that powers the Deals family picks.')

    add_heading(doc, '2. AI Feature Inventory')
    add_table(doc, ['Feature', 'Type', 'Tier Access', 'Description'], [
        ['AI Copilot', 'Conversational AI', 'All (limited)', 'Parenting advice chatbot'],
        ['Smart Feed', 'Ranking/Recommendation', 'All (limited)', 'Personalised home feed'],
        ['Chat Summariser', 'NLP Summarisation', 'Neighbourhood+', 'Catch up on group chats'],
        ['Event Discovery', 'Web Scraping + NLP', 'All', 'AI-discovered local events'],
        ['Listing Generator', 'Text Generation', 'Neighbourhood+', 'Auto-generate marketplace descriptions'],
        ['Travel Concierge', 'Conversational AI', 'Neighbourhood+', 'Family travel planning chatbot'],
        ['AI Matchmaker', 'Recommendation', 'Inner Circle', 'Connect compatible parents'],
        ['Deals Curation', 'Rule-based + NLP', 'All', 'Family-friendly deal recommendations'],
    ])

    add_heading(doc, '3. Deals Curation Engine')
    add_heading(doc, '3.1 Family Relevance Scoring', level=2)
    add_para(doc, 'The Deals "For Families" tab uses keyword-based relevance scoring to surface stores and categories most relevant to parents:')
    add_bullet(doc, 'Primary keywords (high weight): baby, child, kids, toys, family, school')
    add_bullet(doc, 'Store name matching: vertbaudet, scholastic, hamleys, baker ross, start rite, mountain warehouse, picniq, my 1st years')
    add_bullet(doc, 'Category matching: Baby & Child, Education, Sports, Toys and Games')
    
    add_heading(doc, '3.2 Future AI Enhancements', level=2)
    add_bullet(doc, 'Personalised deal recommendations based on user profile (child ages, interests)')
    add_bullet(doc, 'Seasonal deal surfacing (back-to-school, Christmas, summer holidays)')
    add_bullet(doc, 'Collaborative filtering: "Parents like you also saved..."')
    add_bullet(doc, 'Price drop alerts for saved/bookmarked items')
    add_bullet(doc, 'Natural language deal search: "Find deals on pushchairs under GBP200"')

    add_heading(doc, '4. AI Usage Limits by Tier')
    add_table(doc, ['AI Feature', 'Explorer', 'Neighbourhood', 'Inner Circle'], [
        ['Copilot Chats/Day', '3', '25', 'Unlimited'],
        ['Chat Summaries/Day', '1', '10', 'Unlimited'],
        ['Event Discovery', '1/week', 'Daily', 'Unlimited'],
        ['Listing Generator/Month', '0', '10', 'Unlimited'],
        ['Travel Concierge/Day', '0', '15', 'Unlimited'],
        ['AI Matchmaker', 'No', 'No', 'Yes'],
        ['Smart Feed/Day', '2', 'Unlimited', 'Unlimited'],
        ['Deals Views/Day', '10', 'Unlimited', 'Unlimited'],
    ])

    add_heading(doc, '5. Data Privacy & Ethics')
    add_para(doc, 'All AI features comply with UK GDPR and the Children\'s Code. User data is processed locally where possible (on-device models), and cloud AI calls are anonymised. The Deals feature does not share user data with RevGlue - only exit-click tracking occurs.')

    doc.save(f'{OUTPUT_DIR}/Huddl_Connect_AI_System_Design_Document.docx')
    print('  Created: AI System Design Document')


# ═══════════════════════════════════════════════════════════════════
# 5. STORE PUBLISHING GUIDE (Updated)
# ═══════════════════════════════════════════════════════════════════
def create_store_publishing_guide():
    doc = Document()
    
    title = doc.add_heading('Huddl Connect', level=0)
    for run in title.runs:
        run.font.color.rgb = BRAND_ORANGE
        run.font.size = Pt(28)
    add_para(doc, 'Store Publishing & Monetisation Guide', bold=True, size=16)
    add_para(doc, f'Version 2.1 | {TODAY}', italic=True, size=11)
    doc.add_page_break()

    add_heading(doc, '1. App Identifiers')
    add_table(doc, ['Identifier', 'Value'], [
        ['App Name', 'Huddl Connect'],
        ['Package Name (Android)', 'com.huddlconnect.connect'],
        ['Bundle ID (iOS)', 'com.huddlconnect.connect'],
        ['RevGlue Publisher ID', '1202'],
        ['RevGlue Registered Domain', 'Www.huddlparents.com'],
    ])

    add_heading(doc, '2. Revenue Streams')
    add_table(doc, ['Stream', 'Model', 'Expected Revenue'], [
        ['Subscriptions', 'Freemium (GBP5.99-11.99/mo)', 'Primary revenue'],
        ['RevGlue Affiliate Deals', '80% commission per sale', 'Passive secondary revenue'],
        ['Future: Sponsored Events', 'B2B partnerships', 'Planned'],
        ['Future: Promoted Listings', 'Inner Circle feature', 'Planned'],
    ])

    add_heading(doc, '3. RevGlue Affiliate Setup')
    add_heading(doc, '3.1 Account Details', level=2)
    add_table(doc, ['Setting', 'Value'], [
        ['Account Email', 'Conrad.au@gmail.com'],
        ['Publisher ID', '1202'],
        ['Project Type', 'Coupons and Daily Deals UK'],
        ['Registered Domain', 'Www.huddlparents.com'],
        ['Commission Rate', '80% (RevGlue keeps 20%)'],
        ['Payout Method', 'Bank Transfer or PayPal'],
        ['Minimum Payout', 'GBP100'],
    ])

    add_heading(doc, '3.2 Technical Integration', level=2)
    add_para(doc, 'The Flutter app integrates RevGlue\'s iframeapi REST endpoints:')
    add_bullet(doc, 'Base URL: https://www.revglue.com/iframeapi/')
    add_bullet(doc, 'All requests append Publisher ID: /1202')
    add_bullet(doc, 'Exit-click URLs: https://www.revglue.com/revembed/coupon_exitclick/1202/{storeId}')
    add_bullet(doc, 'Data format: JSON')
    add_bullet(doc, 'No authentication required for data endpoints')

    add_heading(doc, '3.3 Revenue Tracking', level=2)
    add_para(doc, 'Revenue is tracked automatically by RevGlue when users click through affiliate exit-click URLs. The RevGlue dashboard at revglue.com shows:')
    add_bullet(doc, 'Daily/weekly/monthly click counts')
    add_bullet(doc, 'Conversion rates per store')
    add_bullet(doc, 'Confirmed vs pending commissions')
    add_bullet(doc, 'Payout history')

    add_heading(doc, '4. Google Play Store Listing')
    add_heading(doc, '4.1 Store Description', level=2)
    add_para(doc, 'Huddl Connect - The parenting community app for Cambridge families. Join local parent groups, discover events, buy and sell baby items, plan family holidays, and save money with exclusive deals from 500+ UK stores. AI-powered features help you find what matters most.', italic=True)
    
    add_heading(doc, '4.2 Key Features for Store Listing', level=2)
    add_bullet(doc, 'Join local parent groups matched to your area and parenting stage')
    add_bullet(doc, 'Discover AI-curated events and meetups near you')
    add_bullet(doc, 'Buy, sell and give away preloved baby items')
    add_bullet(doc, 'Plan family holidays with an AI travel concierge')
    add_bullet(doc, 'Save money with coupons and deals from top UK brands')
    add_bullet(doc, 'Get parenting advice from an AI-powered copilot')

    add_heading(doc, '5. Subscription IAP Configuration')
    add_table(doc, ['Product ID', 'Title', 'Price', 'Period'], [
        ['neighbourhood_monthly', 'Neighbourhood Monthly', 'GBP5.99', 'Monthly'],
        ['neighbourhood_annual', 'Neighbourhood Annual', 'GBP49.99', 'Annual'],
        ['inner_circle_monthly', 'Inner Circle Monthly', 'GBP11.99', 'Monthly'],
        ['inner_circle_annual', 'Inner Circle Annual', 'GBP99.99', 'Annual'],
    ])

    doc.save(f'{OUTPUT_DIR}/Huddl_Connect_Store_Publishing_Guide.docx')
    print('  Created: Store Publishing & Monetisation Guide')


# ═══════════════════════════════════════════════════════════════════
# 6. DEPLOYMENT GUIDE (Updated)
# ═══════════════════════════════════════════════════════════════════
def create_deployment_guide():
    doc = Document()
    
    title = doc.add_heading('Huddl Connect', level=0)
    for run in title.runs:
        run.font.color.rgb = BRAND_ORANGE
        run.font.size = Pt(28)
    add_para(doc, 'Deployment & Configuration Guide', bold=True, size=16)
    add_para(doc, f'Version 2.1 | {TODAY}', italic=True, size=11)
    doc.add_page_break()

    add_heading(doc, '1. Project Overview')
    add_table(doc, ['Component', 'Technology'], [
        ['Framework', 'Flutter 3.35.4 / Dart 3.9.2'],
        ['Backend', 'Firebase (Firestore, Auth, Storage, Messaging)'],
        ['Affiliate Engine', 'RevGlue (Publisher ID 1202)'],
        ['In-App Purchases', 'in_app_purchase 3.2.0'],
        ['Android SDK', 'API Level 35 (Android 15)'],
        ['Java', 'OpenJDK 17.0.2'],
    ])

    add_heading(doc, '2. Environment Setup')
    add_heading(doc, '2.1 Prerequisites', level=2)
    add_bullet(doc, 'Flutter SDK 3.35.4')
    add_bullet(doc, 'Dart SDK 3.9.2')
    add_bullet(doc, 'Android Studio with SDK 35')
    add_bullet(doc, 'Java 17')
    add_bullet(doc, 'Firebase project configured')
    add_bullet(doc, 'RevGlue account with RevEmbed enabled (Publisher ID 1202)')

    add_heading(doc, '2.2 Firebase Configuration', level=2)
    add_bullet(doc, 'Place google-services.json in android/app/')
    add_bullet(doc, 'Ensure firebase_options.dart has Web and Android configs')
    add_bullet(doc, 'Firestore security rules configured for development/production')

    add_heading(doc, '2.3 RevGlue Configuration', level=2)
    add_para(doc, 'No server-side configuration required. The Publisher ID (1202) is embedded in the RevGlueService class. The RevGlue API endpoints are public REST APIs that return JSON data based on the Publisher ID.')

    add_heading(doc, '3. Build Commands')
    add_heading(doc, '3.1 Development (Web Preview)', level=2)
    add_para(doc, 'flutter pub get\nflutter build web --release\npython3 -m http.server 5060 --directory build/web --bind 0.0.0.0', size=10)
    
    add_heading(doc, '3.2 Release (Android APK)', level=2)
    add_para(doc, 'flutter pub get\nflutter build apk --release', size=10)
    
    add_heading(doc, '3.3 Release (Android App Bundle)', level=2)
    add_para(doc, 'flutter pub get\nflutter build appbundle --release', size=10)

    add_heading(doc, '4. Dependencies')
    add_table(doc, ['Package', 'Version', 'Purpose'], [
        ['flutter', '3.35.4', 'Core framework'],
        ['provider', '6.1.5+1', 'State management'],
        ['firebase_core', '3.6.0', 'Firebase core'],
        ['firebase_auth', '5.3.1', 'Authentication'],
        ['cloud_firestore', '5.4.3', 'Cloud database'],
        ['firebase_storage', '12.3.2', 'File storage'],
        ['firebase_messaging', '15.1.3', 'Push notifications'],
        ['http', '1.5.0', 'HTTP client (RevGlue API)'],
        ['url_launcher', '6.2.5', 'Open affiliate links in browser'],
        ['google_fonts', '6.2.1+', 'Typography'],
        ['shared_preferences', '2.5.3', 'Local key-value storage'],
        ['in_app_purchase', '3.2.0', 'Subscription payments'],
        ['cached_network_image', '3.4.1+', 'Image caching'],
        ['image_picker', '1.1.2+', 'Photo capture'],
        ['intl', '0.20.2+', 'Date formatting'],
    ])

    add_heading(doc, '5. GitHub Repository')
    add_table(doc, ['Setting', 'Value'], [
        ['Repository', 'https://github.com/cdcruz007/huddl'],
        ['Branch', 'main'],
        ['CI/CD', 'GitHub Actions (planned)'],
    ])

    doc.save(f'{OUTPUT_DIR}/Huddl_Connect_Deployment_Guide.docx')
    print('  Created: Deployment Guide')


# ═══════════════════════════════════════════════════════════════════
# 7. REVGLUE INTEGRATION & MONETISATION DOCUMENT (NEW)
# ═══════════════════════════════════════════════════════════════════
def create_revglue_integration():
    doc = Document()
    
    title = doc.add_heading('Huddl Connect', level=0)
    for run in title.runs:
        run.font.color.rgb = BRAND_ORANGE
        run.font.size = Pt(28)
    add_para(doc, 'RevGlue Integration & Monetisation Strategy', bold=True, size=16)
    add_para(doc, f'Version 1.0 | {TODAY}', italic=True, size=11)
    doc.add_page_break()

    add_heading(doc, '1. Executive Summary')
    add_para(doc, 'Huddl Connect integrates RevGlue\'s RevEmbed Open API to create a passive affiliate revenue stream. The Deals tab surfaces coupons, voucher codes, and daily deals from 500+ UK e-commerce stores. Every purchase made through Huddl\'s affiliate links earns 80% commission (RevGlue retains 20%). This document covers the technical integration, business model, revenue projections, and operational procedures.')

    add_heading(doc, '2. Business Model')
    add_heading(doc, '2.1 Revenue Share', level=2)
    add_table(doc, ['Party', 'Share', 'Responsibility'], [
        ['Huddl Connect', '80%', 'Surface deals to users via app'],
        ['RevGlue', '20%', 'Maintain retailer relationships, tracking, payouts'],
    ])

    add_heading(doc, '2.2 How Commission Works', level=2)
    add_para(doc, '1. User browses deals in the Huddl app')
    add_para(doc, '2. User taps "Get Offer" or "Use Code & Shop"')
    add_para(doc, '3. App opens RevGlue exit-click URL in device browser')
    add_para(doc, '4. User is redirected to retailer website with affiliate cookie')
    add_para(doc, '5. User completes purchase on retailer site')
    add_para(doc, '6. Affiliate network records the sale and commission')
    add_para(doc, '7. Commission appears in RevGlue dashboard (pending)')
    add_para(doc, '8. After validation period (30-90 days), commission confirmed')
    add_para(doc, '9. Payout when balance reaches GBP100')

    add_heading(doc, '2.3 Competitive Advantage', level=2)
    add_para(doc, 'Unlike generic coupon apps, Huddl surfaces family-relevant deals to an engaged community of parents. This targeted audience has higher conversion rates for baby, children\'s, and family products.')
    add_table(doc, ['Competitor', 'Model', 'Huddl Advantage'], [
        ['Honey (PayPal)', 'Browser extension', 'In-app native experience, community trust'],
        ['VoucherCodes', 'Website', 'Mobile-first, targeted to parents'],
        ['TopCashback', 'Cashback website', 'Integrated into community app, not standalone'],
        ['Peanut Plus', 'GBP8.99/mo subscription', 'Lower price + deals revenue offset'],
    ])

    add_heading(doc, '3. Technical Integration')
    add_heading(doc, '3.1 Account Configuration', level=2)
    add_table(doc, ['Parameter', 'Value'], [
        ['Publisher ID', '1202'],
        ['Account Email', 'Conrad.au@gmail.com'],
        ['Project Type', 'Coupons and Daily Deals UK'],
        ['Registered Domain', 'Www.huddlparents.com'],
        ['API Base URL', 'https://www.revglue.com/iframeapi'],
        ['Sitemap URL', 'https://www.revglue.com/revembed/sitemap/1202'],
    ])

    add_heading(doc, '3.2 API Endpoints', level=2)
    add_table(doc, ['Endpoint', 'Method', 'Description', 'Response'], [
        ['top_stores/1202', 'GET', 'Popular stores with offer counts', 'Array of store objects'],
        ['coupon_allcategories/1202', 'GET', 'All categories with subcategories', 'Nested category array'],
        ['coupon_menu_categories/1202', 'GET', 'Navigation menu categories', 'Category array with sub_cat'],
        ['store_detail/1202/{id}/{page}', 'GET', 'Coupons for a store', 'Coupon object or array'],
        ['homepage_placement_banners/1202', 'GET', 'Featured banners', 'Object with banner entries'],
        ['coupon_exitclick/1202/{storeId}', 'Redirect', 'Affiliate tracking redirect', 'Redirects to retailer'],
    ])

    add_heading(doc, '3.3 Data Models', level=2)
    add_para(doc, 'RevGlueStore:', bold=True)
    add_para(doc, '{ id, title, title_url, offercoupon_str, store_icon, store_large_icon }', size=10)
    add_para(doc, 'RevGlueCoupon:', bold=True)
    add_para(doc, '{ id, store_id, store_title, voucher_code, voucher_title, store_icon, expiry_date, offer_coupon, title }', size=10)
    add_para(doc, 'RevGlueCategory:', bold=True)
    add_para(doc, '{ id, title, cat_url, offercoupon_str, small_icon, large_icon, sub[] }', size=10)

    add_heading(doc, '3.4 Caching Strategy', level=2)
    add_table(doc, ['Data Type', 'Cache TTL', 'Storage', 'Refresh'], [
        ['Top Stores', '30 minutes', 'In-memory', 'Pull-to-refresh or expiry'],
        ['Categories', '30 minutes', 'In-memory', 'Pull-to-refresh or expiry'],
        ['Store Coupons', 'None', 'N/A', 'Fresh on every store tap'],
        ['Banners', 'Session', 'In-memory', 'App restart'],
    ])

    add_heading(doc, '4. Subscription Integration')
    add_para(doc, 'Deals access is tiered to drive subscription upgrades:')
    add_table(doc, ['Feature', 'Explorer (Free)', 'Neighbourhood (GBP5.99/mo)', 'Inner Circle (GBP11.99/mo)'], [
        ['Browse deals', '10 stores/day', 'Unlimited', 'Unlimited'],
        ['Copy coupon codes', 'Yes', 'Yes', 'Yes'],
        ['Cashback tracking', 'No', 'Yes', 'Yes'],
        ['Exclusive/early deals', 'No', 'No', 'Yes'],
        ['Deal alerts (push)', 'No', 'Weekly digest', 'Real-time'],
        ['Saved coupons', '3', '20', 'Unlimited'],
    ])

    add_heading(doc, '5. Revenue Projections')
    add_table(doc, ['Metric', 'Month 1-3', 'Month 4-6', 'Month 7-12', 'Year 2'], [
        ['Monthly Active Users', '500', '2,000', '5,000', '20,000'],
        ['Deals Tab Users (%)', '30%', '40%', '50%', '60%'],
        ['Click-Through Rate', '5%', '7%', '8%', '10%'],
        ['Conversion Rate', '1.5%', '2%', '3%', '4%'],
        ['Avg Order Value', 'GBP25', 'GBP30', 'GBP35', 'GBP40'],
        ['Avg Commission Rate', '5%', '6%', '7%', '8%'],
        ['Monthly Revenue', 'GBP7', 'GBP67', 'GBP420', 'GBP3,840'],
        ['Annual Revenue', '-', '-', 'GBP2,520', 'GBP46,080'],
    ])

    add_heading(doc, '6. Operational Procedures')
    add_heading(doc, '6.1 Dashboard Access', level=2)
    add_para(doc, 'Login to revglue.com to monitor:')
    add_bullet(doc, 'Click counts and conversion rates')
    add_bullet(doc, 'Pending vs confirmed commissions')
    add_bullet(doc, 'Top-performing stores')
    add_bullet(doc, 'Payout history and balance')

    add_heading(doc, '6.2 Optimisation Actions', level=2)
    add_bullet(doc, 'Monthly: Review top-converting stores, feature them in "Popular Stores"')
    add_bullet(doc, 'Quarterly: Update family keyword list based on community trends')
    add_bullet(doc, 'Seasonal: Curate deals for school holidays, Christmas, summer')
    add_bullet(doc, 'Ongoing: Monitor RevGlue API changes and update service accordingly')

    add_heading(doc, '6.3 Risk Mitigation', level=2)
    add_table(doc, ['Risk', 'Mitigation'], [
        ['RevGlue API downtime', 'Cached data shown, graceful error UI with retry button'],
        ['Low conversion rates', 'A/B test deal placement, improve family curation'],
        ['Commission disputes', 'Document all exit-click URLs, monitor RevGlue dashboard'],
        ['API changes', 'Version-pinned integration, monitor RevGlue announcements'],
    ])

    doc.save(f'{OUTPUT_DIR}/Huddl_Connect_RevGlue_Integration_Guide.docx')
    print('  Created: RevGlue Integration & Monetisation Guide')


# ═══════════════════════════════════════════════════════════════════
# 8. TRIPS AI PITCH DECK SUPPLEMENT (Updated)
# ═══════════════════════════════════════════════════════════════════
def create_trips_pitch():
    doc = Document()
    
    title = doc.add_heading('Huddl Trips', level=0)
    for run in title.runs:
        run.font.color.rgb = BRAND_ORANGE
        run.font.size = Pt(28)
    add_para(doc, 'AI Family Travel Concierge - Pitch Deck Supplement', bold=True, size=16)
    add_para(doc, f'Version 2.1 | {TODAY}', italic=True, size=11)
    doc.add_page_break()

    add_heading(doc, 'THE INSIGHT')
    add_heading(doc, 'Why Travel is the Killer Feature for a Parenting Community', level=2)
    add_para(doc, 'Travel planning is the #1 stress point for families. Parents spend 10+ hours researching a single family holiday. Huddl Trips combines AI intelligence with community wisdom to cut that time to minutes.')

    add_heading(doc, 'HUDDL TRIPS')
    add_heading(doc, 'AI Family Travel Concierge - Feature Overview', level=2)
    add_bullet(doc, 'Conversational AI travel planner trained on family travel data')
    add_bullet(doc, 'Destination browser with family-friendly ratings')
    add_bullet(doc, 'AI-generated packing lists based on destination, child ages, and season')
    add_bullet(doc, 'Parents Abroad tips from community members')
    add_bullet(doc, 'Integration with Deals tab for travel-related savings')

    add_heading(doc, 'THE MOAT')
    add_heading(doc, '"Ask Parents Who\'ve Been" - Community-Powered AI Answers', level=2)
    add_para(doc, 'Unlike generic travel AI, Huddl Trips learns from real parent experiences. Community data enriches AI responses with practical, family-tested advice.')

    add_heading(doc, 'STRATEGIC VALUE')
    add_heading(doc, 'Revenue, Retention & Network Effects', level=2)
    add_table(doc, ['Metric', 'Impact'], [
        ['Subscription Conversion', 'Trips is Neighbourhood+ exclusive, driving upgrades'],
        ['Retention', 'Trip planning spans weeks, increasing daily active usage'],
        ['Affiliate Revenue', 'Travel deals in Deals tab (hotels, flights, luggage) earn commission'],
        ['Network Effects', 'Parent tips improve AI for all users'],
        ['Competitive Moat', 'No competitor combines community + AI + deals for family travel'],
    ])

    add_heading(doc, 'ROADMAP FIT')
    add_para(doc, 'Huddl Trips consolidates 4 existing AI features (Concierge, Packing Lists, Parents Abroad, Destination Discovery) into one premium product that drives both subscription revenue and affiliate deal conversions.')

    doc.save(f'{OUTPUT_DIR}/Huddl_Trips_AI_Feature_Pitch_Deck_Supplement.docx')
    print('  Created: Trips AI Pitch Deck Supplement')


# ═══════════════════════════════════════════════════════════════════
# Run all generators
# ═══════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print('Generating Huddl Connect design documents...\n')
    create_functional_design()
    create_software_design()
    create_ux_design()
    create_ai_system_design()
    create_store_publishing_guide()
    create_deployment_guide()
    create_revglue_integration()
    create_trips_pitch()
    print(f'\nAll documents saved to: {OUTPUT_DIR}/')
    print(f'Total: {len(os.listdir(OUTPUT_DIR))} files')
