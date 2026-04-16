#!/usr/bin/env python3
"""
Huddl Connect - Generate ALL updated documentation
Version: 3.0 | June 2026
- No RevGlue / Offers / Deals tab references
- 5-tab navigation: Home, Connect, Discover, Market, Profile
- Three tiers: Welcome (Free), Neighbour (£5.99/mo), Circle (£11.99/mo)
- AI roadmap moved to Appendix in pitch deck & GTM
- Founders: Malgorzata & Conrad
- Revenue streams: Subscriptions + B2B + Marketplace (no affiliate)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

TODAY = "June 2026"
VERSION = "Version 3.0 | June 2026"

# ─────────────────────────────────────────────────────────
# STYLE HELPERS
# ─────────────────────────────────────────────────────────

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p

def add_para(doc, text, bold=False, size=None, color=None, align=None):
    p = doc.add_paragraph(text)
    if align:
        p.alignment = align
    for run in p.runs:
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
    return table

def add_page_break(doc):
    doc.add_page_break()

TEAL = RGBColor(0x00, 0x96, 0x88)
DARK = RGBColor(0x21, 0x21, 0x21)
ACCENT = RGBColor(0x6C, 0x63, 0xFF)

# ─────────────────────────────────────────────────────────
# SHARED CONTENT BLOCKS
# ─────────────────────────────────────────────────────────

APP_OVERVIEW = """Huddl Connect is a hyperlocal parenting community app that keeps parents connected within their own borough — the "next door mum and dad" experience in digital form. Built on Flutter with a Firebase backend, it combines group messaging, local event discovery, a buy/sell/give marketplace, and an AI assistant into one mobile app. There are no national forums, no algorithm-driven strangers — only parents within your area."""

NAV_TABS = [
    ("Home", "Personalised AI-powered smart feed, community announcements, upcoming meetups, group activity, and borough highlights."),
    ("Connect", "Local parent groups (public & private), direct messaging, polls, message threads, saved messages, and AI Group Summaries."),
    ("Discover", "AI-discovered local events, parent-created meetups with RSVP, filters by area/date/age group, map view."),
    ("Market", "Buy, sell, and give away pre-loved baby and kids' items. AI Listing Writer auto-drafts descriptions. Category browse and in-app messaging to sellers."),
    ("Profile", "Edit profile, manage subscription, notification preferences, privacy settings, tutorial, help & legal."),
]

TIERS = [
    ("Welcome (Free)", "£0", "£0", [
        "Join up to 2 local parent groups",
        "Create 1 group",
        "Attend up to 2 meetups/month",
        "Message up to 5 parents directly",
        "Post up to 2 marketplace items",
        "Upload up to 3 photos",
        "Send up to 30 messages/month",
        "AI Chat Helper — 3 conversations/day",
        "AI Events Finder — 1 local event/week",
        "Personalised home feed — 2 refreshes/day",
        "Community Q&A — 3 questions/week",
        "10 bookmarks/month",
    ]),
    ("Neighbour", "£5.99/mo", "£49.99/yr", [
        "Join unlimited local parent groups",
        "Create up to 25 groups (inc. private, invite-only)",
        "Attend & organise unlimited meetups",
        "Unlimited direct messaging",
        "Post up to 15 marketplace items",
        "Upload up to 15 photos",
        "Send unlimited messages",
        "Neighbour profile badge",
        "AI Chat Helper — 25 conversations/day",
        "AI Group Summaries — 10/day",
        "AI Listing Writer — 10 drafts/month",
        "AI Events Finder — daily local event discovery",
        "Unlimited home feed refreshes",
        "Community Q&A — 15 questions/week",
        "AI-generated Q&A answer summaries",
        "50 bookmarks/month",
        "Community participation badges",
    ]),
    ("Circle", "£11.99/mo", "£99.99/yr", [
        "Everything in Neighbour — fully unlimited",
        "Create unlimited groups",
        "Post unlimited marketplace items",
        "Upload up to 50 photos",
        "Circle profile badge (exclusive)",
        "AI Chat Helper — unlimited/day",
        "AI Group Summaries — unlimited",
        "AI Listing Writer — unlimited",
        "AI Events Finder — unlimited daily discovery",
        "AI Meetup Matchmaker — AI-suggested meetups",
        "Unlimited home feed personalisation",
        "Unlimited Q&A questions & AI summaries",
        "Unlimited bookmarks",
    ]),
]

AI_FEATURES_BUILT = [
    ("AI Chat Helper (Copilot)", "Conversational AI parenting assistant powered by Google Gemini. Answers questions about local services, child development, parenting advice — all grounded in borough-specific context. Available to all tiers (quota varies)."),
    ("AI Group Summaries", "Catches up users on missed group chat conversations in one tap. Generates a plain-English summary of recent activity. Neighbour: 10/day. Circle: unlimited."),
    ("AI Listing Writer", "Auto-drafts marketplace listing descriptions from minimal user input. Reduces friction for selling pre-loved items. Neighbour: 10/month. Circle: unlimited."),
    ("AI Events Finder", "Daily discovery of local family-friendly events via web intelligence. Surfaces results filtered by borough, age group, and interests. Welcome: 1/week. Neighbour: daily. Circle: unlimited."),
    ("AI Smart Feed", "Invisible AI that reorders and personalises the Home feed based on activity patterns, group participation, and parenting stage. All tiers (refresh quotas vary)."),
    ("AI Meetup Matchmaker", "Suggests the most relevant local meetups based on the user's interests, location, and children's ages. Circle tier only."),
    ("AI Event Recommendations", "Personalised event suggestions beyond raw discovery — learns preferences over time. Neighbour and Circle tiers."),
    ("Borough AI Context Engine", "Proprietary knowledge layer that grounds all AI responses in hyperlocal data: local schools, parks, services, transport links, and community norms for the user's specific borough."),
    ("Invisible AI / Daily Refresh", "Background AI service that runs nightly to refresh content, pre-compute recommendations, and prime the smart feed. Transparent, non-intrusive."),
]

AI_ROADMAP = [
    ("Voice-to-text parenting journal", "Parent can record voice notes; AI transcribes and organises into milestone diary."),
    ("AI Group Moderator", "Flags inappropriate content, surfaces unanswered questions, recommends relevant threads to join."),
    ("Predictive Meetup Builder", "AI suggests time, venue, and format for new meetups based on group activity patterns."),
    ("AI Health & Development Tracker", "Tracks developmental milestones from chat interactions; surfaces NHS-aligned guidance."),
    ("B2B Partner AI", "Allows local businesses to push AI-curated content to relevant parent segments by borough."),
    ("Cross-Borough Intelligence", "Aggregated anonymised insights across boroughs to surface national parenting trends."),
]

REVENUE_STREAMS = [
    ("Subscriptions", "Neighbour (£5.99/mo | £49.99/yr) and Circle (£11.99/mo | £99.99/yr). Primary revenue stream. Default annual billing with save-30% badge."),
    ("B2B Partner Events", "Local businesses (nurseries, baby classes, GP practices, children's entertainment) pay to promote events and services to hyper-targeted parent audiences in their borough."),
    ("Marketplace Commission (Future)", "Future 5-10% transaction fee on Preloved marketplace sales once payment rails are integrated."),
    ("Data Insights (Future)", "Anonymised, GDPR-compliant borough-level parenting insights sold to local councils, NHS trusts, and family-focused brands."),
]

COMPETITIVE = [
    ("Mumsnet", "National forum. No local focus, no AI, no community groups, no marketplace. Overwhelming, anonymous.", "Borough-scoped, AI-assisted, community-first"),
    ("Peanut", "Social networking for mums. No groups, no events, no marketplace, no hyperlocal borough focus.", "Full group chat, events, marketplace"),
    ("Nextdoor", "General neighbourhood app. Not parent-specific, no parenting AI, no structured groups.", "Parent-only, AI suite, structured groups"),
    ("Facebook Groups", "Fragmented, algorithm-driven. No parenting-specific AI, no marketplace integration, privacy concerns.", "Private, GDPR-compliant, no algorithm noise"),
    ("Huckleberry", "Child development tracking only. No community, no events, no marketplace.", "Full community + AI + commerce"),
    ("Huddl Connect ✓", "Hyperlocal borough focus + AI suite + groups + events + marketplace in one app. Purpose-built for parents.", "—"),
]

# ─────────────────────────────────────────────────────────────────────────────
# 1. GO-TO-MARKET STRATEGY
# ─────────────────────────────────────────────────────────────────────────────

def create_gtm(path):
    doc = Document()

    # Cover
    doc.add_heading("huddl connect", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = doc.add_paragraph("Go-To-Market Strategy")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].bold = True; t.runs[0].font.size = Pt(16)

    t2 = doc.add_paragraph(f"{VERSION} | Confidential")
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("The hyperlocal parenting community — your borough's mums and dads, in one app.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_break(doc)

    # Executive Summary
    add_heading(doc, "1. Executive Summary")
    add_para(doc, APP_OVERVIEW)
    add_para(doc, "This Go-To-Market Strategy outlines how Huddl Connect will acquire, activate, and retain parents across UK boroughs, beginning with a Cambridge hyperlocal launch and expanding borough by borough. The strategy is built around three pillars: organic community growth, borough ambassador activation, and B2B partnerships with local family businesses.")
    add_page_break(doc)

    # Product Overview
    add_heading(doc, "2. What We've Built")
    add_para(doc, "The following features are live and available to users today:", bold=True)
    doc.add_heading("2.1 App Architecture — 5-Tab Navigation", level=2)
    for tab, desc in NAV_TABS:
        add_bullet(doc, f" {desc}", bold_prefix=f"{tab}: ")
    doc.add_heading("2.2 Built AI Features (Live)", level=2)
    for name, desc in AI_FEATURES_BUILT:
        add_bullet(doc, f" {desc}", bold_prefix=f"{name}: ")
    doc.add_heading("2.3 Subscription Tiers (Live)", level=2)
    for name, monthly, annual, features in TIERS:
        doc.add_heading(f"{name} — {monthly} | {annual}/yr", level=3)
        for f in features:
            add_bullet(doc, f)
    add_page_break(doc)

    # Target Market
    add_heading(doc, "3. Target Market")
    doc.add_heading("3.1 Primary Audience", level=2)
    add_para(doc, "Parents and expectant parents living within a specific UK borough who want genuine local connections — not national forums or anonymous strangers. The core user is a parent aged 25–45 with young children (0–12), smartphone-native, time-poor, and seeking local community support, services, and social connection.")
    doc.add_heading("3.2 User Personas", level=2)
    personas = [
        ("Expecting Parent — Sarah", "First-time mum, 30, Cambridge. Researching local NCT groups, baby classes, and pre-loved kit. Needs reassurance and local community."),
        ("New Parent — Marcus", "Dad of newborn, 32, Hackney. Sleep-deprived, wants quick answers and local meetup groups. High AI Chat Helper usage."),
        ("Experienced Parent — Priya", "Mum of 3, 38, Islington. Sells outgrown items, organises meetups, active in multiple school-run groups. Power user: Circle tier."),
        ("Local Business — Dr. Chen", "Runs postnatal classes in Camden. Wants to reach local parents directly. B2B partner event promotions."),
    ]
    for name, desc in personas:
        add_bullet(doc, f" {desc}", bold_prefix=f"{name}: ")
    doc.add_heading("3.3 Market Sizing", level=2)
    add_table(doc,
        ["Segment", "UK Size", "Addressable"],
        [
            ["Parents with children 0–12", "~7.5M households", "~3.5M smartphone-active"],
            ["Expectant parents/year", "~650,000", "~500,000"],
            ["London boroughs alone", "~900,000 parent HHs", "~550,000 primary target"],
            ["Cambridge beachhead", "~18,000 parent HHs", "~12,000 initial TAM"],
        ]
    )
    add_page_break(doc)

    # GTM Phases
    add_heading(doc, "4. Go-To-Market Phases")
    phases = [
        ("Phase 1 — Beachhead (Now)", "Cambridge", [
            "Hyperlocal launch in Cambridge: target 2,000 active users within 90 days",
            "Borough ambassador recruitment: 10 Cambridge mums/dads as community leads",
            "B2B outreach: Cambridge nurseries, baby classes, NCT groups, children's centres",
            "Organic WhatsApp/Facebook group seeding by ambassadors",
            "7-day free trial of Neighbour on every sign-up — no credit card required",
            "Soft paywall at 'aha moments': 3rd group join, 6th DM, AI limit reached",
        ]),
        ("Phase 2 — Greater Cambridgeshire (Q3 2026)", "Cambridgeshire", [
            "Expand to Peterborough, Ely, St Ives, Huntingdon using Cambridge playbook",
            "Local press coverage: Cambridge News, family blogs, parenting podcasts",
            "Paid social: Meta parent audience targeting by postcode",
            "B2B partner events: expand to 50 local businesses by end of phase",
        ]),
        ("Phase 3 — Major Cities (Q4 2026–Q1 2027)", "London, Oxford, Bristol", [
            "Unlock London borough by borough — Hackney, Islington, Wandsworth as first targets",
            "Oxford and Bristol using same ambassador-first model",
            "PR campaign: 'The next-door app for parents'",
            "App Store / Google Play featured category push",
        ]),
        ("Phase 4 — National Rollout (2027)", "UK-wide", [
            "Self-service borough activation: any community can request their borough",
            "National media push and parenting influencer partnerships",
            "B2B national account team for multi-borough businesses (nursery chains, NHS, councils)",
        ]),
    ]
    for phase, geo, bullets in phases:
        doc.add_heading(phase, level=2)
        add_para(doc, f"Geography: {geo}", bold=True)
        for b in bullets:
            add_bullet(doc, b)
    add_page_break(doc)

    # Acquisition Channels
    add_heading(doc, "5. Customer Acquisition Channels")
    channels = [
        ("Borough Ambassadors", "Recruit 5–10 local mums/dads per borough as paid/incentivised community leads. They seed groups, invite contacts, and drive organic word-of-mouth. Cost: low. Impact: high. Conversion: ~35%."),
        ("WhatsApp / Facebook Groups", "Target existing local parent WhatsApp and Facebook groups. Ambassadors post with genuine community value, not advertising. Leverages existing trust networks."),
        ("Paid Social (Meta/Instagram)", "Parent-demographic targeting by postcode radius. Focus on new parents (0–12 months post-birth) and NCT/antenatal community members. CPA target: £8–12."),
        ("GP Surgeries & Health Visitors", "Partner with NHS health visitors and midwife teams to distribute Huddl information at postnatal checks. High trust, direct access to new parents."),
        ("Nurseries & Children's Centres", "B2B partnerships: nurseries use the Events tab to promote their sessions. Staff recommend Huddl to parents. Mutual value exchange."),
        ("App Store Organic", "App Store Optimisation (ASO): 'local parenting app', 'parent community [borough]', 'mums and dads near me'. Aim for top 10 in Parenting category."),
        ("PR & Content", "Local press (Cambridge News, borough papers). Parenting blog partnerships. Guest articles on isolation, community, hyperlocal connection. No paid influencer budget required in Phase 1."),
    ]
    for name, desc in channels:
        add_bullet(doc, f" {desc}", bold_prefix=f"{name}: ")
    add_page_break(doc)

    # Revenue Model
    add_heading(doc, "6. Revenue Model")
    add_para(doc, "Revenue is generated through three streams (no affiliate/RevGlue dependency):", bold=True)
    for name, desc in REVENUE_STREAMS:
        add_bullet(doc, f" {desc}", bold_prefix=f"{name}: ")
    doc.add_heading("6.1 Subscription Revenue Projections", level=2)
    add_table(doc,
        ["Milestone", "Users", "Paying (%)", "ARPU/yr", "ARR"],
        [
            ["Cambridge Launch", "2,000", "12%", "£55", "~£13,200"],
            ["Cambridgeshire Phase 2", "8,000", "18%", "£58", "~£83,520"],
            ["Phase 3 (3 cities)", "35,000", "22%", "£60", "~£462,000"],
            ["Phase 4 (UK national)", "120,000", "25%", "£65", "~£1,950,000"],
        ]
    )
    doc.add_heading("6.2 Conversion Strategy", level=2)
    bullets = [
        "7-day auto-trial of Neighbour on every sign-up — no credit card required",
        "Soft paywalls at 'aha moments': 3rd group join, 6th DM, AI limit reached",
        "Default annual billing with 'Save 30%' badge prominently displayed",
        "Day-5 trial reminder push notification",
        "Day-7 exit survey with 1-month free pause offer on cancellation",
        "In-app upgrade prompts embedded naturally at feature limits",
    ]
    for b in bullets:
        add_bullet(doc, b)
    add_page_break(doc)

    # Competitive Positioning
    add_heading(doc, "7. Competitive Positioning")
    add_para(doc, "Huddl Connect occupies a unique white space: no competitor combines hyperlocal borough focus + parenting AI + community groups + events + marketplace in one app.")
    add_table(doc,
        ["Competitor", "Key Gap", "Huddl Advantage"],
        [
            ["Mumsnet", "National, anonymous, no AI, no local", "Borough-scoped, AI-assisted, community-first"],
            ["Peanut", "No groups, no events, no local borough", "Full group chat, events, marketplace"],
            ["Nextdoor", "Generic neighbourhood, not parent-specific", "Parent-only, AI suite, structured groups"],
            ["Facebook Groups", "Algorithm-driven, privacy concerns", "Private, GDPR-compliant, no algorithm noise"],
            ["Huckleberry", "Dev tracking only, no community", "Full community + AI + commerce"],
        ]
    )
    add_page_break(doc)

    # KPIs
    add_heading(doc, "8. Key Performance Indicators")
    add_table(doc,
        ["KPI", "Target (90 days)", "Target (12 months)"],
        [
            ["Registered users (Cambridge)", "2,000", "8,000"],
            ["Monthly active users (MAU)", "65%+ of registered", "65%+ of registered"],
            ["7-day trial activation rate", ">70%", ">70%"],
            ["Trial-to-paid conversion", ">12%", ">18%"],
            ["Group creation rate", ">30% of MAU", ">40% of MAU"],
            ["AI Chat Helper usage", ">50% of MAU/week", ">60% of MAU/week"],
            ["B2B partner events live", "15", "50"],
            ["Net Promoter Score (NPS)", ">50", ">60"],
        ]
    )
    add_page_break(doc)

    # Team
    add_heading(doc, "9. Team")
    add_para(doc, "Huddl Connect is founded by parents who experienced the exact fragmentation and isolation the platform solves — in their own borough.", bold=False)
    add_table(doc,
        ["Name", "Title", "Background"],
        [
            ["Conrad", "Founder", "Product, engineering, and go-to-market. Built the entire Flutter app. Former tech lead, Cambridge-based parent."],
            ["Malgorzata", "Founder", "Community strategy, UX, and parent-audience insight. Deep roots in local Cambridge parent networks."],
        ]
    )
    add_page_break(doc)

    # ── APPENDIX: AI Roadmap ─────────────────────────────────────────────────
    add_heading(doc, "APPENDIX A — AI Roadmap (Future Features)")
    add_para(doc, "The following AI capabilities are on the product roadmap but are NOT yet built or available in the current app. They are included here for strategic context only.", bold=True)
    for name, desc in AI_ROADMAP:
        add_bullet(doc, f" {desc}", bold_prefix=f"{name}: ")

    add_heading(doc, "APPENDIX B — Document Control")
    add_table(doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ["1.0", "March 2025", "Conrad", "Initial draft"],
            ["2.0", "April 2026", "Conrad", "Post-pilot update, RevGlue integration"],
            ["3.0", "June 2026", "Conrad / Malgorzata", "Removed RevGlue/Deals tab, updated tiers, AI roadmap to appendix, 5-tab navigation"],
        ]
    )

    doc.save(path)
    print(f"✅ GTM saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# 2. INVESTOR PITCH DECK (as .docx)
# ─────────────────────────────────────────────────────────────────────────────

def create_pitch_deck(path):
    doc = Document()

    # Cover Slide
    doc.add_heading("huddl connect", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph("Investor Pitch Deck")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].bold = True; s.runs[0].font.size = Pt(18)
    doc.add_paragraph("The hyperlocal parenting community app").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\"Your borough's mums and dads — in one app.\"").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"{VERSION} | Confidential").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_break(doc)

    # Slide 1 — The Problem
    add_heading(doc, "SLIDE 1 — THE PROBLEM")
    doc.add_heading("Modern Parenting is Fragmented, Isolating & Overwhelming", level=2)
    add_para(doc, "Parents juggle 5+ apps to manage their daily community lives. None of them talk to each other. None of them are local. The result: cognitive overload, social isolation, and missed community moments.")
    problems = [
        "WhatsApp for group chat — but groups are fragmented and unmanaged",
        "Facebook Groups for local community — but full of spam, no structure, privacy concerns",
        "Eventbrite for events — but not parent-specific, no local filtering",
        "Gumtree/Facebook Marketplace for baby items — no dedicated parent experience",
        "Mumsnet for advice — national, anonymous, overwhelming",
        "No single app that does all of this. No app that keeps it local.",
    ]
    for p in problems:
        add_bullet(doc, p)
    add_page_break(doc)

    # Slide 2 — The Solution
    add_heading(doc, "SLIDE 2 — THE SOLUTION")
    doc.add_heading("Huddl Connect: Your Borough's Parent Community — One App", level=2)
    add_para(doc, APP_OVERVIEW)
    doc.add_heading("5-Tab Architecture (Live Today)", level=3)
    for tab, desc in NAV_TABS:
        add_bullet(doc, f" {desc}", bold_prefix=f"{tab}: ")
    add_page_break(doc)

    # Slide 3 — Market Opportunity
    add_heading(doc, "SLIDE 3 — MARKET OPPORTUNITY")
    doc.add_heading("UK Parenting & Family Tech Market", level=2)
    add_table(doc,
        ["Market", "Size"],
        [
            ["UK parents with children 0–12", "~7.5M households"],
            ["Annual new parents", "~650,000/year"],
            ["London parent households", "~900,000 (primary expansion)"],
            ["Cambridge beachhead TAM", "~12,000 parent households"],
            ["UK Family App spend", "£180M+ annually"],
        ]
    )
    doc.add_heading("Beachhead Strategy", level=3)
    add_para(doc, "Cambridge is the ideal launch borough: high density of young families, affluent early-adopter demographic, active local parent communities, strong word-of-mouth networks. Every Cambridge parent knows another. The playbook replicates to any UK borough.")
    add_page_break(doc)

    # Slide 4 — Built Product (THE SPOTLIGHT)
    add_heading(doc, "SLIDE 4 — WHAT WE'VE BUILT (THE SPOTLIGHT)")
    doc.add_heading("Live App — Built, Tested, Ready to Scale", level=2)
    add_para(doc, "This is not a concept. Huddl Connect is a fully functioning Flutter application with a Firebase backend, available on Android, in TestFlight for iOS, and accessible via web preview.", bold=True)

    doc.add_heading("Core Social Features (Live)", level=3)
    core_features = [
        "Local parent groups — auto-assigned by postcode and parenting stage on sign-up",
        "Public and private groups — any parent can create, join, and manage",
        "Group chat with message threads, polls, saved messages, and forward messaging",
        "Direct messaging (DMs) between any two parents in the community",
        "Marketplace — buy, sell, give away pre-loved baby and kids' items",
        "Events & Meetups — discover local events, create and RSVP to parent meetups",
        "Personalised onboarding — postcode, parenting stage, child ages, profile photo",
        "In-app tutorial overlay that walks new users through all 5 tabs",
        "Push notifications for group messages, DMs, and event reminders",
        "Firebase phone authentication with OTP verification (Android & iOS)",
        "Community Q&A board — post questions, get answers from local parents",
        "Bookmark messages and posts for later",
    ]
    for f in core_features:
        add_bullet(doc, f)

    doc.add_heading("AI Features (Live Today)", level=3)
    for name, desc in AI_FEATURES_BUILT:
        add_bullet(doc, f" {desc}", bold_prefix=f"{name}: ")
    add_page_break(doc)

    # Slide 5 — AI Suite
    add_heading(doc, "SLIDE 5 — AI SUITE")
    doc.add_heading("Purpose-Built Parent AI — Not Generic Chatbots", level=2)
    add_para(doc, "Every AI feature in Huddl is grounded in borough-specific context. The AI knows the parent's postcode, children's ages, local community, and activity patterns. It gives hyperlocal answers, not national generalities.")
    add_table(doc,
        ["AI Feature", "What It Does", "Available From"],
        [
            ["AI Chat Helper", "Conversational parenting AI — local, contextual, empathetic", "Welcome (3/day)"],
            ["AI Smart Feed", "Personalises home feed invisibly — surfaces most relevant content", "Welcome (2 refreshes/day)"],
            ["AI Events Finder", "Discovers local family events daily via web intelligence", "Welcome (1/week)"],
            ["AI Group Summaries", "Catches up users on missed chat in one tap", "Neighbour (10/day)"],
            ["AI Listing Writer", "Auto-drafts marketplace descriptions from minimal input", "Neighbour (10/mo)"],
            ["AI Event Recommendations", "Personalised event suggestions based on learned preferences", "Neighbour"],
            ["AI Meetup Matchmaker", "AI-suggests best local meetups for the user's profile", "Circle only"],
            ["Borough AI Context", "Grounds all AI in hyperlocal borough data and community norms", "All tiers"],
            ["Daily AI Refresh", "Background AI that refreshes content and recommendations nightly", "All tiers"],
        ]
    )
    add_para(doc, "THE FLYWHEEL: Every message, group join, and meetup attendance makes the AI smarter. Every AI interaction creates more engagement. The data moat grows with every user.", bold=True)
    add_page_break(doc)

    # Slide 6 — Subscription Model
    add_heading(doc, "SLIDE 6 — SUBSCRIPTION MODEL")
    doc.add_heading("Three-Tier Freemium: Welcome | Neighbour | Circle", level=2)
    add_para(doc, "Value-first freemium designed to convert within 7 days. Default annual billing. No founding member pricing — clean, scalable tiers.")
    add_table(doc,
        ["Tier", "Monthly", "Annual", "Key Benefits"],
        [
            ["Welcome (Free)", "£0", "£0", "2 groups, 2 meetups/mo, 5 DMs, AI Chat (3/day), basic marketplace"],
            ["Neighbour", "£5.99", "£49.99", "Unlimited groups/DMs/meetups, full AI suite, 15 listings, community badges"],
            ["Circle", "£11.99", "£99.99", "Everything unlimited + AI Meetup Matchmaker, 50 photos, Circle badge"],
        ]
    )
    doc.add_heading("Conversion Mechanics (Built & Live)", level=3)
    for b in [
        "7-day auto-trial of Neighbour on every sign-up — no credit card required",
        "Soft paywalls at 'aha moments': 3rd group join, 6th DM, AI limit hit",
        "Default annual billing with 'Save 30%' badge",
        "Day-5 trial reminder push notification",
        "Day-7 exit survey with 1-month free pause on cancellation",
    ]:
        add_bullet(doc, b)
    add_page_break(doc)

    # Slide 7 — Revenue Model
    add_heading(doc, "SLIDE 7 — REVENUE MODEL")
    doc.add_heading("Three Revenue Streams: Subscriptions + B2B + Marketplace", level=2)
    for name, desc in REVENUE_STREAMS:
        add_bullet(doc, f" {desc}", bold_prefix=f"{name}: ")
    doc.add_heading("Revenue Projections", level=3)
    add_table(doc,
        ["Stage", "Users", "Paid %", "ARR"],
        [
            ["Cambridge (live)", "2,000", "12%", "~£13,200"],
            ["Cambridgeshire (Q3 2026)", "8,000", "18%", "~£83,520"],
            ["3-City (Q4 2026)", "35,000", "22%", "~£462,000"],
            ["National (2027)", "120,000", "25%", "~£1.95M"],
        ]
    )
    add_page_break(doc)

    # Slide 8 — Traction & Validation
    add_heading(doc, "SLIDE 8 — TRACTION & VALIDATION")
    doc.add_heading("Proof Points from Cambridge Community", level=2)
    traction = [
        "700+ parents organically engaged in WhatsApp pilot groups — no paid acquisition",
        "Core app built, tested, and in TestFlight for iOS / available on Android",
        "Full AI suite implemented and tested with real parent users",
        "Firebase backend live with Firestore, Auth, Crashlytics, and Storage",
        "5-tab navigation, onboarding flow, subscription management — all built",
        "Phone verification (OTP) working on both Android and iOS (with Firebase test numbers)",
        "B2B event framework ready for nurseries and local businesses",
        "Community Q&A, polls, message threads, saved messages — all live",
    ]
    for t in traction:
        add_bullet(doc, t)
    add_page_break(doc)

    # Slide 9 — Competitive Landscape
    add_heading(doc, "SLIDE 9 — COMPETITIVE LANDSCAPE")
    doc.add_heading("No Single Competitor Covers Community + AI + Marketplace + Events", level=2)
    add_table(doc,
        ["Competitor", "Gap vs Huddl"],
        [[c, g] for c, g, a in COMPETITIVE[:-1]]
    )
    add_para(doc, "Huddl Connect is the only app that combines: hyperlocal borough focus + parenting-specific AI + group chat + events + marketplace.", bold=True)
    add_page_break(doc)

    # Slide 10 — Go-To-Market
    add_heading(doc, "SLIDE 10 — GO-TO-MARKET STRATEGY")
    doc.add_heading("Hyperlocal Launch & Expand Playbook", level=2)
    phases_summary = [
        ("Phase 1 — Cambridge (Now)", "Ambassador seeding, B2B nursery partnerships, organic WhatsApp migration, 2,000 users in 90 days"),
        ("Phase 2 — Greater Cambridgeshire (Q3 2026)", "Replicate Cambridge playbook: Peterborough, Ely, St Ives, Huntingdon"),
        ("Phase 3 — London & Major Cities (Q4 2026)", "Hackney, Islington, Wandsworth; Oxford; Bristol — paid social + PR campaign"),
        ("Phase 4 — National (2027)", "Self-service borough activation, national press, B2B national accounts"),
    ]
    for phase, desc in phases_summary:
        add_bullet(doc, f" {desc}", bold_prefix=f"{phase}: ")
    add_page_break(doc)

    # Slide 11 — Technology
    add_heading(doc, "SLIDE 11 — TECHNOLOGY STACK")
    doc.add_heading("Built to Scale — Cross-Platform, Cloud-Native", level=2)
    add_table(doc,
        ["Layer", "Technology"],
        [
            ["Mobile App", "Flutter 3.35.4 (Dart 3.9.2) — Android primary, iOS (TestFlight), Web preview"],
            ["Backend / Auth", "Firebase (Firestore, Auth, Storage, Crashlytics, Analytics)"],
            ["AI Engine", "Google Gemini API — grounded in borough-specific context"],
            ["State Management", "Provider pattern + service singletons"],
            ["Push Notifications", "Firebase Cloud Messaging (FCM)"],
            ["Payments", "Google Play Billing (Android), Apple IAP (iOS) — subscription management"],
            ["CI/CD", "GitHub (cdcruz007/huddl) — version controlled, documented"],
        ]
    )
    add_page_break(doc)

    # Slide 12 — Product Roadmap (condensed — full roadmap in Appendix)
    add_heading(doc, "SLIDE 12 — NEAR-TERM ROADMAP (2026)")
    doc.add_heading("Next 6 Months — Consolidate & Scale", level=2)
    roadmap_6mo = [
        "iOS App Store submission and approval (TestFlight → Live)",
        "Google Play Store public launch (Android)",
        "B2B Partner Events portal — self-service business event promotion",
        "Borough expansion to Peterborough, Ely, and Huntingdon",
        "Group moderation tools for community leads",
        "Improved push notification personalisation",
        "App Store and Google Play rating campaign",
        "In-app referral programme: 'Invite 3 parents, get 1 month free'",
    ]
    for r in roadmap_6mo:
        add_bullet(doc, r)
    add_para(doc, "For the full AI future roadmap, see Appendix A.", bold=False)
    add_page_break(doc)

    # Slide 13 — Business Model Summary
    add_heading(doc, "SLIDE 13 — BUSINESS MODEL SUMMARY")
    add_table(doc,
        ["Dimension", "Detail"],
        [
            ["Revenue model", "SaaS subscriptions (monthly/annual) + B2B event promotions"],
            ["Pricing", "£5.99/mo or £49.99/yr (Neighbour); £11.99/mo or £99.99/yr (Circle)"],
            ["CAC target", "£8–12 per user (Phase 1 ambassador model: near-zero)"],
            ["LTV target", "£85–120 (2-year average paid subscription)"],
            ["LTV:CAC", ">8:1 at scale"],
            ["Gross margin", "~85% (SaaS model, low marginal cost per user)"],
            ["Payback period", "<6 months at Phase 2 scale"],
        ]
    )
    add_page_break(doc)

    # Slide 14 — The Ask
    add_heading(doc, "SLIDE 14 — THE ASK")
    doc.add_heading("Seed Investment to Scale Cambridge to UK National", level=2)
    add_para(doc, "We are raising seed investment to accelerate borough expansion, grow the ambassador network, and execute the B2B partnership programme.")
    ask_highlights = [
        "PROVEN DEMAND: 700+ parents organically engaged in WhatsApp pilot — zero paid acquisition",
        "LIVE PRODUCT: Full-featured app built and tested — not a concept, not an MVP skeleton",
        "FULL AI SUITE: 7 AI features live today, grounded in hyperlocal borough data",
        "CLEAR REVENUE MODEL: Subscriptions + B2B, no dependency on third-party affiliate deals",
        "SCALABLE PLAYBOOK: Cambridge model replicates to any UK borough — same playbook, new geography",
        "LOW CAC: Ambassador-led organic growth keeps acquisition cost <£12/user",
        "STRONG RETENTION DRIVER: Hyperlocal community creates genuine social lock-in",
        "GDPR-COMPLIANT: UK-built, children's safety code compliant, no data sold",
    ]
    for h in ask_highlights:
        add_bullet(doc, h)
    doc.add_heading("Use of Funds", level=3)
    add_table(doc,
        ["Category", "Allocation"],
        [
            ["Ambassador network (Cambridge + Phase 2)", "25%"],
            ["Paid social acquisition (Meta/Instagram)", "25%"],
            ["B2B sales and partnership team", "20%"],
            ["Product development (iOS, roadmap features)", "20%"],
            ["Operations, legal, marketing", "10%"],
        ]
    )
    add_page_break(doc)

    # Slide 15 — Team
    add_heading(doc, "SLIDE 15 — THE TEAM")
    doc.add_heading("Built by Parents Who Understand the Problem First-Hand", level=2)
    add_para(doc, "Huddl Connect is founded by parents in Cambridge who experienced the exact fragmentation and isolation the platform solves — in their own borough.")
    add_table(doc,
        ["Name", "Title", "Role"],
        [
            ["Conrad", "Founder", "Product & Engineering. Architected and built the entire Flutter application, Firebase backend, and AI integration. Cambridge-based parent."],
            ["Malgorzata", "Founder", "Community & Strategy. Drives parent audience insight, community design, and partner relationships. Deep Cambridge parent network."],
        ]
    )
    add_page_break(doc)

    # Closing Slide
    add_heading(doc, "CLOSING")
    doc.add_paragraph("huddl connect").runs[0].bold = True
    add_para(doc, '"Everything a parent needs. One app. One community. One borough."')
    add_para(doc, "Community + AI + Marketplace + Events")
    add_para(doc, "No competitor has all four. Only Huddl keeps it local.")
    add_para(doc, "Contact: conrad.au@gmail.com | github.com/cdcruz007/huddl")
    add_para(doc, "(c) 2026 Huddl Connect. All rights reserved.")
    add_page_break(doc)

    # ── APPENDIX A: AI Roadmap ───────────────────────────────────────────────
    add_heading(doc, "APPENDIX A — AI ROADMAP (Future Features, Not Yet Built)")
    add_para(doc, "The following AI capabilities are on the product roadmap. They are NOT available in the current app. They represent the strategic AI direction post-seed funding.", bold=True)
    for name, desc in AI_ROADMAP:
        add_bullet(doc, f" {desc}", bold_prefix=f"{name}: ")

    add_heading(doc, "APPENDIX B — Full Subscription Feature Matrix")
    add_table(doc,
        ["Feature", "Welcome", "Neighbour", "Circle"],
        [
            ["Groups joined", "2", "Unlimited", "Unlimited"],
            ["Groups created", "1", "25", "Unlimited"],
            ["Meetups/month", "2", "Unlimited", "Unlimited"],
            ["DM conversations", "5", "Unlimited", "Unlimited"],
            ["Messages/month", "30", "Unlimited", "Unlimited"],
            ["Marketplace listings", "2", "15", "Unlimited"],
            ["Photo uploads", "3", "15", "50"],
            ["Private groups", "No", "Yes", "Yes"],
            ["Profile badge", "No", "Neighbour badge", "Circle badge"],
            ["AI Chat Helper", "3/day", "25/day", "Unlimited"],
            ["AI Group Summaries", "No", "10/day", "Unlimited"],
            ["AI Listing Writer", "No", "10/month", "Unlimited"],
            ["AI Events Finder", "1/week", "Daily", "Unlimited"],
            ["AI Smart Feed", "2 refreshes/day", "Unlimited", "Unlimited"],
            ["AI Event Recommendations", "No", "Yes", "Yes"],
            ["AI Meetup Matchmaker", "No", "No", "Yes"],
            ["Community Q&A questions", "3/week", "15/week", "Unlimited"],
            ["AI Q&A Summaries", "No", "Yes", "Yes"],
            ["Bookmarks", "10/month", "50/month", "Unlimited"],
            ["Community badges", "No", "Yes", "Yes"],
        ]
    )

    add_heading(doc, "APPENDIX C — Document Control")
    add_table(doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ["1.0", "March 2025", "Conrad", "Initial pitch deck"],
            ["2.0", "April 2026", "Conrad", "Post-pilot update with RevGlue/Deals/Trips"],
            ["3.0", "June 2026", "Conrad / Malgorzata", "Removed RevGlue/Deals/Trips, updated tiers, AI roadmap to appendix, founders titles updated"],
        ]
    )

    doc.save(path)
    print(f"✅ Pitch Deck saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# 3. AI SYSTEM DESIGN DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────

def create_ai_sdd(path):
    doc = Document()
    doc.add_heading("Huddl Connect", 0)
    doc.add_paragraph("AI System Design Document")
    doc.add_paragraph(VERSION)
    add_page_break(doc)

    add_heading(doc, "Document Control")
    add_table(doc, ["Version", "Date", "Author", "Changes"],
        [["3.0", "June 2026", "Conrad / Malgorzata", "Removed RevGlue/Deals AI, updated feature inventory, AI roadmap to appendix"]])

    add_heading(doc, "1. Executive Summary")
    add_para(doc, "Huddl Connect embeds AI throughout the user experience to create a uniquely intelligent parenting community. All AI features are grounded in hyperlocal borough context — the app knows the user's postcode, parenting stage, children's ages, and local community. This document covers all live AI features, model architecture, data flow, usage limits by subscription tier, and privacy compliance.")

    add_heading(doc, "2. AI Feature Inventory (Live)")
    add_table(doc,
        ["Feature", "Status", "Model/Engine", "Description"],
        [
            ["AI Chat Helper (Copilot)", "LIVE", "Google Gemini API", "Conversational parenting AI. Answers questions grounded in borough context."],
            ["AI Smart Feed", "LIVE", "On-device scoring + Gemini", "Personalises Home feed. Invisible, non-intrusive reordering."],
            ["AI Events Finder", "LIVE", "Gemini + web discovery", "Daily local event discovery. Borough and age-group filtered."],
            ["AI Group Summaries", "LIVE", "Gemini summarisation", "Catches up users on missed chat. Plain-English group summaries."],
            ["AI Listing Writer", "LIVE", "Gemini text generation", "Auto-drafts marketplace listing descriptions from minimal input."],
            ["AI Event Recommendations", "LIVE", "Collaborative filtering", "Personalised event suggestions based on learned preferences."],
            ["AI Meetup Matchmaker", "LIVE", "Gemini + user profile matching", "Suggests best local meetups for user's profile. Circle tier only."],
            ["Borough AI Context Engine", "LIVE", "GeminiSystemPromptBuilder", "Injects hyperlocal borough data into all AI prompts."],
            ["Daily AI Refresh", "LIVE", "Scheduled Gemini calls", "Nightly background refresh. Pre-computes feed and recommendations."],
            ["Invisible AI Learning Engine", "LIVE", "Activity tracking + scoring", "Learns user preferences silently from app interactions."],
        ]
    )

    add_heading(doc, "3. AI Architecture")
    doc.add_heading("3.1 Borough Context Engine", level=2)
    add_para(doc, "The GeminiSystemPromptBuilder assembles a rich system prompt for every AI call. It injects:")
    for item in [
        "User's borough and postcode (from onboarding)",
        "Parenting stage (Aspiring, Expecting, New Parent, Experienced Parent)",
        "Children's ages and developmental stage",
        "Local borough directory: schools, parks, children's centres, NHS services",
        "Community activity: recent group messages, meetup attendance, marketplace listings",
        "HyperlocalRules: borough-specific context constraints and tone guidelines",
        "Empathy and safety guardrails: UK safeguarding and children's code alignment",
        "AI knowledge base: curated parenting information grounded in NHS guidance",
    ]:
        add_bullet(doc, item)

    doc.add_heading("3.2 AI Copilot (Chat Helper) Flow", level=2)
    add_para(doc, "1. User sends message in AI Copilot screen.\n2. AiCopilotService constructs conversation history.\n3. GeminiSystemPromptBuilder assembles borough-grounded system prompt.\n4. Gemini API called via AiApiHelper (with retry and timeout handling).\n5. Response parsed, categorised (health, development, local services, etc.).\n6. CopilotActions generated if relevant (e.g., 'Browse local events', 'Find a meetup').\n7. Response displayed with source note and feedback mechanism.\n8. Usage counter incremented via SubscriptionService.")

    doc.add_heading("3.3 AI Smart Feed Flow", level=2)
    add_para(doc, "1. DailyAiRefreshService runs nightly to pre-compute feed scores.\n2. AiFeedService scores each content item (group activity, event relevance, marketplace match).\n3. Home screen applies InvisibleAiService reordering on load.\n4. Feed presented with subtle 'AI-suggested' labels where AI influenced ranking.\n5. User thumbs feedback captured and fed back to learning engine.")

    doc.add_heading("3.4 AI Events Finder Flow", level=2)
    add_para(doc, "1. AiEventDiscoveryService triggered daily (per-tier quota).\n2. Gemini API queried with borough, date range, and age-group filters.\n3. Discovered events enriched with local venue data from borough directory.\n4. Results stored in Firestore events collection.\n5. AiEventRecommenderService applies personalisation layer for Neighbour/Circle tiers.")

    add_heading(doc, "4. AI Usage Limits by Tier")
    add_table(doc,
        ["AI Feature", "Welcome (Free)", "Neighbour", "Circle"],
        [
            ["AI Chat Helper", "3 conversations/day", "25 conversations/day", "Unlimited"],
            ["AI Smart Feed refresh", "2/day", "Unlimited", "Unlimited"],
            ["AI Events Finder", "1 event/week", "Daily discovery", "Unlimited"],
            ["AI Group Summaries", "Not available", "10/day", "Unlimited"],
            ["AI Listing Writer", "Not available", "10/month", "Unlimited"],
            ["AI Event Recommendations", "Not available", "Included", "Included"],
            ["AI Meetup Matchmaker", "Not available", "Not available", "Unlimited"],
            ["Borough AI Context", "All tiers", "All tiers", "All tiers"],
            ["Daily AI Refresh", "All tiers", "All tiers", "All tiers"],
        ]
    )

    add_heading(doc, "5. Data Privacy & Ethics")
    privacy = [
        "All AI features comply with UK GDPR and the Children's Code",
        "User data is processed with borough-scoped isolation — no cross-borough data leakage",
        "BoroughScopeGuard enforces that all queries are scoped to the user's own borough",
        "GdprBoroughDataService manages data export and deletion requests",
        "No personal data is sent to third-party AI providers beyond what is necessary for the query",
        "AI responses are never stored against user profiles without explicit consent",
        "Conversation history is session-scoped and not retained server-side",
        "All AI model calls are logged to Firebase Crashlytics for error monitoring only",
    ]
    for p in privacy:
        add_bullet(doc, p)

    add_heading(doc, "6. AI Services File Reference")
    add_table(doc,
        ["Service File", "Purpose"],
        [
            ["ai_api_helper.dart", "Gemini API wrapper with retry logic and timeout handling"],
            ["ai_copilot_service.dart", "AI Chat Helper — conversation management and response parsing"],
            ["ai_chat_summariser_service.dart", "Group chat summarisation service"],
            ["ai_event_discovery_service.dart", "Daily local event discovery via Gemini"],
            ["ai_event_recommender_service.dart", "Personalised event recommendations"],
            ["ai_feed_service.dart", "Smart home feed scoring and reordering"],
            ["ai_knowledge_base_service.dart", "Curated parenting knowledge base for AI grounding"],
            ["ai_learning_engine_service.dart", "User preference learning from app interactions"],
            ["ai_listing_service.dart", "Marketplace listing description generation"],
            ["ai_matchmaker_service.dart", "AI Meetup Matchmaker — profile-based meetup suggestions"],
            ["borough_ai_context.dart", "Borough-specific AI context data"],
            ["borough_analytics_service.dart", "Borough usage analytics (anonymised)"],
            ["daily_ai_refresh_service.dart", "Nightly background AI refresh scheduler"],
            ["discover_ai_service.dart", "Discover tab AI integration"],
            ["gemini_system_prompt_builder.dart", "Assembles borough-grounded Gemini system prompts"],
            ["invisible_ai_service.dart", "Invisible feed AI — non-intrusive personalisation"],
            ["meetup_ai_service.dart", "AI-powered meetup creation assistance"],
            ["messages_ai_service.dart", "AI features within group messaging context"],
        ]
    )

    add_page_break(doc)
    add_heading(doc, "APPENDIX — AI Future Roadmap (Not Yet Built)")
    add_para(doc, "The following features are on the roadmap but are NOT live in the current app.", bold=True)
    for name, desc in AI_ROADMAP:
        add_bullet(doc, f" {desc}", bold_prefix=f"{name}: ")

    doc.save(path)
    print(f"✅ AI SDD saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# 4. UX DESIGN DOCUMENTATION
# ─────────────────────────────────────────────────────────────────────────────

def create_ux_doc(path):
    doc = Document()
    doc.add_heading("Huddl Connect", 0)
    doc.add_paragraph("UX Design Documentation")
    doc.add_paragraph(VERSION)
    add_page_break(doc)

    add_heading(doc, "Document Control")
    add_table(doc, ["Version", "Date", "Author", "Changes"],
        [["3.0", "June 2026", "Conrad / Malgorzata", "Removed Deals/RevGlue tab UX, updated to 5-tab navigation, aligned with current build"]])

    add_heading(doc, "1. Executive Summary")
    add_para(doc, "Huddl Connect is designed as a warm, approachable mobile experience for parents. The UX prioritises simplicity, trust, and the 'next door neighbour' feeling — as if you are talking to the parent next door, not a stranger on the internet. Every screen and interaction is optimised for time-poor parents who need quick access to community, local events, and secondhand items.")

    add_heading(doc, "2. Design Principles")
    principles = [
        ("Trust First", "Every interaction reinforces safety and community trust. Verified phone numbers, moderated content, and transparent data practices. Parents must feel safe sharing family information."),
        ("Reduce Cognitive Load", "Parents are busy. Surface the right content at the right time. AI-powered recommendations replace manual searching. 5-tab navigation means everything is reachable in 2 taps."),
        ("Keep It Local", "Every feature is scoped to the user's borough. No national noise. The 'next door mum and dad' feel is maintained throughout every screen."),
        ("Celebrate Milestones", "The parenting journey is full of milestones. The app acknowledges and celebrates them — from first DM to first meetup attended."),
        ("Invisible AI", "AI never interrupts or overwhelms. It works in the background, surfacing better content, summarising missed chat, and suggesting meetups — only when invited."),
    ]
    for name, desc in principles:
        doc.add_heading(name, level=2)
        add_para(doc, desc)

    add_heading(doc, "3. Information Architecture")
    doc.add_heading("3.1 Navigation Model — 5-Tab Bottom Bar", level=2)
    add_para(doc, "The bottom navigation uses a custom floating pill-shaped bar with 5 equally-spaced icons. Each tab has an outlined (inactive) and filled (active) icon state. The bar uses subtle shadow and rounded corners (28px radius) for a modern feel.")
    add_table(doc,
        ["Tab", "Icon", "Screens"],
        [
            ["Home", "House", "Home feed, announcements, upcoming meetups, group activity, borough badge"],
            ["Connect", "People/Chat", "Groups list, group chat, DM inbox, create group, polls, saved messages"],
            ["Discover", "Compass/Map", "Events list, meetup list, map view, create meetup, event/meetup detail"],
            ["Market", "Tag/Shop", "Marketplace listing grid, item detail, create listing, seller DM"],
            ["Profile", "Person", "Edit profile, subscription management, notifications, privacy, tutorial, legal"],
        ]
    )

    doc.add_heading("3.2 Complete Screen Inventory", level=2)
    add_table(doc,
        ["Screen", "Tab", "Key Elements"],
        [
            ["Splash Screen", "System", "Huddl logo, loading animation, Firebase init"],
            ["Onboarding Carousel", "System", "5-card swipeable intro to app features"],
            ["Phone Number Entry", "Auth", "Country picker (+44 UK default), phone input, Continue CTA"],
            ["OTP Verification", "Auth", "6-digit code input, resend timer (60s), loading states, retry flow"],
            ["Name Input", "Onboarding", "First name input, validation, progress bar"],
            ["Password Setup", "Onboarding", "Secure password, confirmation, strength indicator"],
            ["Parent Type", "Onboarding", "Aspiring / Expecting / New Parent / Experienced — tap to select"],
            ["Stage of Life", "Onboarding", "Subcategory selection within parent type"],
            ["About You", "Onboarding", "Bio text field, optional"],
            ["Child Info", "Onboarding", "Add child name and DOB — repeatable, add another CTA"],
            ["Due Date", "Onboarding", "Date picker for expecting parents"],
            ["Postcode Entry", "Onboarding", "UK postcode input, borough auto-detection via API"],
            ["Add Photo", "Onboarding", "Camera/gallery picker, circle crop, skip option"],
            ["Welcome Complete", "Onboarding", "Confirmation screen, auto-navigate to main app"],
            ["Home", "Home", "AI smart feed, announcements banner, upcoming meetups row, group activity cards, borough badge"],
            ["Groups", "Connect", "Public groups list, user's groups section, create group FAB"],
            ["Group Chat", "Connect", "Message list, input bar, attach media, polls, thread replies, AI Summary button"],
            ["DM Chat", "Connect", "1:1 message thread, read receipts, typing indicator"],
            ["Events", "Discover", "Tabbed: AI Events / Meetups. Map/List toggle. Age filter. Date filter."],
            ["Create Meetup", "Discover", "Title, description, date/time, location, max attendees"],
            ["Marketplace", "Market", "Grid of listing cards. Category filter bar. Search. Create listing FAB."],
            ["Item Detail", "Market", "Photo carousel, description, price, seller info, Message Seller CTA"],
            ["Create Listing", "Market", "Photos (up to limit), title, description (AI Writer button), price, category"],
            ["Profile", "Profile", "Avatar, name, parenting stage, bio, subscription badge, edit options"],
            ["Manage Subscription", "Profile", "Current tier, billing period, next renewal, upgrade/cancel options"],
            ["Subscription Plans", "Profile", "Three-tier comparison with monthly/annual toggle and purchase flow"],
            ["Not Available", "System", "Borough not yet launched — waitlist screen"],
        ]
    )

    add_heading(doc, "4. Onboarding UX Flow")
    add_para(doc, "The onboarding flow is a multi-step wizard that builds a rich user profile before the main app is shown. Each step is a full-screen card with a progress bar at the top.")
    steps = [
        ("Step 1", "Carousel introduction (swipeable, skippable)"),
        ("Step 2", "Phone number entry → OTP verification (Firebase Auth)"),
        ("Step 3", "Name + password setup"),
        ("Step 4", "Parent type selection (Aspiring / Expecting / New / Experienced)"),
        ("Step 5", "Stage of life sub-selection"),
        ("Step 6", "About you (optional bio)"),
        ("Step 7", "Child info (name + DOB — repeatable)"),
        ("Step 8", "Due date (expecting parents only)"),
        ("Step 9", "Postcode entry — borough auto-detected"),
        ("Step 10", "Profile photo (optional)"),
        ("Step 11", "Welcome complete → main app with auto-assigned groups"),
    ]
    for step, desc in steps:
        add_bullet(doc, f" {desc}", bold_prefix=f"{step}: ")

    add_heading(doc, "5. Tutorial Overlay")
    add_para(doc, "The in-app tutorial overlay walks new users through all 5 tabs with a swipeable card interface. It launches automatically on first login and can be re-triggered from Profile settings.")
    add_para(doc, "Each tutorial card: highlights the tab icon, provides a 2-sentence description, includes a visual illustration of the screen, and has Skip / Next / Done controls.")

    add_heading(doc, "6. Subscription Upgrade UX")
    add_para(doc, "Soft paywalls are embedded at 'aha moments' — natural points where users hit a feature limit:")
    paywall_moments = [
        "Joining a 3rd group when on Welcome tier",
        "Starting a 6th DM conversation on Welcome tier",
        "Reaching daily AI Chat Helper limit",
        "Attempting to create a meetup on Welcome tier (meetups require Neighbour)",
        "Attempting to create a private group on Welcome tier",
        "Attempting to create a 3rd marketplace listing on Welcome tier",
    ]
    for m in paywall_moments:
        add_bullet(doc, m)
    add_para(doc, "The upgrade prompt is a bottom sheet with: feature gate message, tier comparison (2-column), monthly/annual toggle, and a single primary CTA. It never blocks the user from completing their current action — it appears after the action is denied.")

    add_heading(doc, "7. Accessibility")
    for item in [
        "Minimum touch target: 44×44px for all interactive elements",
        "Colour contrast: WCAG 2.1 AA minimum (4.5:1 for text)",
        "Screen reader support: Semantic labels on all icons and buttons",
        "Font scaling: Respects system text size preferences",
        "Reduced motion: Animations respect system accessibility settings",
        "SafeArea: All screens wrapped in SafeArea to avoid system UI overlap",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. Colour & Typography")
    add_para(doc, "Primary: Teal (#009688) — community, warmth, trust")
    add_para(doc, "Accent: Indigo/Purple (#6C63FF) — AI features, premium")
    add_para(doc, "Background: White (#FFFFFF) + Light Grey (#F5F5F5)")
    add_para(doc, "Text: Near-black (#212121) + Medium grey (#757575)")
    add_para(doc, "Typography: Google Fonts — Poppins (headings), Inter (body)")
    add_para(doc, "Material Design 3 components throughout — Cards, FABs, Chips, Bottom Sheets")

    doc.save(path)
    print(f"✅ UX Doc saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# 5. FUNCTIONAL DESIGN DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────

def create_fdd(path):
    doc = Document()
    doc.add_heading("Huddl Connect", 0)
    doc.add_paragraph("Functional Design Document (FDD)")
    doc.add_paragraph(VERSION)
    doc.add_paragraph("Classification: Confidential")
    add_page_break(doc)

    add_heading(doc, "Document Control")
    add_table(doc, ["Version", "Date", "Author", "Changes"],
        [["3.0", "June 2026", "Conrad / Malgorzata", "Removed RevGlue/Deals/Trips tab. Updated navigation to 5 tabs. Updated tiers. AI roadmap to appendix."]])

    add_heading(doc, "1. Introduction & Purpose")
    add_para(doc, "Huddl Connect is a mobile-first community platform for parents in UK boroughs. It combines group messaging, local event discovery, a buy/sell/give marketplace, and a purpose-built AI assistant into one app. The platform's defining characteristic is hyperlocal scoping — every feature is borough-bounded to create the 'next door mum and dad' experience.")
    doc.add_heading("Intended Audience", level=2)
    for a in ["Development team (Flutter/Dart)", "Product owners and stakeholders", "QA engineers", "UX/UI designers", "Investor and business development teams"]:
        add_bullet(doc, a)
    doc.add_heading("Scope", level=2)
    add_para(doc, "This document covers all user-facing features, business rules, subscription tiers, and technical specifications for the current live build. Deals/RevGlue features and Trips tab are NOT part of this build.")

    add_heading(doc, "2. Product Overview")
    doc.add_heading("Core Value Proposition", level=2)
    add_para(doc, "A single app that replaces fragmented parenting tools: WhatsApp groups, Facebook Marketplace, Eventbrite, Mumsnet, and generic classifieds. Every feature is scoped to the user's borough — keeping community genuinely local.")
    doc.add_heading("Primary Navigation — 5-Tab Architecture", level=2)
    add_table(doc,
        ["Tab Index", "Tab Name", "Primary Screen", "Key Features"],
        [
            ["0", "Home", "HomeScreen", "AI smart feed, announcements, meetups, group activity, borough badge"],
            ["1", "Connect", "GroupsScreen", "Group chat, DMs, polls, threads, AI Summaries, saved messages"],
            ["2", "Discover", "EventsScreen", "AI events, parent meetups, map view, RSVP, create meetup"],
            ["3", "Market", "MarketplaceScreen", "Buy/sell/give listings, AI Listing Writer, item detail, seller DM"],
            ["4", "Profile", "ProfileScreen", "Edit profile, subscription, notifications, privacy, tutorial"],
        ]
    )
    doc.add_heading("Platform Targets", level=2)
    for p in ["Android (primary) — Google Play Store", "iOS — Apple App Store (TestFlight stage)", "Web (preview/testing)"]:
        add_bullet(doc, p)
    doc.add_heading("Technology Stack", level=2)
    add_table(doc,
        ["Component", "Technology"],
        [
            ["Mobile Framework", "Flutter 3.35.4 (Dart 3.9.2)"],
            ["Backend / Auth", "Firebase (Firestore, Auth, Storage, Crashlytics, Analytics, FCM)"],
            ["AI Engine", "Google Gemini API"],
            ["State Management", "Provider + service singletons"],
            ["Local Storage", "SharedPreferences + Hive"],
            ["Payments", "Google Play Billing / Apple IAP"],
        ]
    )

    add_heading(doc, "3. User Personas")
    personas = [
        ("Aspiring Parent — Alex", "Researching parenthood, joining communities for advice and support. Taster usage of AI Chat Helper."),
        ("Expecting Parent — Sarah", "Preparing for baby arrival, finding local groups and buying essentials. Heavy Marketplace and Events usage."),
        ("New Parent — Marcus", "Sleep-deprived but eager to connect. Uses Connect and Discover heavily. High AI Chat Helper usage for quick answers."),
        ("Experienced Parent — Priya", "Multiple children, sells outgrown items on Marketplace, organises meetups, creates private groups. Circle tier power user."),
        ("Local Service Provider — Dr. Chen", "Offers parenting classes, uses Events to promote. B2B partner use case."),
    ]
    for name, desc in personas:
        doc.add_heading(name, level=2)
        add_para(doc, desc)

    add_heading(doc, "4. Feature Specifications")
    doc.add_heading("4.1 Home (Smart Feed Dashboard)", level=2)
    for f in [
        "AI-powered smart feed — personalised content reordered by relevance",
        "Community announcements carousel (moderated by community leads)",
        "Borough badge with user's local area displayed prominently",
        "Upcoming meetups horizontal scroll",
        "Group activity cards — recent messages from joined groups",
        "New public groups recommendation row",
        "Invitation cards (if user has pending invitations)",
        "Learning maturity indicator — tracks AI personalisation progress",
        "Daily AI refresh indicator — shows when feed was last AI-refreshed",
    ]:
        add_bullet(doc, f)

    doc.add_heading("4.2 Connect (Groups & DMs)", level=2)
    for f in [
        "Auto-assigned groups based on postcode and parenting stage at sign-up",
        "Public group discovery and joining",
        "Private group creation (Neighbour+ tier)",
        "Group chat: message threads, image/file attachments, polls, emoji reactions",
        "Saved messages: bookmark important messages for later retrieval",
        "Forward message to other groups or contacts",
        "Direct messaging with any parent in the community",
        "Online status indicator in DMs",
        "AI Chat Summariser: generate plain-English summary of missed group chat",
        "Group polls: create, vote, and view results in-line",
        "Group members screen with admin controls",
        "Group details screen (edit group, invite members, leave group)",
    ]:
        add_bullet(doc, f)

    doc.add_heading("4.3 Discover (Events & Meetups)", level=2)
    for f in [
        "AI-discovered local events (daily web scraping — borough filtered)",
        "B2B partner events from local businesses (nurseries, classes, services)",
        "Parent-created meetups with RSVP (name, date, location, max attendees)",
        "Filters: area, date range, age group, event category",
        "Map view and list view toggle",
        "Event detail screen with full description and RSVP CTA",
        "Meetup detail screen with attendee list and messaging",
        "Create meetup screen (Neighbour+ tier)",
        "AI event recommendations (Neighbour+ tier)",
    ]:
        add_bullet(doc, f)

    doc.add_heading("4.4 Market (Pre-loved Marketplace)", level=2)
    for f in [
        "Buy, sell, and give away baby and children's items",
        "AI Listing Writer: auto-generates item description from category and title",
        "Category browsing: Clothing, Toys, Furniture, Equipment, Books, Other",
        "Search with real-time filtering",
        "Item detail: photo carousel (up to limit), description, price (or Free), seller info",
        "Direct messaging to seller (in-app, no phone number sharing required)",
        "Photo uploads: 3 (Welcome), 15 (Neighbour), 50 (Circle)",
        "Listings per user: 2 (Welcome), 15 (Neighbour), Unlimited (Circle)",
        "Mark item as sold or remove listing",
    ]:
        add_bullet(doc, f)

    doc.add_heading("4.5 Profile & Settings", level=2)
    for f in [
        "Edit profile: name, bio, photo, children info, parenting stage",
        "Subscription management: view current tier, next renewal, upgrade/downgrade/cancel",
        "Subscription plans screen: three-tier comparison with monthly/annual billing toggle",
        "Notification preferences: group messages, DMs, events, system announcements",
        "Privacy settings: profile visibility, data sharing preferences",
        "Run Tutorial option (re-triggers the in-app tutorial overlay)",
        "Help & Support link",
        "About and version info",
        "Legal links: Privacy Policy, Terms of Service",
        "Sign Out",
    ]:
        add_bullet(doc, f)

    add_heading(doc, "5. Subscription Model")
    add_para(doc, "Three-tier subscription model with freemium conversion strategy. No founding member pricing.")
    for name, monthly, annual, features in TIERS:
        doc.add_heading(f"{name} — {monthly} / {annual}", level=2)
        for f in features:
            add_bullet(doc, f)

    add_heading(doc, "6. Data Architecture")
    doc.add_heading("6.1 Firestore Collections", level=2)
    add_table(doc,
        ["Collection", "Key Fields", "Purpose"],
        [
            ["users", "uid, name, phone, postcode, borough, parenting_stage, children, photo_url, tier, created_at", "User profiles"],
            ["groups", "id, name, description, type (public/private), borough, admin_uid, member_count", "Community groups"],
            ["messages", "id, group_id, sender_uid, text, timestamp, thread_id, type", "Group and DM messages"],
            ["events", "id, title, description, date, location, borough, organiser_type, rsvp_count", "Events and meetups"],
            ["listings", "id, seller_uid, title, description, price, category, photos, borough, status", "Marketplace listings"],
            ["announcements", "id, title, body, borough, created_at, expires_at", "Community announcements"],
            ["subscriptions", "uid, tier, billing_period, start_date, renewal_date, is_active", "Subscription state"],
        ]
    )
    doc.add_heading("6.2 Borough Scoping", level=2)
    add_para(doc, "All Firestore queries are filtered by borough field. The BoroughScopeGuard service enforces this at the service layer — no query can return data outside the user's assigned borough. This is the core data architecture decision that makes Huddl genuinely hyperlocal.")

    add_heading(doc, "7. Non-Functional Requirements")
    add_table(doc,
        ["Requirement", "Target"],
        [
            ["App launch time (cold)", "<3 seconds"],
            ["Feed load time", "<2 seconds"],
            ["AI Chat Helper response", "<5 seconds"],
            ["Push notification delivery", "<30 seconds"],
            ["Availability", "99.5% uptime (Firebase SLA backed)"],
            ["GDPR compliance", "Full UK GDPR + Children's Code compliance"],
            ["Data residency", "EU region Firebase project"],
        ]
    )

    add_page_break(doc)
    add_heading(doc, "APPENDIX A — AI Roadmap (Future Features)")
    add_para(doc, "The following features are NOT live in the current build. They are on the product roadmap for post-seed development.", bold=True)
    for name, desc in AI_ROADMAP:
        add_bullet(doc, f" {desc}", bold_prefix=f"{name}: ")

    doc.save(path)
    print(f"✅ FDD saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# 6. SOFTWARE DESIGN DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────

def create_sdd(path):
    doc = Document()
    doc.add_heading("Huddl Connect", 0)
    doc.add_paragraph("Software Design Document (SDD)")
    doc.add_paragraph(VERSION)
    add_page_break(doc)

    add_heading(doc, "Document Control")
    add_table(doc, ["Version", "Date", "Author", "Changes"],
        [["3.0", "June 2026", "Conrad / Malgorzata", "Removed RevGlue/Deals. Updated navigation to 5 tabs. Updated to current lib/ structure."]])

    add_heading(doc, "1. Introduction")
    doc.add_heading("1.1 Purpose", level=2)
    add_para(doc, "This document describes the software architecture, component design, data models, and service integrations for Huddl Connect. It is the authoritative technical reference for the Flutter application.")
    doc.add_heading("1.2 Scope", level=2)
    add_para(doc, "Covers all application layers: UI (screens/widgets), business logic (services), data models, navigation, state management, Firebase integration, and AI integration. No RevGlue, no Deals tab, no Trips tab.")
    doc.add_heading("1.3 Technology Stack", level=2)
    add_table(doc,
        ["Component", "Technology", "Version"],
        [
            ["Framework", "Flutter", "3.35.4"],
            ["Language", "Dart", "3.9.2"],
            ["Backend", "Firebase (Firestore, Auth, Storage, Crashlytics, FCM)", "firebase_core: 3.6.0"],
            ["AI Engine", "Google Gemini API", "Via ai_api_helper.dart"],
            ["State Management", "Provider pattern", "6.1.5+1"],
            ["Local Storage", "SharedPreferences + Hive", "2.5.3 + 2.2.3"],
            ["Fonts", "Google Fonts (Poppins + Inter)", "6.2.1"],
            ["Payments", "Google Play Billing + Apple IAP", "Platform native"],
        ]
    )

    add_heading(doc, "2. System Architecture")
    doc.add_heading("2.1 High-Level Architecture", level=2)
    for layer in [
        "Presentation Layer: Flutter screens, widgets, theme (lib/screens/, lib/widgets/)",
        "Business Logic Layer: Services (singleton pattern) (lib/services/)",
        "Data Layer: Firebase Firestore, local storage (SharedPreferences/Hive) (lib/models/)",
        "Navigation Layer: Named routes via AppRouter (lib/config/router.dart)",
        "Config Layer: Firebase options, theme, constants (lib/config/, lib/theme/, lib/constants/)",
    ]:
        add_bullet(doc, layer)

    doc.add_heading("2.2 Navigation Architecture", level=2)
    add_para(doc, "MainShell uses an IndexedStack with 5 child screens (Home, Groups, Events, Marketplace, Profile) and a custom bottom navigation bar. The AppRouter handles all named route navigation for modal/detail screens.")

    add_heading(doc, "3. Service Layer")
    add_para(doc, "All services use the singleton pattern for shared state. Key services:")
    add_table(doc,
        ["Service", "Purpose"],
        [
            ["FirebaseAuthService", "Firebase Auth singleton — phone OTP, sign-in, sign-out, profile creation"],
            ["SubscriptionService", "Tier management, feature gating, usage counters, purchase flow"],
            ["OnboardingDataService", "User profile data from onboarding — postcode, parenting stage, children"],
            ["DefaultGroupService", "Auto-assigns users to borough groups on sign-up"],
            ["PostcodeService", "Postcode → borough resolution API"],
            ["GroupService (via screens)", "CRUD for groups — Firestore-backed"],
            ["DMService", "Direct message threads — Firestore real-time streams"],
            ["EventService", "Events CRUD and discovery — Firestore + AI"],
            ["MeetupService", "Meetup creation, RSVP, and management"],
            ["RehomeService", "Marketplace listing CRUD (file: rehome_service.dart)"],
            ["AiCopilotService", "AI Chat Helper — Gemini API integration"],
            ["AiFeedService", "Smart feed scoring and personalisation"],
            ["AiEventDiscoveryService", "Daily local event discovery via Gemini"],
            ["AiChatSummariserService", "Group chat summarisation"],
            ["AiListingService", "Marketplace AI Listing Writer"],
            ["AiMatchmakerService", "AI Meetup Matchmaker (Circle tier)"],
            ["GeminiSystemPromptBuilder", "Assembles borough-grounded Gemini prompts"],
            ["BoroughScopeGuard", "Enforces borough-scoped data access"],
            ["TutorialService", "In-app tutorial overlay lifecycle management"],
            ["AnnouncementService", "Community announcements — Firestore-backed"],
            ["InvitationService", "Group invitation management"],
            ["PollService", "Group polls CRUD and voting"],
            ["SavedMessageService", "Bookmark messages for later"],
            ["FeedbackService", "User feedback on AI responses"],
            ["PaymentService", "Subscription purchase flow — Play Billing / Apple IAP"],
            ["BiometricAuthService", "Optional biometric lock for the app"],
            ["BackupRestoreService", "User data export and restore"],
        ]
    )

    add_heading(doc, "4. Screen Component Design")
    doc.add_heading("4.1 MainShell Architecture", level=2)
    add_para(doc, "IndexedStack preserves all 5 tab states in memory. Tab switches are O(1) — no rebuild cost. The custom bottom nav bar is a floating pill with 70px height, 28px border radius, and teal active colour.")

    doc.add_heading("4.2 HomeScreen Architecture", level=2)
    add_para(doc, "The HomeScreen implements the 'Invisible AI Redesign' pattern: a unified smart feed replaces 10 separate content sections. AiFeedService scores and reorders feed items on load. DailyAiRefreshService pre-computes scores nightly. The feed uses a _SmartFeedItem discriminated union type to render group activity, meetup cards, event cards, and announcement cards uniformly.")

    doc.add_heading("4.3 GroupsScreen / GroupChatScreen Architecture", level=2)
    add_para(doc, "GroupsScreen uses Firestore real-time snapshots for group list. GroupChatScreen uses a StreamBuilder on the messages sub-collection. Message input supports text, polls, image attachments, and thread replies. AI Summariser button available to Neighbour/Circle tiers — triggers AiChatSummariserService with the last N messages.")

    doc.add_heading("4.4 MarketplaceScreen Architecture", level=2)
    add_para(doc, "Grid layout with category filter chips at top. CreateListingScreen uses RehomeService. AI Listing Writer button triggers AiListingService.generateDescription(). Photo uploads use Firebase Storage with per-tier limit enforcement.")

    add_heading(doc, "5. Subscription & Feature Gating")
    add_para(doc, "SubscriptionService manages three tiers (Explorer/Welcome, Neighbourhood/Neighbour, InnerCircle/Circle) with feature limits tracked via usage counters. Feature gate checks use boolean getters (e.g. canJoinGroup, canCreatePrivateGroup, canUseAiSummaries). Gates trigger UpgradePrompt bottom sheet via UpgradePromptWidget.")

    add_heading(doc, "6. Project File Structure")
    structure = """lib/
  main.dart
  config/
    firebase_options.dart
    router.dart
  constants/
  models/
    group.dart
    subscription.dart
  screens/
    main_shell.dart
    home/
      home_screen.dart
    groups/
      groups_screen.dart
      group_chat_screen.dart
      dm_chat_screen.dart
      create_group_screen.dart
      group_details_screen.dart
      group_members_screen.dart
      group_polls_screen.dart
      poll_detail_screen.dart
      create_poll_screen.dart
      saved_messages_for_group_screen.dart
      thread_reply_screen.dart
      new_dm_screen.dart
      forward_message_sheet.dart
      image_gallery_picker.dart
    events/
      events_screen.dart
      event_detail_screen.dart
      meetup_detail_screen.dart
      create_event_screen.dart
      create_meetup_screen.dart
    marketplace/
      marketplace_screen.dart
      item_detail_screen.dart
    rehome/
      create_listing_screen.dart
      home_screen.dart  (marketplace shell)
      journey_map_screen.dart
    profile/
      profile_screen.dart
      backup_restore_screen.dart
    subscription/
      subscription_plans_screen.dart
      subscription_checkout_screen.dart
      manage_subscription_screen.dart
    ai/
      ai_copilot_screen.dart
      ai_listing_generator_sheet.dart
      ai_matchmaker_sheet.dart
    onboarding/
      splash_screen.dart
      onboarding_carousel_screen.dart
      phone_number_screen.dart
      verification_screen.dart
      name_input_screen.dart
      password_screen.dart
      parent_type_screen.dart
      stage_of_life_screen.dart
      about_you_screen.dart
      child_info_screen.dart
      due_date_screen.dart
      postcode_screen.dart
      add_photo_screen.dart
      welcome_complete_screen.dart
      not_available_screen.dart
    legal/
      privacy_policy_detail_screen.dart
      terms_of_service_screen.dart
    debug/
  services/
    [54 service files — see AI SDD for AI services list]
  theme/
    huddl_colors.dart
    huddl_theme.dart
  utils/
  widgets/
    huddl_widgets.dart
    upgrade_prompt.dart
    borough_badge.dart
    learning_maturity_indicator.dart
    tutorial/
      tutorial_overlay.dart"""
    add_para(doc, structure)

    doc.save(path)
    print(f"✅ SDD saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# 7. DEPLOYMENT GUIDE
# ─────────────────────────────────────────────────────────────────────────────

def create_deployment_guide(path):
    doc = Document()
    doc.add_heading("Huddl Connect", 0)
    doc.add_paragraph("Deployment Guide")
    doc.add_paragraph(VERSION)
    add_page_break(doc)

    add_heading(doc, "Document Control")
    add_table(doc, ["Version", "Date", "Author", "Changes"],
        [["3.0", "June 2026", "Conrad / Malgorzata", "Removed RevGlue/Deals references. Updated to 5-tab navigation. Added iOS TestFlight notes."]])

    add_heading(doc, "1. Environment Setup")
    add_table(doc,
        ["Component", "Version", "Notes"],
        [
            ["Flutter SDK", "3.35.4", "DO NOT UPDATE — locked for stability"],
            ["Dart", "3.9.2", "DO NOT UPDATE — locked for stability"],
            ["Android Studio / SDK", "API Level 35", "DO NOT UPDATE"],
            ["Java", "OpenJDK 17.0.2", "DO NOT UPDATE"],
            ["Xcode", "Latest stable", "Required for iOS builds"],
            ["Firebase CLI", "Latest", "For Firebase deployment"],
            ["CocoaPods", "Latest", "Required for iOS pod install"],
        ]
    )

    add_heading(doc, "2. Firebase Configuration")
    for step in [
        "Place google-services.json at android/app/google-services.json",
        "Place GoogleService-Info.plist at ios/Runner/GoogleService-Info.plist",
        "Ensure firebase_options.dart exists at lib/config/firebase_options.dart",
        "Firestore database must be created in Firebase Console before first run",
        "Firestore security rules must be configured (see Section 3)",
        "Firebase Auth must have Phone provider enabled",
        "Firebase Storage must be enabled with appropriate rules",
        "Crashlytics: enabled in firebase_options.dart and main.dart",
        "Firebase Analytics: enabled",
        "FCM: configured for push notifications (Android + iOS)",
    ]:
        add_bullet(doc, step)

    add_heading(doc, "3. Firestore Security Rules")
    add_para(doc, "Development rules (allow all reads and writes within authenticated context):")
    add_para(doc, """rules_version = '2';
service cloud.firestore {
  match /databases/{database}/documents {
    match /{document=**} {
      allow read, write: if request.auth != null;
    }
  }
}""")
    add_para(doc, "Production rules should be scoped per collection with borough validation.")

    add_heading(doc, "4. Android Build")
    add_para(doc, "Commands to build and deploy Android app:")
    for cmd in [
        "cd /home/user/flutter_app",
        "flutter pub get",
        "flutter analyze  # Must pass before build",
        "flutter build apk --release  # Debug APK for testing",
        "flutter build appbundle --release  # AAB for Play Store",
    ]:
        add_bullet(doc, cmd)
    add_para(doc, "Signing config: android/key.properties + android/release-key.jks")
    add_para(doc, "Package name: must match google-services.json package_name exactly")

    add_heading(doc, "5. iOS Build")
    for step in [
        "cd ios && pod install && cd ..",
        "Open ios/Runner.xcworkspace in Xcode",
        "Set signing team and bundle identifier",
        "Ensure AppDelegate.swift sets appVerificationDisabledForTesting BEFORE GeneratedPluginRegistrant.register (test builds only)",
        "Product → Archive → Distribute App → TestFlight",
        "Wait for App Store Connect processing (~10 minutes)",
        "Add internal testers in App Store Connect",
    ]:
        add_bullet(doc, step)

    add_heading(doc, "6. Web Preview")
    for cmd in [
        "cd /home/user/flutter_app",
        "flutter build web --release",
        "python3 -m http.server 5060 --directory build/web --bind 0.0.0.0",
    ]:
        add_bullet(doc, cmd)
    add_para(doc, "CORS server command for iframe-embedded preview:")
    add_para(doc, "cd build/web && python3 -c \"import http.server, socketserver; [CORS handler]\" &")

    add_heading(doc, "7. Key Dependencies (pubspec.yaml)")
    add_table(doc,
        ["Package", "Version", "Purpose"],
        [
            ["firebase_core", "3.6.0", "Firebase Core SDK"],
            ["cloud_firestore", "5.4.3", "Firestore Database"],
            ["firebase_auth", "5.3.1", "Authentication"],
            ["firebase_storage", "12.3.2", "Cloud Storage"],
            ["firebase_messaging", "15.1.3", "Push Notifications"],
            ["firebase_analytics", "11.3.3", "Analytics"],
            ["firebase_crashlytics", "4.1.3", "Crash Reporting"],
            ["provider", "6.1.5+1", "State Management"],
            ["shared_preferences", "2.5.3", "Key-Value Storage"],
            ["hive", "2.2.3", "Document Database"],
            ["hive_flutter", "1.1.0", "Hive Flutter Integration"],
            ["google_fonts", "6.2.1", "Poppins + Inter fonts"],
            ["dynamic_color", "1.7.0", "Material You dynamic theming"],
        ]
    )

    add_heading(doc, "8. Environment Variables & Secrets")
    add_para(doc, "Never commit secrets to git. Store in:")
    for item in [
        "android/key.properties — keystore path and passwords",
        "android/app/google-services.json — Firebase Android config (gitignored)",
        "ios/Runner/GoogleService-Info.plist — Firebase iOS config (gitignored)",
        "lib/config/firebase_options.dart — Firebase web/platform config (review before commit)",
        "Gemini API key — stored in Firebase Remote Config or environment variable, never hardcoded",
    ]:
        add_bullet(doc, item)

    doc.save(path)
    print(f"✅ Deployment Guide saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# 8. STORE PUBLISHING GUIDE
# ─────────────────────────────────────────────────────────────────────────────

def create_store_guide(path):
    doc = Document()
    doc.add_heading("Huddl Connect", 0)
    doc.add_paragraph("App Store Publishing Guide")
    doc.add_paragraph(VERSION)
    add_page_break(doc)

    add_heading(doc, "Document Control")
    add_table(doc, ["Version", "Date", "Author", "Changes"],
        [["3.0", "June 2026", "Conrad / Malgorzata", "Removed RevGlue/Deals/Trips references. Updated store listing to reflect 5-tab app and current features."]])

    add_heading(doc, "1. App Store Metadata")
    doc.add_heading("1.1 App Name & Description", level=2)
    add_para(doc, "App Name: huddl connect")
    add_para(doc, "Short Description (Google Play):")
    add_para(doc, "Your borough's parent community. Groups, events, marketplace & AI — all hyperlocal.")
    add_para(doc, "Long Description:")
    add_para(doc, """Huddl Connect is the hyperlocal parenting community app — your borough's mums and dads, in one place.

WHAT'S INSIDE:
• Connect — join local parent groups, chat, send DMs, create polls
• Discover — find local events and parent meetups near you, or organise your own
• Market — buy, sell, and give away pre-loved baby and kids' items
• Home — your personalised local parenting feed, powered by AI
• AI Tools — AI Chat Helper, Group Summaries, Listing Writer, Events Finder & more

WHY HUDDL?
✓ Hyperlocal — scoped to your borough. No national noise.
✓ AI-powered — 7 AI features built for parents, not generic chatbots
✓ Private & safe — phone-verified accounts, no anonymous strangers
✓ Free to start — 7-day Neighbour trial on sign-up, no card required

SUBSCRIPTION TIERS:
• Welcome — Free forever. 2 groups, 2 meetups/mo, AI Chat Helper (3/day)
• Neighbour — £5.99/mo or £49.99/yr. Unlimited groups, DMs, full AI suite
• Circle — £11.99/mo or £99.99/yr. Everything unlimited + AI Meetup Matchmaker

Built by Cambridge parents, for parents everywhere.""")

    doc.add_heading("1.2 Keywords", level=2)
    keywords = ["local parenting app", "parent community", "mum groups near me", "local baby events", "pre-loved baby items", "parenting AI", "local parents", "mums and dads app", "borough community", "parent meetups"]
    add_para(doc, ", ".join(keywords))

    doc.add_heading("1.3 Category & Rating", level=2)
    add_table(doc,
        ["Field", "Value"],
        [
            ["Primary Category", "Social Networking (Google Play) / Lifestyle (Apple)"],
            ["Secondary Category", "Parenting"],
            ["Age Rating", "4+ (Apple) / Everyone (Google Play)"],
            ["Content Rating", "No mature content"],
            ["Privacy Policy URL", "Required — must be provided"],
        ]
    )

    add_heading(doc, "2. Google Play Store Setup")
    for step in [
        "Create app in Google Play Console at play.google.com/console",
        "Upload signed AAB: flutter build appbundle --release",
        "Complete content rating questionnaire (age-appropriate: Everyone)",
        "Complete data safety form: data collected (name, phone, postcode, children info), purpose (app functionality), encryption in transit",
        "Set up in-app purchases: Neighbour Monthly, Neighbour Annual, Circle Monthly, Circle Annual",
        "Add store listing screenshots: phone (1080x1920), tablet (optional)",
        "Add feature graphic: 1024x500px",
        "Submit for review",
    ]:
        add_bullet(doc, step)

    add_heading(doc, "3. Apple App Store Setup")
    for step in [
        "Create app record in App Store Connect at appstoreconnect.apple.com",
        "Archive in Xcode and upload via Organiser",
        "Set up in-app purchases matching Play Store products",
        "Complete App Privacy questionnaire",
        "Add screenshots: 6.7-inch, 6.5-inch, 5.5-inch, iPad (if universal)",
        "Fill App Store listing: name, subtitle, description, keywords, support URL",
        "Submit for review",
        "NOTE: Firebase phone auth requires Apple's APNs push notification certificate to be configured",
    ]:
        add_bullet(doc, step)

    add_heading(doc, "4. In-App Purchase Product IDs")
    add_table(doc,
        ["Product", "Product ID", "Price", "Type"],
        [
            ["Neighbour Monthly", "neighbourhood_monthly", "£5.99", "Auto-renewable subscription"],
            ["Neighbour Annual", "neighbourhood_annual", "£49.99", "Auto-renewable subscription"],
            ["Circle Monthly", "inner_circle_monthly", "£11.99", "Auto-renewable subscription"],
            ["Circle Annual", "inner_circle_annual", "£99.99", "Auto-renewable subscription"],
        ]
    )

    add_heading(doc, "5. Privacy Policy Requirements")
    for req in [
        "Data collected: name, phone number (for auth), postcode/borough, parenting stage, children's ages (not names in production), profile photo",
        "Data NOT collected: children's full names stored server-side, financial data (handled by platform billing), precise location",
        "Data sharing: not sold to third parties. Anonymous analytics to Firebase Analytics.",
        "Children's Code compliance: app is for parents (18+), not directly for children",
        "GDPR compliance: right to access, right to delete, data portability via backup/restore feature",
        "Data retention: user data deleted within 30 days of account deletion request",
    ]:
        add_bullet(doc, req)

    add_heading(doc, "6. App Store Optimisation (ASO)")
    add_table(doc,
        ["Element", "Strategy"],
        [
            ["Title", "Include 'local parents' and 'community' for search ranking"],
            ["Subtitle", "Short, benefit-focused: 'Your borough's parent community'"],
            ["Keywords", "Rotate quarterly: local, parenting, community, mums, dads, borough, groups, events, baby, kids"],
            ["Screenshots", "Show key screens: Home feed, Group chat, Events map, Marketplace grid, AI Chat"],
            ["Preview video", "30-second walkthrough: sign up → join group → discover event → list item → AI chat"],
            ["Ratings", "In-app review prompt after first successful group join and first AI chat"],
        ]
    )

    doc.save(path)
    print(f"✅ Store Guide saved: {path}")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN: Run all generators
# ─────────────────────────────────────────────────────────────────────────────

def main():
    os.makedirs("/home/user/flutter_app/docs", exist_ok=True)
    docs_dir = "/home/user/flutter_app/docs"

    print("🚀 Generating all updated Huddl Connect documentation...")
    print("=" * 60)

    create_gtm(f"{docs_dir}/Huddl_Connect_Go_To_Market_Strategy.docx")
    create_pitch_deck(f"{docs_dir}/Huddl_Connect_Investor_Pitch_Deck.docx")
    create_ai_sdd(f"{docs_dir}/Huddl_Connect_AI_System_Design_Document.docx")
    create_ux_doc(f"{docs_dir}/Huddl_Connect_UX_Design_Documentation.docx")
    create_fdd(f"{docs_dir}/Huddl_Connect_Functional_Design_Document.docx")
    create_sdd(f"{docs_dir}/Huddl_Connect_Software_Design_Document.docx")
    create_deployment_guide(f"{docs_dir}/Huddl_Connect_Deployment_Guide.docx")
    create_store_guide(f"{docs_dir}/Huddl_Connect_Store_Publishing_Guide.docx")

    print("=" * 60)
    print("✅ ALL DOCUMENTS GENERATED SUCCESSFULLY")
    print(f"📁 Location: {docs_dir}/")

if __name__ == "__main__":
    main()
