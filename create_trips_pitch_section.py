"""
Add Huddl Trips AI Feature section to the existing pitch deck
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p

def add_styled_paragraph(doc, text, bold=False, italic=False, size=11, color=None, alignment=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p

def create_styled_table(doc, headers, rows, col_widths=None):
    """Create a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Style header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'FF975C')
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # Style rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            bg = 'FFFFFF' if r_idx % 2 == 0 else 'FFF8F0'
            set_cell_shading(cell, bg)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    
    return table

# ── Load existing pitch deck ──
src = '/home/user/uploaded_files/huddl_pitch_deck_restored_20260309170311.pdf'

# Create the supplementary document
doc = Document()

# ── Page setup ──
section = doc.sections[0]
section.page_width = Inches(11.69)
section.page_height = Inches(8.27)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.5)

orange = RGBColor(0xFF, 0x97, 0x5C)
dark = RGBColor(0x43, 0x46, 0x4D)
teal = RGBColor(0x19, 0x9A, 0x85)
white_rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ════════════════════════════════════════════════════════════════════════
# SLIDE 1: TITLE SLIDE
# ════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('huddl')
run.font.size = Pt(48)
run.font.color.rgb = orange
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TRIPS')
run.font.size = Pt(36)
run.font.color.rgb = dark
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('AI Family Travel Concierge')
run.font.size = Pt(24)
run.font.color.rgb = teal
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run('The only travel assistant that knows your child\'s age,\nyour community\'s reviews, and what\'s in the nearest shop\n— because it was built by parents, for parents.')
run.font.size = Pt(14)
run.font.color.rgb = dark
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('PITCH DECK SUPPLEMENT — 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x94, 0x94, 0x94)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# SLIDE 2: THE INSIGHT — WHY TRAVEL?
# ════════════════════════════════════════════════════════════════════════
add_styled_heading(doc, 'THE INSIGHT', level=1, color=orange)
add_styled_heading(doc, 'Why Travel is the Killer Feature for a Parenting Community', level=2, color=dark)

add_styled_paragraph(doc, 'In our Cambridge WhatsApp pilot (700+ parents, 40 groups, 18,000+ messages in a single group), one of the most organically popular chat groups was the Parent Travels group — where parents asked questions about travelling to places with children.', size=12)

add_styled_paragraph(doc, 'This is not a coincidence. It reveals a massive unserved need:', size=12, bold=True)

# Key stats table
create_styled_table(doc, 
    ['The Problem', 'The Data', 'The Impact'],
    [
        ['Parents juggle 10+ apps to plan one family trip', 'TripAdvisor + Google Maps + NHS Travel Health + Airline policies + Weather + 15 browser tabs', 'Cognitive overload — the #1 pain point in our pitch deck'],
        ['Generic travel advice doesn\'t account for children', 'No travel app knows your child is 14 months old and needs a cot', 'Parents waste hours researching age-inappropriate options'],
        ['Trust deficit with anonymous reviews', 'Parents trust "Sarah from our group went to Malaga with a 2-year-old" over TripAdvisor', 'Community intelligence is 3x more trusted than anonymous reviews'],
        ['Travel amplifies parent loneliness', 'Being in a foreign place with young children = peak isolation', 'The exact crisis our platform was built to solve'],
    ]
)

doc.add_paragraph()
add_styled_paragraph(doc, 'Travelling with kids sits at the intersection of every pain point huddl already solves: cognitive overload, information fragmentation, trust deficit, and loneliness amplification.', size=12, bold=True, color=teal)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# SLIDE 3: THE FEATURE — OVERVIEW
# ════════════════════════════════════════════════════════════════════════
add_styled_heading(doc, 'HUDDL TRIPS', level=1, color=orange)
add_styled_heading(doc, 'AI Family Travel Concierge — Feature Overview', level=2, color=dark)

add_styled_paragraph(doc, 'A context-aware AI assistant that combines community-sourced parent intelligence with real-time personalisation based on your specific children\'s ages, needs, and stage of life.', size=12, italic=True)

doc.add_paragraph()

features = [
    ['1. "Ask Parents Who\'ve Been"', 'Community-Powered AI Answers', 'When a parent asks "Is Tenerife good for a 14-month-old in March?", the AI mines community conversation history, finds parents who visited, anonymises and synthesises their tips, and links back to the community for direct connection.'],
    ['2. "My Family" Context Engine', 'Personalised Intelligence', 'The AI knows your children\'s ages, dietary needs, nap schedules, stage of life, and postcode. It filters recommendations, adjusts itineraries around nap windows, flags airline restrictions for pregnant mums, and suggests age-appropriate activities.'],
    ['3. "Pack My Bag" Generator', 'Smart Packing Lists', 'Auto-generates personalised packing lists based on destination, duration, children\'s ages, and weather. Flags what to buy locally vs. bring from home. Integrates with the Marketplace — "3 parents in your area are lending travel cots this month."'],
    ['4. Live Trip Companion', 'Real-Time Assistant', 'While travelling: nearest nappy shop open now, nearest indoor play area for a meltdown, restaurants with highchairs within 500m rated by huddl parents, rainy-day activity pivots — all personalised to your children\'s ages.'],
    ['5. "Parents Abroad" Hub', 'Temporary Travel Communities', 'When families arrive at a destination, they join a temporary travel hub. "4 huddl families are in Malaga this week — want to connect?" Turns solo holidays into micro-community experiences. Auto-creates trip reviews for the next parent.'],
]

create_styled_table(doc,
    ['Feature', 'Category', 'Description'],
    features
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# SLIDE 4: COMMUNITY-POWERED AI — THE MOAT
# ════════════════════════════════════════════════════════════════════════
add_styled_heading(doc, 'THE MOAT', level=1, color=orange)
add_styled_heading(doc, '"Ask Parents Who\'ve Been" — Community-Powered AI Answers', level=2, color=dark)

add_styled_paragraph(doc, 'This is what makes huddl Trips impossible to replicate by any travel app.', size=13, bold=True, color=teal)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('How it works:')
run.font.size = Pt(12)
run.bold = True

steps = [
    ('Parent asks:', '"Is Tenerife good for a 14-month-old in March?"'),
    ('AI mines community data:', 'Finds that 6 parents in the Cambridge group discussed Tenerife last summer'),
    ('AI synthesises:', '"Based on 6 huddl parents who visited Tenerife with under-2s: 5/6 recommended it. Top tips: Book Siam Park\'s baby area in advance, the south coast beaches have calmer waves, Lidl Tenerife stocks Ella\'s Kitchen pouches"'),
    ('AI connects:', '"Sarah shared photos from her trip — would you like to message her?"'),
    ('Flywheel:', 'Every trip creates reviews that make the AI smarter for the next parent'),
]

for label, detail in steps:
    p = doc.add_paragraph()
    run1 = p.add_run(f'{label} ')
    run1.font.size = Pt(11)
    run1.bold = True
    run1.font.color.rgb = orange
    run2 = p.add_run(detail)
    run2.font.size = Pt(11)
    run2.font.color.rgb = dark

doc.add_paragraph()

add_styled_paragraph(doc, 'WHY NO COMPETITOR CAN COPY THIS:', size=12, bold=True, color=orange)

create_styled_table(doc,
    ['Competitor', 'What They Have', 'What They\'re Missing'],
    [
        ['TripAdvisor', 'Millions of reviews', 'No child-age context, no community trust, anonymous reviewers'],
        ['Booking.com', 'Hotel inventory', 'Doesn\'t know your child is 14 months and needs a cot'],
        ['Peanut', 'Parent social network', 'No travel features, no integrated marketplace, women-only'],
        ['Google Maps', 'Location data', 'No parent community, no age-appropriate filtering'],
        ['huddl Trips', 'Community trust + child context + local intelligence', 'THE ONLY PLATFORM WITH ALL THREE'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# SLIDE 5: MY FAMILY CONTEXT ENGINE
# ════════════════════════════════════════════════════════════════════════
add_styled_heading(doc, 'CONTEXT ENGINE', level=1, color=orange)
add_styled_heading(doc, '"My Family" — Personalised Travel Intelligence', level=2, color=dark)

add_styled_paragraph(doc, 'The AI leverages data already collected during onboarding — zero additional input required.', size=12, italic=True)

create_styled_table(doc,
    ['Data Point', 'Source', 'How Trips Uses It'],
    [
        ['Child ages (14mo, 3yr)', 'Onboarding: Child Info', 'Filters age-appropriate activities, highchair restaurants, playground suitability'],
        ['Stage of life (expecting)', 'Onboarding: Stage of Life', 'Adjusts recs — no hiking trails for 9-month pregnant mum, airline cut-off dates'],
        ['Dietary needs / allergies', 'Profile settings', 'Flags restaurants with allergen menus, locates nearest pharmacy'],
        ['Due date', 'Onboarding: Due Date', 'Travel insurance warnings, airline restrictions, hospital proximity at destination'],
        ['Postcode / borough', 'Onboarding: Postcode', 'Suggests direct flights from nearest airports, school holiday dates'],
        ['Nap schedule', 'Health tracker (future)', 'Builds itineraries around nap windows: "Beach morning, quiet time 1pm at villa"'],
        ['Community connections', 'Chat groups', 'Finds parents who\'ve visited, connects for advice, suggests travel buddies'],
    ]
)

doc.add_paragraph()
add_styled_paragraph(doc, 'ZERO ADDITIONAL ONBOARDING — all context is already captured in the existing user journey. This is the power of an all-in-one platform vs. a standalone travel app.', size=12, bold=True, color=teal)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# SLIDE 6: STRATEGIC VALUE & REVENUE
# ════════════════════════════════════════════════════════════════════════
add_styled_heading(doc, 'STRATEGIC VALUE', level=1, color=orange)
add_styled_heading(doc, 'Revenue, Retention & Network Effects', level=2, color=dark)

create_styled_table(doc,
    ['Strategic Angle', 'Impact', 'Detail'],
    [
        ['Revenue Driver', 'Premium Feature', 'Perfect for Growth (£7.99/mo) or Premium (£9.99/mo) tier. Parents will upgrade during holiday planning season (Jan-Mar, Sep-Oct)'],
        ['Retention Hook', '4-6 week engagement cycles', 'Parents plan trips 2-4x/year. Each trip = 4-6 weeks of daily app engagement during planning, booking, and review phases'],
        ['Marketplace Integration', 'Highest-value category', 'Travel gear (prams, car seats, travel cots) is the highest-value preloved category. "Borrow a travel cot from a local parent" drives marketplace GMV'],
        ['Data Flywheel', 'Self-reinforcing moat', 'Every trip review makes the AI smarter → better recs → more users → more reviews. Classic network effect investors want to see'],
        ['Seasonal Revenue Spikes', 'Predictable surges', 'Jan-Mar (summer planning) and Sep-Oct (half-term/Christmas) create subscription upgrade waves aligned with marketing spend'],
        ['B2B Opportunity', 'New revenue stream', 'Family-friendly hotels, airlines, and resorts would pay to be featured as "huddl recommended" — affiliate/sponsored placement revenue'],
        ['Competitive Moat', 'Unreplicable advantage', 'Community trust + child context + local intelligence. No single competitor has all three. This is a true Blue Ocean feature'],
    ]
)

doc.add_paragraph()

# Revenue projection mini-table
add_styled_paragraph(doc, 'REVENUE PROJECTION — HUDDL TRIPS UPSELL', size=12, bold=True, color=orange)

create_styled_table(doc,
    ['Metric', 'Conservative', 'Base Case', 'Optimistic'],
    [
        ['Users engaging with Trips', '15% of MAU', '25% of MAU', '40% of MAU'],
        ['Conversion to paid tier', '8%', '15%', '22%'],
        ['Avg. uplift per user/year', '£24', '£48', '£72'],
        ['Marketplace GMV from travel gear', '£5K/mo', '£15K/mo', '£40K/mo'],
        ['Affiliate revenue (hotels, flights)', '£2K/mo', '£8K/mo', '£20K/mo'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# SLIDE 7: ROADMAP INTEGRATION
# ════════════════════════════════════════════════════════════════════════
add_styled_heading(doc, 'ROADMAP FIT', level=1, color=orange)
add_styled_heading(doc, 'Huddl Trips Consolidates 4 Existing AI Features Into 1 Premium Product', level=2, color=dark)

add_styled_paragraph(doc, 'Rather than shipping 4 separate utilities, Huddl Trips packages them as one cohesive premium experience with a single narrative: "Travel smarter with parents who\'ve been there."', size=12, italic=True)

create_styled_table(doc,
    ['Feature #', 'Original Feature', 'How Trips Integrates It', 'Phase'],
    [
        ['#26', 'Vacation Planner', 'The AI itinerary engine — builds toddler-friendly schedules with community intelligence', 'Phase 1 (2026)'],
        ['#76', 'Travel Safety Alerts', 'Real-time health/safety monitoring for travel destinations', 'Phase 1 (2026)'],
        ['#94', 'Relocation Assistant', 'Neighbourhood scouting logic repurposed for trip destinations — finding parks, pharmacies, doctors', 'Phase 2 (2027)'],
        ['#98', 'Family Reunion Planner', 'Multi-family coordination for group holidays and meetups abroad', 'Phase 2 (2027)'],
        ['NEW', 'Community Travel Intelligence', 'Mining group chat history for destination-specific parent reviews and tips', 'Phase 1 (2026)'],
        ['NEW', 'Pack My Bag Generator', 'AI packing lists personalised to child age, destination, and weather', 'Phase 1 (2026)'],
        ['NEW', 'Parents Abroad Hub', 'Temporary travel communities connecting huddl families at the same destination', 'Phase 2 (2027)'],
        ['NEW', 'Live Trip Companion', 'Real-time "meltdown mode" assistance — nearest play area, nappy shop, restaurant with highchairs', 'Phase 2 (2027)'],
    ]
)

doc.add_paragraph()

add_styled_paragraph(doc, 'IMPLEMENTATION TIMELINE', size=12, bold=True, color=orange)

create_styled_table(doc,
    ['Phase', 'Timeline', 'Features', 'Investment'],
    [
        ['Phase 1 — MVP', 'Q3-Q4 2026', 'AI Concierge chat, Pack My Bag, Community travel reviews, Destination search', 'Frontend + AI integration (built into existing Flutter app)'],
        ['Phase 2 — Social', 'Q1-Q2 2027', 'Parents Abroad hub, Live Trip Companion, Multi-family coordination', 'Location services + real-time features'],
        ['Phase 3 — Monetise', 'Q3-Q4 2027', 'Affiliate partnerships, Sponsored destinations, B2B hotel/airline integrations', 'Business development + API integrations'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# SLIDE 8: USER JOURNEY
# ════════════════════════════════════════════════════════════════════════
add_styled_heading(doc, 'USER JOURNEY', level=1, color=orange)
add_styled_heading(doc, 'From Question to Booking — A Parent\'s Trip Through huddl Trips', level=2, color=dark)

journey_steps = [
    ['1. Discovery', 'Parent browses Trips tab or asks in chat group "Where should I take a toddler in February?"', 'Home Screen → Trips Tab'],
    ['2. AI Response', 'Concierge synthesises community reviews: "6 huddl parents visited Tenerife with under-2s. 5/6 recommend it."', 'Trips → AI Chat'],
    ['3. Destination Deep-Dive', 'Taps destination card → sees community reviews, safety alerts, weather, age-appropriate activities', 'Trips → Destination Detail'],
    ['4. Itinerary Builder', 'AI generates personalised itinerary around child\'s nap schedule, dietary needs, and parent interests', 'Destination → Itinerary'],
    ['5. Pack My Bag', 'Auto-generated packing list: "28 nappies + 6 spare, SPF50 sun cream, Calpol sachets for hand luggage"', 'Destination → Packing List'],
    ['6. Marketplace Integration', '"3 parents in Cambridge are lending travel cots — borrow one?" Links to Preloved marketplace', 'Packing List → Marketplace'],
    ['7. Parents Abroad', 'Arrives in Tenerife → "4 huddl families are here this week. Beach playdate Tuesday?"', 'Live Trip → Social Hub'],
    ['8. Live Companion', 'Day 3: child has meltdown → "Nearest indoor play: 8 min walk. Park with swings: 3 min"', 'Live Trip → Emergency Help'],
    ['9. Trip Review', 'Auto-prompted review on return. Tips fed back into AI for next parent', 'Post-Trip → Community Intelligence'],
    ['10. Flywheel', 'Review appears when next parent asks about Tenerife. Community gets smarter with every trip', 'Ongoing → Network Effect'],
]

create_styled_table(doc,
    ['Step', 'Experience', 'Screen / Location'],
    journey_steps
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════
# SLIDE 9: CLOSING
# ════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
run = p.add_run('huddl Trips')
run.font.size = Pt(42)
run.font.color.rgb = orange
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('"Travel smarter with parents who\'ve been there."')
run.font.size = Pt(20)
run.font.color.rgb = dark
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('The only travel assistant that combines:\n\n')
run.font.size = Pt(16)
run.font.color.rgb = dark

bullets = [
    'Community Trust — real reviews from parents in your neighbourhood',
    'Child Context — knows your child\'s age, needs, and nap schedule',  
    'Local Intelligence — what\'s nearby, what\'s open, what\'s age-appropriate',
]

for bullet in bullets:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'✓  {bullet}')
    run.font.size = Pt(14)
    run.font.color.rgb = teal
    run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(50)
run = p.add_run('No competitor has all three. Only huddl can.')
run.font.size = Pt(16)
run.font.color.rgb = orange
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run('© 2026 huddl. All rights reserved.')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x94, 0x94, 0x94)

# ── Save ──
output_path = '/home/user/flutter_app/Huddl_Trips_AI_Feature_Pitch_Deck_Supplement.docx'
doc.save(output_path)
file_size = os.path.getsize(output_path)
print(f"✅ Saved: {output_path}")
print(f"📄 Size: {file_size / 1024:.1f} KB")
print(f"📊 Slides: 9 pages")

