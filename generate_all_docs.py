"""
Huddl Connect — Complete Document Generator (v3.0)
Generates all 8 updated documents based on actual built app state.
No RevGlue/Offers/Deals. 5-tab navigation. 3 subscription tiers.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

VERSION = "3.0"
DATE = "April 2026"
CONFIDENTIAL = "Confidential"
OUTPUT_DIR = "/home/user/flutter_app/docs"

# ─── Colour palette ───────────────────────────────────────────────────────────
ORANGE   = RGBColor(0xFF, 0x6B, 0x35)   # Huddl brand orange
DARK     = RGBColor(0x1A, 0x1A, 0x2E)   # Near-black
MID      = RGBColor(0x44, 0x44, 0x66)   # Mid-grey/blue
LIGHT    = RGBColor(0xF5, 0xF5, 0xF5)   # Off-white bg hint
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
GREEN    = RGBColor(0x2E, 0x7D, 0x32)
BLUE     = RGBColor(0x15, 0x65, 0xC0)

# ─── Helper functions ─────────────────────────────────────────────────────────

def new_doc():
    doc = Document()
    # Margins
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)
    return doc

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p

def body(doc, text, bold=False, color=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix + ": ")
        r.bold = True
        r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    return p

def divider(doc):
    doc.add_paragraph("─" * 80)

def cover_block(doc, title, subtitle, version, date, authors):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = ORANGE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(14)
    r2.font.color.rgb = MID

    doc.add_paragraph()
    for line in [f"Version {version}  |  {date}  |  {CONFIDENTIAL}", authors]:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(line)
        r3.font.size = Pt(10)
        r3.font.color.rgb = MID
    doc.add_page_break()

def tier_table(doc):
    """Render the 3-tier subscription comparison table."""
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Feature"
    hdr[1].text = "Explorer (Free)"
    hdr[2].text = "Neighbourhood £5.99/mo"
    hdr[3].text = "Inner Circle £11.99/mo"
    for i, cell in enumerate(hdr):
        set_cell_bg(cell, "1A1A2E" if i == 0 else "FF6B35")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)

    rows = [
        ("Groups – join",              "2",          "Unlimited",    "Unlimited"),
        ("Groups – create",            "1",          "25",           "Unlimited"),
        ("Private groups",             "✗",          "✓",            "✓"),
        ("Meetups / month",            "2",          "Unlimited",    "Unlimited"),
        ("Direct messages",            "5",          "Unlimited",    "Unlimited"),
        ("Marketplace listings",       "2",          "15",           "Unlimited"),
        ("Photo uploads",              "3",          "15",           "50"),
        ("Messages / month",           "30",         "Unlimited",    "Unlimited"),
        ("Custom profile badge",       "✗",          "✓",            "✓"),
        ("AI Copilot chats / day",     "3",          "25",           "Unlimited"),
        ("AI Event Discovery / week",  "1",          "7",            "Unlimited"),
        ("AI Chat Summaries / day",    "✗",          "10",           "Unlimited"),
        ("AI Listing Generator /mo",   "✗",          "10",           "Unlimited"),
        ("AI Smart Feed refreshes",    "2/day",      "Unlimited",    "Unlimited"),
        ("AI Meetup Matchmaker",       "✗",          "✗",            "Unlimited"),
        ("AI Event Recommendations",   "✗",          "✓",            "✓"),
        ("Community Q&A / week",       "3",          "15",           "Unlimited"),
        ("Bookmarks / month",          "10",         "50",           "Unlimited"),
        ("Annual price",               "Free",       "£49.99/yr",    "£99.99/yr"),
    ]
    for feat, e, n, c in rows:
        row = table.add_row().cells
        row[0].text = feat
        row[1].text = e
        row[2].text = n
        row[3].text = c
        for j, cell in enumerate(row):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    if j == 0:
                        run.font.bold = True

# ══════════════════════════════════════════════════════════════════════════════
# 1. GO TO MARKET STRATEGY
# ══════════════════════════════════════════════════════════════════════════════

def create_gtm(path):
    doc = new_doc()
    cover_block(doc,
        "Huddl Connect",
        "Go To Market Strategy — The Digital Village for Modern Parents",
        VERSION, DATE,
        "Prepared by: Malgorzata D'Cruz, Co-Founder & CEO  |  Conrad D'Cruz, Co-Founder & Advisor")

    # ── 1. Executive Summary ──────────────────────────────────────────────────
    heading(doc, "1. Executive Summary", 1, ORANGE)
    body(doc,
        "Huddl Connect is a hyper-local community platform built for parents. "
        "Our singular mission is to replicate the 'next-door mum and dad' feeling "
        "in every borough across the UK — giving parents immediate access to their "
        "local tribe through genuine, borough-scoped connections.\n\n"
        "The app is live, built on Flutter/Firebase, and ships on iOS (TestFlight) "
        "and Android. It is not a concept — it is a working product with real "
        "features, real users, and a clear monetisation path via a three-tier "
        "subscription model.", size=10)

    heading(doc, "Key Stats & Context", 2)
    for b in [
        "860,000 parents in the UK feel lonely every single day (Oxford Academic, 2024)",
        "600,000 face serious mental health issues directly linked to isolation",
        "34% of new parents report feeling more isolated than before having children (AXA Health, 2025)",
        "£2.6B addressable market in UK parenting apps and community platforms",
        "No direct competitor combines hyper-local groups, AI, marketplace, and events in one app",
    ]:
        bullet(doc, b)

    # ── 2. What We Have Built ─────────────────────────────────────────────────
    heading(doc, "2. What We Have Built — The Product Today", 1, ORANGE)
    body(doc,
        "The following represents the live, shipped features of Huddl Connect. "
        "This is our spotlight — the real product investors can download and experience today.")

    heading(doc, "2.1 Five-Tab Navigation Architecture", 2)
    tabs = [
        ("Home",     "Personalised smart feed powered by invisible AI. Surfaces relevant groups, "
                     "meetups, events, marketplace items, and community announcements — all scoped "
                     "to the user's verified borough."),
        ("Connect",  "Borough-scoped group messaging. Users are auto-assigned to local groups on "
                     "sign-up based on postcode and parenting stage. Supports public and private "
                     "groups, polls, threads, saved messages, media attachments, and 1-to-1 DMs."),
        ("Discover", "Local events and parent-created meetups. AI discovers and surfaces events "
                     "happening in the user's borough daily. Parents can RSVP, create their own "
                     "meetups, and browse by category, date, and age group."),
        ("Market",   "Preloved marketplace for baby and children's items. Parents buy, sell, and "
                     "give away items locally. AI Listing Generator writes product descriptions "
                     "automatically. Direct messaging to sellers built in."),
        ("Profile",  "User profile, subscription management, backup/restore, and settings. "
                     "Shows tier badge, bio, parenting stage, and children's ages."),
    ]
    for tab, desc in tabs:
        bullet(doc, desc, bold_prefix=tab)

    heading(doc, "2.2 Borough-First Architecture (Core Differentiator)", 2)
    body(doc,
        "Every piece of content — groups, events, meetups, marketplace listings, "
        "and AI recommendations — is scoped to the user's verified postcode/borough. "
        "This is enforced at the service layer via BoroughScopeGuard, ensuring "
        "parents only see content relevant to their immediate community. "
        "This 'next-door mum and dad' feel is what separates Huddl from national "
        "platforms like Mumsnet or Facebook Groups.")

    heading(doc, "2.3 Onboarding Flow", 2)
    body(doc, "A frictionless, multi-step onboarding captures:")
    for b in [
        "Name and profile photo",
        "Parenting stage (aspiring, expecting, new parent, experienced)",
        "Children's ages and due dates",
        "Postcode (verified to borough level)",
        "Phone number (verified via SMS OTP / Firebase Auth)",
        "Parent type and bio",
    ]:
        bullet(doc, b)
    body(doc,
        "On completion, users are auto-assigned to up to 3 relevant local groups "
        "and shown a personalised welcome screen. A tutorial overlay walks through all 5 tabs.")

    heading(doc, "2.4 AI Features — Live in the App Today", 2)
    body(doc,
        "AI is embedded across every tab of the app as an invisible layer — "
        "not a chatbot bolted on, but intelligence woven into the experience. "
        "All AI is powered by Google Gemini via a structured prompt architecture.")

    ai_features = [
        ("AI Copilot",
         "Conversational parenting assistant. Answers questions about local resources, "
         "parenting challenges, child development, and borough-specific information. "
         "Context-aware — knows the user's borough, parenting stage, and children's ages. "
         "3 chats/day on Explorer, 25/day on Neighbourhood, unlimited on Inner Circle."),
        ("AI Smart Feed",
         "Personalised home feed that learns from user behaviour. Reorders content cards, "
         "surfaces timely nudges, and predicts what the parent needs next. "
         "Runs a daily refresh cycle."),
        ("AI Event Discovery",
         "Automated daily discovery of local events relevant to parents in the user's borough. "
         "Web-crawled and scored for family relevance. 1 discovery/week on Explorer, "
         "7/week on Neighbourhood, unlimited on Inner Circle."),
        ("AI Event Recommendations",
         "Personalised scoring of discovered events based on user profile — child ages, "
         "interests, parenting stage. Available on Neighbourhood and Inner Circle tiers."),
        ("AI Chat Summariser",
         "Catches users up on group conversations they have missed. Generates a concise "
         "summary of unread messages in any group. 10 summaries/day on Neighbourhood, "
         "unlimited on Inner Circle."),
        ("AI Listing Generator",
         "Writes marketplace listing titles and descriptions automatically from a photo "
         "and basic item details. 10 generations/month on Neighbourhood, unlimited on Inner Circle."),
        ("AI Meetup Matchmaker",
         "Scores parent compatibility and suggests optimal meetup pairings. "
         "Exclusive to Inner Circle tier. Unlimited requests/month."),
        ("Invisible AI Layer",
         "Predictive pre-fill across forms, adaptive content reordering, contextual "
         "intelligence throughout the app — users benefit from AI without seeing an "
         "AI interface. Powered by the AI Learning Engine and Knowledge Base services."),
        ("Borough AI Context",
         "All AI responses are grounded in the user's specific borough — local parks, "
         "schools, services, events, and community knowledge. Built via a dedicated "
         "borough context service and Gemini system prompt builder."),
    ]
    for name, desc in ai_features:
        bullet(doc, desc, bold_prefix=name)

    heading(doc, "2.5 Subscription Tiers — Live Pricing", 2)
    tier_table(doc)
    doc.add_paragraph()

    heading(doc, "2.6 Additional Built Features", 2)
    extra = [
        ("Group Polls",            "Create and vote on polls within any group chat."),
        ("Message Threading",      "Reply to specific messages in threads — keeps conversations organised."),
        ("Saved Messages",         "Bookmark messages across groups for later reference."),
        ("Media Attachments",      "Send images and attachments within group chats and DMs."),
        ("Message Search",         "Full-text search across all group messages."),
        ("Journey Map",            "Visual timeline of the user's parenting journey — milestones and memories."),
        ("Backup & Restore",       "Export and import user data for profile continuity."),
        ("Block / Report",         "Safety tools — block users and report content."),
        ("Tutorial Overlay",       "Interactive first-run walkthrough of all 5 tabs."),
        ("Biometric Auth",         "Face ID / fingerprint unlock on supported devices."),
        ("Borough Analytics",      "Internal analytics on borough engagement and group activity."),
        ("GDPR Controls",          "Full data export, deletion, and privacy settings compliant with UK GDPR."),
        ("Rehome Listings",        "Dedicated rehome journey for items being given away free to local families."),
    ]
    for name, desc in extra:
        bullet(doc, desc, bold_prefix=name)

    # ── 3. Target Market ──────────────────────────────────────────────────────
    heading(doc, "3. Target Market", 1, ORANGE)
    heading(doc, "3.1 Primary Audience", 2)
    for b in [
        "Parents aged 25–45 in UK urban/suburban boroughs",
        "New parents (0–2 years) — highest isolation, highest motivation to connect",
        "Expecting parents — preparing, seeking local community pre-birth",
        "Parents with children up to 16 years — ongoing community need",
        "All parent types: single parents, same-sex couples, blended families",
    ]:
        bullet(doc, b)

    heading(doc, "3.2 Initial Geography — Cambridge & Expansion", 2)
    body(doc,
        "Launch borough: Cambridge, UK. Borough-scoped architecture means each "
        "new city/borough is a self-contained community that activates independently. "
        "Expansion path: Cambridge → London boroughs → UK cities → International.")

    heading(doc, "3.3 Total Addressable Market", 2)
    tam = [
        ("UK parents (total)",          "~13 million"),
        ("UK smartphone parents",       "~11.5 million"),
        ("Addressable (community app)", "~4–6 million"),
        ("Revenue @ 5% penetration, avg £7/mo", "~£28M ARR"),
        ("Revenue @ 15% penetration",  "~£84M ARR"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Segment"
    hdr[1].text = "Estimate"
    for cell in hdr:
        set_cell_bg(cell, "FF6B35")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
    for seg, val in tam:
        row = table.add_row().cells
        row[0].text = seg
        row[1].text = val
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

    # ── 4. Revenue Model ──────────────────────────────────────────────────────
    heading(doc, "4. Revenue Model & Monetisation", 1, ORANGE)
    heading(doc, "4.1 Subscription Revenue (Primary)", 2)
    body(doc,
        "Three-tier freemium model. Explorer (free) provides a genuine taste "
        "of the community — enough to create real value and drive organic upgrades. "
        "Paid tiers unlock full AI suite, unlimited social features, and premium tools.")
    for b in [
        "Explorer: Free — 2 groups, 2 meetups/month, 3 AI copilot chats/day",
        "Neighbourhood: £5.99/month or £49.99/year — full AI suite, unlimited social",
        "Inner Circle: £11.99/month or £99.99/year — everything + AI Matchmaker, unlimited AI",
        "Annual plans save users ~30% vs monthly — drives LTV and reduces churn",
    ]:
        bullet(doc, b)

    heading(doc, "4.2 B2B Revenue (Secondary)", 2)
    for b in [
        "Featured event slots for local businesses (nurseries, play centres, classes)",
        "Sponsored group creation for local service providers",
        "Borough partnership packages for councils and NHS trusts",
        "API access for third-party integrations (future)",
    ]:
        bullet(doc, b)

    heading(doc, "4.3 Revenue Projections", 2)
    proj = [
        ("Year 1", "500 paid users", "~£43K ARR", "Product-market fit, Cambridge launch"),
        ("Year 2", "5,000 paid users", "~£430K ARR", "London borough expansion"),
        ("Year 3", "25,000 paid users", "~£2.1M ARR", "UK city rollout + B2B"),
        ("Year 4", "100,000 paid users", "~£8.4M ARR", "National scale + international"),
        ("Year 5", "400,000 paid users", "~£33M ARR", "Category leader"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(["Year", "Paid Users", "ARR", "Milestone"]):
        hdr[i].text = h
        set_cell_bg(hdr[i], "1A1A2E")
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
    for yr, users, arr, ms in proj:
        row = table.add_row().cells
        for i, val in enumerate([yr, users, arr, ms]):
            row[i].text = val
            for para in row[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

    # ── 5. Go-To-Market Strategy ──────────────────────────────────────────────
    heading(doc, "5. Go-To-Market Strategy", 1, ORANGE)
    heading(doc, "5.1 Phase 1 — Borough Seeding (Months 1–6)", 2)
    body(doc, "Focus: Cambridge. Goal: 1,000 active users, 200 paid.")
    for b in [
        "NCT group partnerships — partner with National Childbirth Trust Cambridge chapter",
        "Nursery and playgroup partnerships — promote via noticeboards and newsletters",
        "Mums & dads Facebook group migration — offer a better, safer alternative",
        "School gate presence — street marketing at primary school drop-offs",
        "NHS midwife and health visitor referral programme",
        "Local influencer seeding — parent bloggers and Instagram accounts in Cambridge",
        "Press outreach — Cambridge News, local parenting publications",
    ]:
        bullet(doc, b)

    heading(doc, "5.2 Phase 2 — Borough Replication (Months 6–18)", 2)
    for b in [
        "Activate 5 London boroughs using the same seeding playbook",
        "Launch B2B sales for featured events and sponsored content",
        "Referral programme — existing users invite neighbours (borough stays local)",
        "App Store Optimisation (ASO) — target 'local parents', 'mums near me' keywords",
        "Paid social — hyper-targeted Meta ads by postcode",
        "PR — national parenting press (MadeForMums, Netmums, Primary Times)",
    ]:
        bullet(doc, b)

    heading(doc, "5.3 Phase 3 — National Scale (Months 18–36)", 2)
    for b in [
        "UK city-by-city rollout using borough activation playbook",
        "Council and NHS trust B2B partnerships (safeguarding, health visiting, SEND)",
        "App Store featuring campaign — apply for editorial placement",
        "Corporate employee benefit packages (employer-sponsored subscriptions)",
        "International exploration — Ireland, Australia, Canada (English-speaking boroughs)",
    ]:
        bullet(doc, b)

    heading(doc, "5.4 Competitive Positioning", 2)
    comps = [
        ("WhatsApp/Facebook Groups", "National, not local; no AI; no safety; no structure"),
        ("Mumsnet",                  "Forum-based, not real-time; no local matching; aging UX"),
        ("Nextdoor",                 "Neighbours not parents; no parenting AI; no marketplace focus"),
        ("Peanut",                   "Dating-style swiping; national; no groups; no events; no marketplace"),
        ("Huddl Connect",            "Borough-first, AI-powered, full-stack parenting community — built"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(["Platform", "Limitation", ""]):
        hdr[i].text = h
        set_cell_bg(hdr[i], "FF6B35")
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
    for plat, lim, *_ in comps:
        row = table.add_row().cells
        row[0].text = plat
        row[1].text = lim
        row[2].text = "✓ Huddl wins" if plat == "Huddl Connect" else "✗ Gap"
        for j, cell in enumerate(row):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    if plat == "Huddl Connect":
                        run.font.color.rgb = GREEN
    doc.add_paragraph()

    # ── 6. Crowdfunding ───────────────────────────────────────────────────────
    heading(doc, "6. Crowdfunding & Fundraising Strategy", 1, ORANGE)
    heading(doc, "6.1 Raise Overview", 2)
    for b in [
        "Target raise: £250,000–£500,000 (seed/pre-seed)",
        "Platforms: Crowdcube (primary), Seedrs (secondary), angels",
        "Valuation basis: product built + TAM + team + early traction",
        "Equity offered: 10–15% at £3–5M pre-money valuation",
        "Use of funds: marketing/growth (50%), hiring (30%), infrastructure (20%)",
    ]:
        bullet(doc, b)

    heading(doc, "6.2 Campaign Process", 2)
    for b in [
        "Pre-launch: build crowd of 500+ registered interest (email list, socials)",
        "Soft-launch to existing users and warm network first — target 30% funded in week 1",
        "Public launch with PR push — parenting press, tech press, investor networks",
        "Video pitch: founders on camera, app demo, user testimonials",
        "Stretch goals: £350K = Android launch; £500K = 3 London boroughs; £750K = full UK",
    ]:
        bullet(doc, b)

    heading(doc, "6.3 Funding Sources", 2)
    for b in [
        "Crowdfunding community investors (target: 200+ backers)",
        "Angel investors — parenting sector, EdTech, community platform experience",
        "VC seed funds — UK-focused consumer social / family tech",
        "Grant funding — Innovate UK, Local Authority digital innovation grants",
        "Strategic investors — NHS trusts, children's charities, council partnerships",
    ]:
        bullet(doc, b)

    # ── 7. Team ───────────────────────────────────────────────────────────────
    heading(doc, "7. Team & Founders", 1, ORANGE)
    for b in [
        ("Malgorzata D'Cruz — Co-Founder & CEO",
         "Early Childhood Education (ECE) expert. Deep domain expertise in child development, "
         "parent support, and community building. Identified the parent loneliness gap first-hand. "
         "Drives product vision, community strategy, and user research."),
        ("Conrad D'Cruz — Co-Founder & Advisor",
         "MBA, Cambridge. Commercial and strategic leadership. Drives investor relations, "
         "business development, and growth strategy. Technical oversight of platform architecture."),
    ]:
        bullet(doc, b[1], bold_prefix=b[0])

    # ── 8. Strategic Advantages ───────────────────────────────────────────────
    heading(doc, "8. Strategic Advantages", 1, ORANGE)
    for b in [
        "Product is built — investors can download and use it today; no concept risk",
        "Borough-first architecture creates powerful network effects within local communities",
        "AI embedded throughout — not a feature, but the fabric of the experience",
        "Three subscription tiers with clear value laddering drives predictable SaaS revenue",
        "Founders have domain expertise (ECE) + commercial expertise (MBA Cambridge)",
        "No direct competitor combines all five pillars: local groups, AI, events, marketplace, profiles",
        "Firebase + Flutter tech stack — scalable, cross-platform, low infrastructure cost",
        "GDPR-compliant and Children's Code aware from day one",
    ]:
        bullet(doc, b)

    # ── APPENDIX: AI Roadmap ──────────────────────────────────────────────────
    doc.add_page_break()
    heading(doc, "APPENDIX A — AI Roadmap (Future Features)", 1, MID)
    body(doc,
        "The following AI capabilities are on the product roadmap and are NOT yet "
        "live in the current build. They represent the platform's long-term AI vision "
        "and form the basis of the scalability story for investors.", color=MID)
    doc.add_paragraph()

    roadmap = [
        ("Personalised Deal Recommendations",
         "Based on user profile (child ages, interests) — price drop alerts, seasonal surfacing"),
        ("Collaborative Filtering",
         "'Parents like you also saved...' — cross-borough recommendations"),
        ("Natural Language Search",
         "Search events, groups, marketplace items in plain English"),
        ("AI Moderation",
         "Automated content moderation and safety scoring for all posts and messages"),
        ("Predictive Churn Prevention",
         "Identify users at risk of churning and trigger personalised re-engagement"),
        ("AI Group Health Score",
         "Surface groups that are becoming inactive and prompt admins to re-engage"),
        ("AI Parenting Milestone Tracker",
         "Proactive milestone alerts and personalised advice for each stage"),
        ("Voice Input",
         "Hands-free parenting — voice commands for tired parents with hands full"),
        ("AI Trip Planner",
         "Full family holiday planning with AI concierge, packing lists, family-friendly ratings"),
        ("Multi-borough Matching",
         "Connect parents moving between boroughs with new local community immediately"),
        ("B2B AI Dashboard",
         "Analytics and AI-driven insights for local business partners"),
        ("100+ Additional AI Micro-Features",
         "Context-aware nudges, smart notifications, adaptive UI, predictive search, "
         "sentiment analysis on community health, AI-generated weekly digest emails"),
    ]
    for name, desc in roadmap:
        bullet(doc, desc, bold_prefix=name)

    doc.save(path)
    print(f"✓ GTM saved: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# 2. PITCH DECK (Word format)
# ══════════════════════════════════════════════════════════════════════════════

def create_pitch_deck(path):
    doc = new_doc()
    cover_block(doc,
        "Huddl Connect",
        "Investor Pitch Deck 2026 — Solving the Global Parent Loneliness Epidemic",
        VERSION, DATE,
        "Malgorzata D'Cruz, Co-Founder & CEO  |  Conrad D'Cruz, Co-Founder & Advisor")

    # Slide 1 — Title
    heading(doc, "Slide 1 — The Mission", 1, ORANGE)
    body(doc,
        "Huddl Connect is the hyper-local community platform built for all parents "
        "to find their local tribe — the 'next-door mum and dad' feeling, "
        "digitally delivered, borough by borough across the UK.", bold=True, size=12)
    body(doc, "Available on iOS (TestFlight) and Android. Built. Real. Working.")
    doc.add_paragraph()

    # Slide 2 — Problem
    heading(doc, "Slide 2 — The Problem: The Isolation Crisis", 1, ORANGE)
    stats = [
        ("860,000",  "Parents feel lonely every day in the UK alone"),
        ("600,000",  "Face serious mental health issues linked to isolation"),
        ("34%",      "Of new parents feel more isolated than before having children"),
        ("£2.6B",    "Addressable market for parenting community platforms in the UK"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Statistic"
    hdr[1].text = "Context"
    for cell in hdr:
        set_cell_bg(cell, "FF6B35")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
    for stat, ctx in stats:
        row = table.add_row().cells
        row[0].text = stat
        row[1].text = ctx
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    body(doc,
        "Despite Facebook, WhatsApp, and Mumsnet — parents remain isolated. "
        "These tools are national, noisy, and not built for local community. "
        "No platform solves the hyper-local parent connection problem.")

    # Slide 3 — Solution
    heading(doc, "Slide 3 — The Solution: Your Digital Village", 1, ORANGE)
    body(doc,
        "Huddl Connect scopes every interaction to your borough — your neighbours, "
        "your local events, your local marketplace. AI makes it effortless. "
        "The result: a platform that feels like the village that modern parents have lost.", bold=True)
    doc.add_paragraph()
    for b in [
        "Borough-locked content — only see parents in your area",
        "Auto-matched to local groups on sign-up based on postcode + parenting stage",
        "AI surfaces what you need before you search for it",
        "One app replaces WhatsApp groups + Facebook Marketplace + Eventbrite + Mumsnet",
    ]:
        bullet(doc, b)

    # Slide 4 — Product Demo
    heading(doc, "Slide 4 — The Product: 5 Tabs, One Community", 1, ORANGE)
    body(doc, "Live, working app. Download it. Use it. This is not a mockup.")
    doc.add_paragraph()
    tabs = [
        ("🏠 Home",     "AI-personalised smart feed. Borough announcements, nearby meetups, group activity, marketplace highlights — curated for you, updated daily."),
        ("💬 Connect",  "Borough-scoped group chats. Auto-assigned groups + user-created public/private groups. DMs, polls, threads, saved messages, media sharing."),
        ("🗺 Discover", "Local events + parent meetups. AI discovers events daily in your borough. RSVP, create your own meetup, browse by category and child age."),
        ("🛍 Market",   "Preloved marketplace. Buy, sell, rehome baby and children's items locally. AI writes your listing descriptions. Direct message sellers."),
        ("👤 Profile",  "Your parenting profile. Subscription management, journey map, backup & restore, biometric auth, GDPR data controls."),
    ]
    for tab, desc in tabs:
        bullet(doc, desc, bold_prefix=tab)

    # Slide 5 — AI Differentiator
    heading(doc, "Slide 5 — AI: Invisible Intelligence", 1, ORANGE)
    body(doc,
        "Huddl doesn't have 'an AI feature'. AI IS the product. "
        "Every tab is powered by Google Gemini-backed intelligence, "
        "running invisibly in the background to make every parent's experience "
        "feel effortless and personal.", bold=True)
    doc.add_paragraph()
    ai_live = [
        "AI Copilot — borough-aware parenting assistant (live)",
        "AI Smart Feed — personalised home feed with daily refresh (live)",
        "AI Event Discovery — daily local event crawl + scoring (live)",
        "AI Chat Summariser — catch up on missed group messages (live)",
        "AI Listing Generator — writes marketplace listings from a photo (live)",
        "AI Meetup Matchmaker — parent compatibility scoring (live, Inner Circle)",
        "AI Event Recommender — personalised event scoring (live)",
        "Invisible AI Layer — predictive pre-fill, adaptive reordering (live)",
        "Borough AI Context — all AI grounded in local knowledge (live)",
    ]
    for item in ai_live:
        bullet(doc, item)

    # Slide 6 — Subscription & Revenue
    heading(doc, "Slide 6 — Business Model: SaaS Subscriptions", 1, ORANGE)
    body(doc, "Three clear tiers. Free to get in. Easy to upgrade. Sticky once in.", bold=True)
    doc.add_paragraph()
    tier_table(doc)
    doc.add_paragraph()
    body(doc, "Annual plans (£49.99 / £99.99) reduce churn and increase LTV by ~30%.")

    # Slide 7 — Market Opportunity
    heading(doc, "Slide 7 — Market Opportunity", 1, ORANGE)
    for b in [
        "13M UK parents — 11.5M smartphone users",
        "~4–6M addressable for a hyper-local community app",
        "Global: 2B+ parents worldwide — same isolation problem, same solution",
        "£2.6B UK TAM growing at ~18% CAGR driven by community + AI",
        "Adjacent revenue: B2B partnerships, council contracts, employer benefits",
    ]:
        bullet(doc, b)

    # Slide 8 — Traction
    heading(doc, "Slide 8 — Traction & Validation", 1, ORANGE)
    for b in [
        "App is live and available on iOS (TestFlight) and Android",
        "Full feature set built: 5 tabs, AI suite, subscription billing, onboarding",
        "Borough-scoped architecture tested in Cambridge",
        "Firebase backend with real user authentication (phone OTP)",
        "Subscription infrastructure built on Apple/Google IAP",
        "54 services built covering AI, community, payments, GDPR, analytics",
        "Positive user feedback from TestFlight cohort on UX and local feel",
    ]:
        bullet(doc, b)

    # Slide 9 — Go To Market
    heading(doc, "Slide 9 — Go To Market", 1, ORANGE)
    phases = [
        ("Phase 1 (Months 1–6)",
         "Cambridge borough seeding. NCT partnerships, nursery/school gate marketing, "
         "NHS midwife referrals, local influencers. Target: 1,000 active, 200 paid."),
        ("Phase 2 (Months 6–18)",
         "5 London boroughs. B2B sales launch. Referral programme. Paid social by postcode. "
         "Target: 10,000 active, 2,000 paid."),
        ("Phase 3 (Months 18–36)",
         "UK city rollout. Council + NHS B2B. National press. Corporate benefits. "
         "Target: 50,000 active, 10,000 paid."),
    ]
    for phase, desc in phases:
        bullet(doc, desc, bold_prefix=phase)

    # Slide 10 — Competitive Landscape
    heading(doc, "Slide 10 — Why Now, Why Us", 1, ORANGE)
    body(doc,
        "The market has failed parents. Every existing platform is national, "
        "noisy, or single-purpose. No competitor has built borough-first, "
        "AI-native, full-stack parenting community — until Huddl.", bold=True)
    doc.add_paragraph()
    for b in [
        "WhatsApp/Facebook — national, no AI, no safety, no structure",
        "Mumsnet — forum, not real-time, no local matching",
        "Nextdoor — neighbours not parents, no parenting AI",
        "Peanut — dating-style, national, no groups/events/marketplace",
        "Huddl — borough-first + AI-native + full-stack = unfair advantage",
    ]:
        bullet(doc, b)

    # Slide 11 — Financial Projections
    heading(doc, "Slide 11 — Financial Projections", 1, ORANGE)
    proj = [
        ("Year 1", "500 paid",    "£43K ARR",   "Cambridge PMF"),
        ("Year 2", "5,000 paid",  "£430K ARR",  "London expansion"),
        ("Year 3", "25,000 paid", "£2.1M ARR",  "UK rollout + B2B"),
        ("Year 4", "100K paid",   "£8.4M ARR",  "National scale"),
        ("Year 5", "400K paid",   "£33M ARR",   "Category leader"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(["Year", "Users", "ARR", "Stage"]):
        hdr[i].text = h
        set_cell_bg(hdr[i], "1A1A2E")
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
    for yr, u, arr, stage in proj:
        row = table.add_row().cells
        for i, val in enumerate([yr, u, arr, stage]):
            row[i].text = val
            for para in row[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

    # Slide 12 — Use of Funds
    heading(doc, "Slide 12 — Use of Funds", 1, ORANGE)
    body(doc, "Target raise: £250,000–£500,000 (seed/pre-seed)")
    doc.add_paragraph()
    for b in [
        "50% — Marketing & growth (borough seeding, paid social, partnerships)",
        "30% — Team (community manager, growth lead, part-time engineering)",
        "20% — Infrastructure, legal, compliance (GDPR, App Store fees, Firebase)",
    ]:
        bullet(doc, b)
    doc.add_paragraph()
    body(doc,
        "Stretch milestones: £350K = full Android launch; "
        "£500K = 3 London boroughs activated; £750K = full UK rollout plan")

    # Slide 13 — Valuation
    heading(doc, "Slide 13 — Valuation & Deal", 1, ORANGE)
    for b in [
        "Pre-money valuation: £3M–£5M",
        "Equity offered: 10–15%",
        "Comparable: Peanut raised $12M Series A at ~£30M valuation (product-only, no AI)",
        "Huddl advantage: full product built + AI native + borough architecture + UK-first",
        "Crowdfunding platform: Crowdcube (primary), Seedrs (secondary)",
    ]:
        bullet(doc, b)

    # Slide 14 — Vision
    heading(doc, "Slide 14 — The Vision: A Borough in Every Pocket", 1, ORANGE)
    body(doc,
        "In 5 years, every parent in the UK opens Huddl when they want to know "
        "what's happening near them, who to connect with, or what to sell/buy locally. "
        "The borough becomes the unit of community. AI makes it feel personal. "
        "The network effect locks in the moat.", bold=True, size=11)
    doc.add_paragraph()
    body(doc,
        "International: same borough-first model deployed in Ireland, Australia, Canada. "
        "B2B: councils, NHS, employers pay for borough-level insights and community engagement. "
        "Platform: open API for local service providers to build on Huddl's community graph.")

    # Slide 15 — Team (UPDATED: both are Co-Founders)
    heading(doc, "Slide 15 — The Team", 1, ORANGE)
    body(doc, "Two founders. Complementary expertise. Shared mission.", bold=True)
    doc.add_paragraph()

    founders = [
        ("Malgorzata D'Cruz",
         "Co-Founder & CEO",
         "Early Childhood Education (ECE) Expert. Deep domain expertise in child development, "
         "parent wellbeing, and community building. Identified the parent isolation crisis "
         "through direct professional experience. Leads product vision, community strategy, "
         "user research, and partnerships."),
        ("Conrad D'Cruz",
         "Co-Founder & Advisor",
         "MBA, Cambridge. Commercial and strategic leader with experience in business development "
         "and growth strategy. Leads investor relations, B2B partnerships, and platform "
         "commercialisation. Technical oversight of architecture and scaling strategy."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    for i, (name, title, bio) in enumerate(founders):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, "1A1A2E")
        p = cell.paragraphs[0]
        r = p.add_run(name)
        r.font.bold = True
        r.font.color.rgb = ORANGE
        r.font.size = Pt(12)
        p2 = cell.add_paragraph()
        r2 = p2.add_run(title)
        r2.font.color.rgb = WHITE
        r2.font.bold = True
        r2.font.size = Pt(10)
        p3 = cell.add_paragraph()
        r3 = p3.add_run(bio)
        r3.font.color.rgb = WHITE
        r3.font.size = Pt(9)
    doc.add_paragraph()

    # Slide 16 — Call to Action
    heading(doc, "Slide 16 — Join the Village", 1, ORANGE)
    body(doc,
        "The app is built. The market is massive. The team is focused. "
        "The timing is right.", bold=True, size=12)
    doc.add_paragraph()
    for b in [
        "Download Huddl Connect today and experience it yourself",
        "Join our crowdfunding campaign — invest in the future of local parenting communities",
        "Partner with us — borough activations, B2B, council programmes",
        "Contact: Malgorzata D'Cruz — malgorzata@huddlconnect.com",
        "Website: www.huddlconnect.com",
    ]:
        bullet(doc, b)
    doc.add_paragraph()
    body(doc,
        "© 2026 Huddl Connect. All rights reserved. "
        "Sources: [1] Oxford Academic (2024) [2] AXA Health National Parent Survey (2025)",
        color=MID, size=8)

    # APPENDIX: AI Roadmap
    doc.add_page_break()
    heading(doc, "APPENDIX A — AI Roadmap (Future Vision)", 1, MID)
    body(doc,
        "The features below are NOT yet live. They represent the long-term AI roadmap "
        "and form the basis of the platform's scalability and moat argument.", color=MID)
    doc.add_paragraph()
    future = [
        ("AI Trip Planner",          "Full family holiday AI concierge with packing lists and family ratings"),
        ("Natural Language Search",  "Search groups, events, marketplace in plain English"),
        ("AI Moderation",            "Automated content safety scoring across all posts and messages"),
        ("Predictive Churn Engine",  "Identify at-risk users, trigger personalised re-engagement"),
        ("AI Group Health Score",    "Alert admins when groups go quiet; suggest re-engagement tactics"),
        ("AI Milestone Tracker",     "Proactive parenting milestone alerts with personalised advice"),
        ("Voice Input",              "Hands-free mode for busy parents"),
        ("Multi-borough Matching",   "Help parents moving area connect with their new community immediately"),
        ("B2B AI Dashboard",         "Borough-level analytics and insights for local business partners"),
        ("Collaborative Filtering",  "'Parents like you also joined...' cross-borough discovery"),
        ("100+ AI Micro-Features",   "Context-aware nudges, smart notifications, adaptive UI, sentiment analysis"),
    ]
    for name, desc in future:
        bullet(doc, desc, bold_prefix=name)

    doc.save(path)
    print(f"✓ Pitch Deck saved: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# 3. AI SYSTEM DESIGN DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

def create_ai_doc(path):
    doc = new_doc()
    cover_block(doc, "Huddl Connect",
        "AI System Design Document",
        VERSION, DATE, "Classification: Confidential — Internal Engineering")

    heading(doc, "1. Executive Summary", 1, ORANGE)
    body(doc,
        "Huddl Connect embeds AI throughout the user experience as an invisible, "
        "contextual intelligence layer. AI is not a standalone feature — it is woven "
        "into every tab. All AI is powered by Google Gemini via the Gemini API, "
        "orchestrated through a structured prompt-building architecture and "
        "per-tier usage gating. The app does NOT use RevGlue or any affiliate "
        "deals/offers AI — that functionality has been removed.")

    heading(doc, "2. AI Feature Inventory — Live Features", 1, ORANGE)
    features = [
        ("AI Copilot",             "ai_copilot_service.dart",         "Conversational parenting assistant. Borough-aware, persona-aware, child-age-aware. Uses gemini_system_prompt_builder.dart for context injection."),
        ("AI Smart Feed",          "ai_feed_service.dart",            "Personalised home feed curation. Reorders content cards by predicted relevance. Daily refresh cycle via daily_ai_refresh_service.dart."),
        ("AI Event Discovery",     "ai_event_discovery_service.dart", "Daily web-crawl of local events in the user's borough. Scores for family relevance. Results cached by borough."),
        ("AI Event Recommender",   "ai_event_recommender_service.dart","Personalised event scoring based on child ages, interests, parenting stage."),
        ("AI Chat Summariser",     "ai_chat_summariser_service.dart", "Generates concise summaries of unread group messages. Respects per-tier daily limits."),
        ("AI Listing Generator",   "ai_listing_service.dart",         "Generates marketplace listing titles and descriptions from item details and photos. Per-tier monthly limit."),
        ("AI Meetup Matchmaker",   "ai_matchmaker_service.dart",      "Parent compatibility scoring. Suggests optimal meetup pairings. Inner Circle only."),
        ("Invisible AI Layer",     "invisible_ai_service.dart",       "Predictive pre-fill, adaptive reordering, contextual intelligence throughout the app."),
        ("AI Learning Engine",     "ai_learning_engine_service.dart", "Learns from user behaviour to improve recommendations over time."),
        ("AI Knowledge Base",      "ai_knowledge_base_service.dart",  "Borough-specific local knowledge base powering all AI context responses."),
        ("Borough AI Context",     "borough_ai_context.dart",         "Injects borough-specific context into all Gemini prompts. Local parks, schools, services."),
        ("Discover AI",            "discover_ai_service.dart",        "AI layer for the Discover tab — event scoring, meetup recommendations."),
        ("Meetup AI",              "meetup_ai_service.dart",          "AI assistance for meetup creation and participant matching."),
        ("Messages AI",            "messages_ai_service.dart",        "In-chat AI assistance — smart replies, message insights."),
        ("Gemini Prompt Builder",  "gemini_system_prompt_builder.dart","Central service constructing all Gemini system prompts with user context."),
        ("AI API Helper",          "ai_api_helper.dart",              "Shared helper for Gemini API calls with error handling and retry logic."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(["AI Feature", "Service File", "Description"]):
        hdr[i].text = h
        set_cell_bg(hdr[i], "FF6B35")
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
    for name, file, desc in features:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = file
        row[2].text = desc
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph()

    heading(doc, "3. AI Usage Limits by Tier", 1, ORANGE)
    tier_table(doc)
    doc.add_paragraph()

    heading(doc, "4. Architecture", 1, ORANGE)
    heading(doc, "4.1 Gemini Integration", 2)
    for b in [
        "Model: Google Gemini (via Gemini API / Vertex AI)",
        "Prompt construction: gemini_system_prompt_builder.dart injects user borough, parenting stage, children's ages, and tier",
        "All prompts are structured with system context + user message",
        "Borough knowledge injected via borough_ai_context.dart and borough_cache_service.dart",
        "API calls centralised through ai_api_helper.dart with retry and error handling",
        "All AI usage tracked by SubscriptionService for per-tier gating",
    ]:
        bullet(doc, b)

    heading(doc, "4.2 Borough-Scoped AI", 2)
    body(doc,
        "Every AI call is scoped to the user's verified borough. "
        "borough_scope_guard.dart enforces that no AI response surfaces "
        "content, events, or people outside the user's borough. "
        "borough_cache_service.dart caches borough-level data to "
        "reduce API calls and improve response time.")

    heading(doc, "4.3 Daily Refresh Architecture", 2)
    body(doc,
        "daily_ai_refresh_service.dart runs a background refresh cycle that: "
        "re-scores the smart feed, runs event discovery for the user's borough, "
        "updates group recommendations, and pre-fetches AI responses for "
        "likely next actions. This makes the 'invisible AI' feel instant.")

    heading(doc, "5. Data Privacy & Ethics", 1, ORANGE)
    for b in [
        "All AI features comply with UK GDPR and the Children's Code",
        "User data processed locally where possible (on-device model inference for low-sensitivity tasks)",
        "Gemini API calls include only anonymised, non-PII context (borough, stage, ages — never name/phone)",
        "Users can opt out of AI personalisation via Profile > Settings",
        "AI-generated content is labelled with subtle 'AI suggested' indicators",
        "No AI data sold to third parties or used for advertising profiling",
        "GDPR data export and deletion includes AI preference data",
        "No RevGlue, no affiliate tracking, no deals/offers AI — removed from platform",
    ]:
        bullet(doc, b)

    heading(doc, "6. Removed Features (No Longer in App)", 1, ORANGE)
    body(doc,
        "The following were previously documented but have been removed from the platform:", color=MID)
    for b in [
        "RevGlue affiliate integration (deals, coupons, offers tab) — REMOVED",
        "Deals/Offers tab — REMOVED from navigation",
        "AI Offers Service (ai_offers_service.dart) — service file exists but is not wired to any UI",
        "Trips/Family Travel tab — REMOVED from main navigation (service files remain as future capability)",
        "7-tab navigation — REPLACED by 5-tab navigation (Home, Connect, Discover, Market, Profile)",
    ]:
        bullet(doc, b)

    doc.save(path)
    print(f"✓ AI Doc saved: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# 4. UX DESIGN DOCUMENTATION
# ══════════════════════════════════════════════════════════════════════════════

def create_ux_doc(path):
    doc = new_doc()
    cover_block(doc, "Huddl Connect",
        "UX Design Documentation",
        VERSION, DATE, "Classification: Confidential — Internal Design & Engineering")

    heading(doc, "1. Executive Summary", 1, ORANGE)
    body(doc,
        "Huddl Connect is designed as a warm, approachable mobile experience for parents. "
        "The UX prioritises simplicity, trust, and the feeling of belonging to a local community. "
        "The design philosophy is 'next-door mum and dad' — every screen should feel "
        "like it was made by a neighbour, not a corporation. AI is invisible — it works "
        "for the user without demanding attention.")

    heading(doc, "2. Design Principles", 1, ORANGE)
    principles = [
        ("Borough First",        "Every interaction is rooted in the user's local community. Content, people, events, and AI responses are all scoped to their verified borough."),
        ("Invisible AI",         "AI improves the experience without asking the user to interact with it. Smart feed, auto-assigned groups, predictive content — all seamless."),
        ("Trust First",          "Safety and verification are foundational. Phone OTP verification, block/report tools, moderated content, and transparent data practices."),
        ("Reduce Cognitive Load","Parents are busy. Surface the right content at the right time. Minimal navigation, clear hierarchy, no cognitive clutter."),
        ("Warm & Human",         "The visual language is warm orange, friendly typography, rounded corners. Not a cold tech product — a community home."),
        ("Progressive Disclosure","Start simple. Reveal depth as users explore. Subscription upgrades surface naturally at the point of need, not intrusively."),
    ]
    for name, desc in principles:
        bullet(doc, desc, bold_prefix=name)

    heading(doc, "3. Information Architecture", 1, ORANGE)
    heading(doc, "3.1 Navigation Model — 5-Tab Bottom Bar", 2)
    body(doc,
        "The bottom navigation uses a custom floating pill-shaped bar with 5 tabs. "
        "Each tab has an outlined (inactive) and filled (active) icon variant. "
        "The bar floats above content with a subtle shadow — does not obstruct content. "
        "Haptic feedback on tab switch.")
    doc.add_paragraph()
    tabs = [
        ("0 — Home",     "HomeScreen",        "AI smart feed, borough announcements, nearby activity"),
        ("1 — Connect",  "GroupsScreen",       "Group chats, DMs, polls, threads"),
        ("2 — Discover", "EventsScreen",       "Local events, parent meetups, RSVP"),
        ("3 — Market",   "MarketplaceScreen",  "Buy/sell preloved, rehome listings"),
        ("4 — Profile",  "ProfileScreen",      "User profile, subscription, settings"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(["Tab", "Screen", "Purpose"]):
        hdr[i].text = h
        set_cell_bg(hdr[i], "FF6B35")
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
    for tab, screen, purpose in tabs:
        row = table.add_row().cells
        row[0].text = tab
        row[1].text = screen
        row[2].text = purpose
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

    heading(doc, "3.2 Onboarding Flow", 2)
    onboarding = [
        "SplashScreen — brand entry, Firebase init",
        "OnboardingCarouselScreen — product introduction (swipeable)",
        "NameInputScreen — first name and surname",
        "PhoneNumberScreen — country code picker + phone entry",
        "PasswordScreen — secure password entry",
        "VerificationScreen — SMS OTP verification (Firebase Auth)",
        "ParentTypeScreen — select parent type (mum, dad, same-sex, single, etc.)",
        "StageOfLifeScreen — parenting stage selection",
        "ChildInfoScreen — add children's ages",
        "DueDateScreen — due date entry (for expecting parents)",
        "PostcodeScreen — postcode entry → borough resolution",
        "AboutYouScreen — bio and interests",
        "AddPhotoScreen — optional profile photo",
        "WelcomeCompleteScreen — auto-assigned groups shown, tutorial triggered",
    ]
    for step in onboarding:
        bullet(doc, step)

    heading(doc, "3.3 Complete Screen Inventory", 2)
    screens = [
        ("Onboarding",    "splash, carousel, name, phone, password, verification, parent type, stage of life, child info, due date, postcode, about you, add photo, welcome complete"),
        ("Home",          "HomeScreen (smart feed), JourneyMapScreen (milestone timeline)"),
        ("Connect",       "GroupsScreen, GroupChatScreen, GroupDetailsScreen, GroupMembersScreen, CreateGroupScreen, DMChatScreen, NewDMScreen, ThreadReplyScreen, SavedMessagesScreen, GroupPollsScreen, PollDetailScreen, CreatePollScreen, ForwardMessageSheet, ImageGalleryPicker"),
        ("Discover",      "EventsScreen, EventDetailScreen, CreateEventScreen, MeetupDetailScreen, CreateMeetupScreen"),
        ("Market",        "MarketplaceScreen, ItemDetailScreen, CreateListingScreen (Rehome), RehomeHomeScreen"),
        ("Profile",       "ProfileScreen, ManageSubscriptionScreen, SubscriptionPlansScreen, SubscriptionCheckoutScreen, BackupRestoreScreen"),
        ("AI",            "AiCopilotScreen, AiListingGeneratorSheet, AiMatchmakerSheet"),
        ("Legal",         "PrivacyPolicyScreen, TermsOfServiceScreen"),
        ("Auth",          "Login/Register flows via onboarding"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(["Section", "Screens"]):
        hdr[i].text = h
        set_cell_bg(hdr[i], "1A1A2E")
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
    for section, scr in screens:
        row = table.add_row().cells
        row[0].text = section
        row[1].text = scr
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph()

    heading(doc, "4. Design System", 1, ORANGE)
    heading(doc, "4.1 Colour Palette", 2)
    colours = [
        ("Primary (Orange)",       "#FF6B35", "CTAs, active nav, headings, brand moments"),
        ("Dark Navy",              "#1A1A2E", "Table headers, dark backgrounds"),
        ("Mid Grey-Blue",          "#444466", "Subtitles, secondary text"),
        ("Success Green",          "#2E7D32", "Confirmations, positive states"),
        ("Error Red",              "#C62828", "Errors, warnings"),
        ("Background (White)",     "#FFFFFF", "Screen backgrounds"),
        ("Surface",                "#F5F5F5", "Card backgrounds, input fields"),
        ("Disabled",               "#BDBDBD", "Disabled buttons, inactive elements"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(["Role", "Hex", "Usage"]):
        hdr[i].text = h
        set_cell_bg(hdr[i], "FF6B35")
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
    for role, hex_c, usage in colours:
        row = table.add_row().cells
        row[0].text = role
        row[1].text = hex_c
        row[2].text = usage
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

    heading(doc, "4.2 Typography", 2)
    for b in [
        "Primary font: Google Fonts — Poppins (headings, nav labels, CTAs)",
        "Body font: System default / Poppins (body text, descriptions)",
        "Heading 1: 26–28pt, Bold, Dark Navy or Orange",
        "Heading 2: 18–20pt, SemiBold",
        "Body: 14–16pt, Regular",
        "Caption: 10–12pt, Regular, Mid Grey",
        "Nav labels: 9pt, SemiBold (active) / Regular (inactive)",
    ]:
        bullet(doc, b)

    heading(doc, "4.3 Component Patterns", 2)
    for b in [
        "Cards: 16px border radius, subtle shadow (blur 24, spread 0, offset Y4), white background",
        "Buttons: 12px border radius, 54pt height, full-width primary CTAs",
        "Input fields: Bottom-border only style on onboarding screens; full border on forms",
        "Bottom nav: Floating pill, 70pt height, 28px border radius, 16px horizontal padding",
        "FAB: Orange, 56px, bottom-right, used in Connect and Discover tabs",
        "Upgrade prompts: Orange accent, non-blocking bottom sheet, clear value proposition",
        "AI indicator: Subtle sparkle icon (✨) or 'AI suggested' label in caption style",
        "Borough badge: Pill-shaped, orange outline, shows verified borough name",
    ]:
        bullet(doc, b)

    heading(doc, "5. Onboarding Tutorial", 1, ORANGE)
    body(doc,
        "First-run tutorial overlay (TutorialOverlay widget) walks users through all 5 tabs. "
        "Triggered on WelcomeCompleteScreen completion. Can be re-triggered from Profile. "
        "Swipeable card interface with tab switching via onTabSwitch callback.")

    heading(doc, "6. Subscription UX", 1, ORANGE)
    body(doc,
        "Upgrade prompts are contextual — they appear at the exact moment a user "
        "hits a tier limit (e.g. tries to join a 3rd group on Explorer). "
        "They are non-blocking (bottom sheet), show the specific benefit unlocked, "
        "and present both monthly and annual price options with annual savings highlighted.")
    for b in [
        "SubscriptionPlansScreen — full tier comparison",
        "SubscriptionCheckoutScreen — IAP payment flow (Apple/Google)",
        "ManageSubscriptionScreen — view current plan, scheduled changes, cancel",
        "Upgrade prompt widget — inline, contextual, non-blocking",
    ]:
        bullet(doc, b)

    heading(doc, "7. Accessibility", 1, ORANGE)
    for b in [
        "Minimum touch target: 44×44px for all interactive elements",
        "Colour contrast: WCAG 2.1 AA minimum (4.5:1 for text)",
        "Screen reader support: Semantic labels on all icons and buttons",
        "Font scaling: Respects system text size preferences",
        "Reduced motion: Animations respect system accessibility settings",
        "Haptic feedback: Tactile confirmation on nav switches, button presses",
    ]:
        bullet(doc, b)

    heading(doc, "8. Removed UI Elements (No Longer Present)", 1, ORANGE)
    body(doc, "The following UI elements have been removed and must not be referenced:", color=MID)
    for b in [
        "Deals/Offers tab — REMOVED. Was tab 5 of a previous 7-tab navigation. No longer exists.",
        "Trips tab — REMOVED. Was tab 6 of a previous 7-tab navigation. No longer exists.",
        "RevGlue store cards, coupon modals, and affiliate CTAs — REMOVED",
        "7-tab navigation — REPLACED by 5-tab navigation",
        "NestedScrollView with pinned TabBar for Deals — REMOVED",
    ]:
        bullet(doc, b)

    doc.save(path)
    print(f"✓ UX Doc saved: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# 5. FUNCTIONAL DESIGN DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

def create_fdd(path):
    doc = new_doc()
    cover_block(doc, "Huddl Connect",
        "Functional Design Document (FDD)",
        VERSION, DATE, "Classification: Confidential — Internal Engineering & Product")

    heading(doc, "1. Introduction & Purpose", 1, ORANGE)
    body(doc,
        "Huddl Connect is a mobile-first community platform for parents in the UK. "
        "It combines group messaging, local events, AI-powered discovery, a preloved "
        "marketplace, and subscription-based AI tools — all scoped to the user's "
        "verified borough. This document covers all user-facing features, business "
        "rules, and subscription tiers. RevGlue/Deals/Offers functionality has been "
        "removed from the platform entirely.")

    heading(doc, "2. Product Overview", 1, ORANGE)
    heading(doc, "2.1 Core Value Proposition", 2)
    body(doc,
        "A single app that replaces fragmented parenting tools: WhatsApp groups, "
        "Facebook Marketplace, Eventbrite, and Mumsnet — with a borough-first, "
        "AI-native experience that feels like your local community in your pocket.")

    heading(doc, "2.2 Navigation Architecture — 5 Tabs", 2)
    for b in ["Home (AI smart feed)", "Connect (groups + DMs)", "Discover (events + meetups)", "Market (marketplace + rehome)", "Profile (account + subscription)"]:
        bullet(doc, b)

    heading(doc, "2.3 Platform Targets", 2)
    for b in ["iOS — Apple App Store (TestFlight live)", "Android — Google Play Store", "Web — preview/testing only"]:
        bullet(doc, b)

    heading(doc, "2.4 Technology Stack", 2)
    for b in [
        "Frontend: Flutter 3.35.4 / Dart 3.9.2",
        "Backend: Firebase (Firestore, Auth, Storage, Crashlytics)",
        "AI: Google Gemini API",
        "Auth: Firebase Phone Auth (SMS OTP)",
        "Payments: Apple IAP + Google Play Billing",
        "State management: Provider pattern",
        "Local storage: Hive + shared_preferences",
    ]:
        bullet(doc, b)

    heading(doc, "3. User Personas", 1, ORANGE)
    personas = [
        ("Aspiring Parent — Alex",      "Researching parenthood, joining groups for advice. Uses AI Copilot for questions."),
        ("Expecting Parent — Sarah",    "Preparing for baby. Finds local groups, browses meetups, uses marketplace for essentials."),
        ("New Parent — Marcus",         "Sleep-deprived, eager to connect. Heavy Connect and Discover user. Uses AI Chat Summariser."),
        ("Experienced Parent — Priya",  "Multiple children. Sells on marketplace, organises meetups, heavy AI power user."),
        ("Local Service Provider — Dr Chen", "Offers parenting classes. Uses Events to promote. B2B partner candidate."),
    ]
    for name, desc in personas:
        bullet(doc, desc, bold_prefix=name)

    heading(doc, "4. Feature Specifications", 1, ORANGE)

    heading(doc, "4.1 Home (AI Smart Feed)", 2)
    for b in [
        "AI-personalised content feed — borough-scoped, updated daily",
        "Borough announcements and community news",
        "Nearby upcoming meetups and events",
        "Group activity highlights",
        "Marketplace item previews",
        "AI nudge cards — contextual suggestions based on user behaviour",
        "Journey Map access — visual parenting milestone timeline",
        "Quick action shortcuts (create meetup, list item)",
    ]:
        bullet(doc, b)

    heading(doc, "4.2 Connect (Groups & DMs)", 2)
    for b in [
        "Auto-assigned groups based on postcode and parenting stage (up to 3 on sign-up)",
        "Browse and join public borough groups",
        "Create public or private groups (Neighbourhood+ only for private)",
        "Group chat with real-time messaging via Firebase",
        "Message threading — reply to specific messages",
        "Polls — create and vote within groups",
        "Saved messages — bookmark messages for later",
        "Media attachments — send images within chats",
        "Message search — full-text search across all groups",
        "1-to-1 Direct Messages with online status indicator",
        "Forward messages to other groups or DMs",
        "Block and report users",
        "AI Chat Summariser — catch up on missed conversations (Neighbourhood+)",
    ]:
        bullet(doc, b)

    heading(doc, "4.3 Discover (Events & Meetups)", 2)
    for b in [
        "AI-discovered local events — daily web scrape scoped to user's borough",
        "B2B partner events from local businesses and service providers",
        "Parent-created meetups with RSVP functionality",
        "Filters: date range, age group, category, area",
        "Event detail: description, location, organiser, RSVP count",
        "Create Meetup screen with title, description, date/time, location, max attendees",
        "Create Event screen for business/partner events",
        "AI Event Recommender — personalised event scoring (Neighbourhood+)",
    ]:
        bullet(doc, b)

    heading(doc, "4.4 Market (Preloved + Rehome)", 2)
    for b in [
        "Browse and search preloved baby and children's items",
        "Create listing — title, description, price, category, photos (up to 5)",
        "AI Listing Generator — auto-generates title and description (Neighbourhood+)",
        "Direct message to seller from listing detail",
        "Rehome listings — give away items free to local families",
        "Rehome Journey Map — track items given away",
        "Category browsing: clothing, toys, equipment, furniture, books, etc.",
        "Saved/bookmarked listings",
        "My listings management",
    ]:
        bullet(doc, b)

    heading(doc, "4.5 Profile & Account", 2)
    for b in [
        "User profile: name, photo, bio, parenting stage, children, borough",
        "Edit profile",
        "Subscription management (view tier, upgrade, cancel, scheduled changes)",
        "Backup and restore user data (JSON export/import)",
        "Biometric authentication (Face ID / fingerprint)",
        "GDPR data export and deletion",
        "Tutorial re-trigger",
        "Legal: Privacy Policy, Terms of Service",
        "Tier badge display",
    ]:
        bullet(doc, b)

    heading(doc, "4.6 AI Copilot", 2)
    for b in [
        "Conversational parenting assistant — answers questions about local resources, child development, parenting challenges",
        "Borough-aware — responses reference the user's specific borough",
        "Persona-aware — knows parenting stage and children's ages",
        "Powered by Google Gemini with structured system prompts",
        "3 chats/day on Explorer, 25/day on Neighbourhood, unlimited on Inner Circle",
        "Accessible via home screen sparkle button and AI tab",
    ]:
        bullet(doc, b)

    heading(doc, "5. Subscription Model", 1, ORANGE)
    tier_table(doc)
    doc.add_paragraph()

    heading(doc, "5.1 Billing Rules", 2)
    for b in [
        "Explorer (free) activates immediately on sign-up — no card required",
        "Upgrading from Explorer activates new tier immediately via IAP",
        "Switching between paid tiers is scheduled — new tier activates on next renewal date",
        "Cancellation sets cancelledAtPeriodEnd = true; access continues until renewal date",
        "Annual plans show ~30% saving vs monthly equivalent",
        "7-day free trial on first paid subscription",
    ]:
        bullet(doc, b)

    heading(doc, "6. Removed Features", 1, ORANGE)
    body(doc, "The following features have been removed and are NOT part of the current product:", color=MID)
    for b in [
        "RevGlue affiliate integration — REMOVED",
        "Deals/Offers tab — REMOVED",
        "Trips/Family Travel tab — REMOVED from main navigation",
        "7-tab navigation — replaced by 5-tab navigation",
        "Founding member pricing — REMOVED",
    ]:
        bullet(doc, b)

    doc.save(path)
    print(f"✓ FDD saved: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# 6. SOFTWARE DESIGN DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

def create_sdd(path):
    doc = new_doc()
    cover_block(doc, "Huddl Connect",
        "Software Design Document (SDD)",
        VERSION, DATE, "Classification: Confidential — Internal Engineering")

    heading(doc, "1. Introduction", 1, ORANGE)
    heading(doc, "1.1 Purpose", 2)
    body(doc,
        "Describes the software architecture, component design, data models, "
        "and service integrations of Huddl Connect. Intended audience: engineering team, "
        "technical co-founders, and technical due diligence reviewers.")

    heading(doc, "1.2 Technology Stack", 2)
    for b in [
        "Flutter 3.35.4 / Dart 3.9.2 — cross-platform mobile + web",
        "Firebase Firestore — real-time database",
        "Firebase Auth — phone OTP authentication",
        "Firebase Storage — media file storage",
        "Firebase Crashlytics — crash reporting",
        "Google Gemini API — all AI features",
        "Apple IAP + Google Play Billing — subscription payments",
        "Hive 2.2.3 + hive_flutter 1.1.0 — local document storage",
        "shared_preferences 2.5.3 — key-value local storage",
        "Provider 6.1.5 — state management",
        "http 1.5.0 — HTTP client",
    ]:
        bullet(doc, b)

    heading(doc, "2. System Architecture", 1, ORANGE)
    heading(doc, "2.1 High-Level Architecture", 2)
    for b in [
        "Presentation Layer: Flutter screens, widgets, theme (HuddlColors, Google Fonts Poppins)",
        "Business Logic Layer: 54 singleton services (AI, community, payments, auth, GDPR)",
        "Data Layer: Firebase Firestore (cloud), Hive (local), shared_preferences (settings)",
        "Navigation Layer: Named routes via AppRouter (config/router.dart)",
        "AI Layer: Google Gemini API via ai_api_helper.dart + gemini_system_prompt_builder.dart",
    ]:
        bullet(doc, b)

    heading(doc, "2.2 Navigation Architecture", 2)
    body(doc,
        "MainShell uses an IndexedStack with 5 child screens and a custom floating "
        "bottom navigation bar. Named routes handle deep-linking and screen transitions. "
        "MainShell.shellKey provides global tab switching from child screens.")

    heading(doc, "3. Service Layer — Complete Service Inventory", 1, ORANGE)
    services = [
        ("AI Services", [
            "ai_api_helper.dart — Gemini API calls, retry logic",
            "ai_chat_summariser_service.dart — group message summarisation",
            "ai_copilot_service.dart — conversational parenting assistant",
            "ai_event_discovery_service.dart — daily local event crawl",
            "ai_event_recommender_service.dart — personalised event scoring",
            "ai_feed_service.dart — smart home feed curation",
            "ai_knowledge_base_service.dart — borough knowledge base",
            "ai_learning_engine_service.dart — behavioural learning",
            "ai_listing_service.dart — marketplace listing generation",
            "ai_matchmaker_service.dart — parent compatibility scoring",
            "borough_ai_context.dart — borough context injection",
            "discover_ai_service.dart — Discover tab AI",
            "gemini_system_prompt_builder.dart — prompt construction",
            "invisible_ai_service.dart — invisible AI layer",
            "meetup_ai_service.dart — meetup AI assistance",
            "messages_ai_service.dart — in-chat AI",
            "daily_ai_refresh_service.dart — daily AI refresh cycle",
        ]),
        ("Community Services", [
            "community_feed_service.dart — feed data aggregation",
            "announcement_service.dart — borough announcements",
            "default_group_service.dart — auto-group assignment on sign-up",
            "group_prepopulation_service.dart — group seeding",
            "meetup_service.dart — meetup CRUD",
            "event_service.dart — event CRUD",
            "meetup_prepopulation_service.dart — meetup seeding",
            "dm_service.dart — direct messaging",
            "poll_service.dart — group polls",
            "saved_message_service.dart — message bookmarking",
            "message_search_service.dart — full-text search",
            "attachment_service.dart — media attachments",
            "media_attach_service.dart — media upload",
            "member_photo_service.dart — profile photos",
            "invitation_service.dart — group invitations",
            "block_service.dart — user blocking",
        ]),
        ("Auth & User Services", [
            "firebase_auth_service.dart — phone OTP auth, user profile",
            "onboarding_data_service.dart — onboarding state",
            "biometric_auth_service.dart — Face ID / fingerprint",
            "permission_service.dart — runtime permissions",
            "otp_service.dart — OTP helpers",
            "tutorial_service.dart — first-run tutorial state",
        ]),
        ("Marketplace Services", [
            "rehome_service.dart — rehome listing CRUD",
            "firestore_service.dart — Firestore abstraction layer",
        ]),
        ("Subscription & Payments", [
            "subscription_service.dart — tier management, feature gating, usage tracking",
            "payment_service.dart — Apple IAP + Google Play Billing",
        ]),
        ("Infrastructure Services", [
            "borough_scope_guard.dart — enforces borough-scoped content",
            "borough_cache_service.dart — borough data caching",
            "borough_analytics_service.dart — borough engagement analytics",
            "gdpr_borough_data_service.dart — GDPR data export/deletion",
            "backup_restore_service.dart — JSON backup/restore",
            "postcode_service.dart — postcode → borough resolution",
            "browser_storage.dart — web-compatible local storage",
            "backend_api_service.dart — backend API abstraction",
            "feedback_service.dart — user feedback",
        ]),
    ]
    for section, items in services:
        heading(doc, section, 2)
        for item in items:
            bullet(doc, item)

    heading(doc, "4. Subscription & Feature Gating", 1, ORANGE)
    body(doc,
        "SubscriptionService manages three tiers (Explorer, Neighbourhood, Inner Circle) "
        "with feature limits defined in TierLimits. Usage is tracked in memory and "
        "persisted via BrowserStorage. Feature gates check both boolean access "
        "and usage count limits before allowing actions.")
    tier_table(doc)
    doc.add_paragraph()

    heading(doc, "5. Removed Integrations", 1, ORANGE)
    body(doc, "The following integrations have been removed from the codebase and docs:", color=MID)
    for b in [
        "RevGlueService — file exists (revglue_service.dart) but is NOT wired to any screen or navigation",
        "AiOffersService — file exists (ai_offers_service.dart) but is NOT wired to any screen or navigation",
        "Deals/Offers tab — removed from MainShell navigation",
        "Trips tab — removed from MainShell navigation",
        "All RevGlue API endpoints and Publisher ID references — no longer active",
    ]:
        bullet(doc, b)

    doc.save(path)
    print(f"✓ SDD saved: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# 7. DEPLOYMENT GUIDE
# ══════════════════════════════════════════════════════════════════════════════

def create_deployment(path):
    doc = new_doc()
    cover_block(doc, "Huddl Connect",
        "Deployment & Configuration Guide",
        VERSION, DATE, "Classification: Confidential — Internal Engineering")

    heading(doc, "1. Project Overview", 1, ORANGE)
    body(doc,
        "Huddl Connect is a Flutter mobile app targeting iOS and Android, "
        "with a web preview build for testing. Backend: Firebase. AI: Google Gemini. "
        "Payments: Apple IAP + Google Play Billing. No RevGlue integration.")

    heading(doc, "2. Environment Setup", 1, ORANGE)
    heading(doc, "2.1 Prerequisites", 2)
    for b in [
        "Flutter SDK 3.35.4 (locked — do not upgrade)",
        "Dart SDK 3.9.2 (locked — do not upgrade)",
        "Android Studio with Android SDK API 35",
        "Java 17 (OpenJDK 17.0.2)",
        "Xcode 15+ (for iOS builds)",
        "Firebase project configured (Firestore, Auth, Storage, Crashlytics)",
        "Google Gemini API key configured",
        "Apple Developer account (for IAP + TestFlight)",
        "Google Play Developer account (for Play Billing)",
    ]:
        bullet(doc, b)

    heading(doc, "2.2 Firebase Configuration", 2)
    for b in [
        "Place google-services.json in android/app/",
        "Place GoogleService-Info.plist in ios/Runner/",
        "Ensure firebase_options.dart has Web and Android configurations",
        "Firestore security rules configured for development/production",
        "Firebase Auth: Phone provider enabled",
        "Firebase Storage: rules configured for user media",
    ]:
        bullet(doc, b)

    heading(doc, "2.3 Gemini API Configuration", 2)
    for b in [
        "Google Gemini API key stored securely (not hardcoded in source)",
        "API key injected via --dart-define at build time",
        "Rate limiting handled in ai_api_helper.dart",
        "Model version: Gemini (latest stable)",
    ]:
        bullet(doc, b)

    heading(doc, "3. Build Commands", 1, ORANGE)
    heading(doc, "3.1 Development (Web Preview)", 2)
    body(doc, "flutter pub get\nflutter build web --release\npython3 -m http.server 5060 --directory build/web --bind 0.0.0.0")

    heading(doc, "3.2 Release (Android APK)", 2)
    body(doc, "flutter pub get\nflutter build apk --release")

    heading(doc, "3.3 Release (Android App Bundle)", 2)
    body(doc, "flutter pub get\nflutter build appbundle --release")

    heading(doc, "3.4 iOS (Xcode Archive)", 2)
    body(doc, "flutter pub get\ncd ios && pod install && cd ..\nOpen ios/Runner.xcworkspace in Xcode\nProduct → Archive → Distribute App → TestFlight")

    heading(doc, "4. Dependencies (Key Packages)", 1, ORANGE)
    deps = [
        ("firebase_core",          "3.6.0",  "Firebase Core SDK"),
        ("firebase_auth",          "5.x",    "Phone OTP authentication"),
        ("cloud_firestore",        "5.4.3",  "Real-time database"),
        ("firebase_storage",       "12.3.2", "Media file storage"),
        ("firebase_crashlytics",   "4.1.3",  "Crash reporting"),
        ("google_mobile_ads",      "5.3.1",  "AdMob (future use)"),
        ("provider",               "6.1.5",  "State management"),
        ("hive",                   "2.2.3",  "Local document database"),
        ("hive_flutter",           "1.1.0",  "Hive Flutter integration"),
        ("shared_preferences",     "2.5.3",  "Key-value storage"),
        ("http",                   "1.5.0",  "HTTP client"),
        ("google_fonts",           "6.x",    "Poppins typeface"),
        ("in_app_purchase",        "3.x",    "Apple IAP + Google Play Billing"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(["Package", "Version", "Purpose"]):
        hdr[i].text = h
        set_cell_bg(hdr[i], "FF6B35")
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
    for pkg, ver, purpose in deps:
        row = table.add_row().cells
        row[0].text = pkg
        row[1].text = ver
        row[2].text = purpose
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

    heading(doc, "5. GitHub Repository", 1, ORANGE)
    for b in [
        "Repository: github.com/cdcruz007/huddl",
        "Branch: main",
        "Docs folder: /docs/ — all design and business documents",
        ".gitignore: excludes build/, android/build, ios/Pods/, .dart_tool/",
        "pubspec.lock: committed for dependency consistency",
    ]:
        bullet(doc, b)

    heading(doc, "6. Removed Configuration (No Longer Required)", 1, ORANGE)
    body(doc, "The following configuration is no longer needed:", color=MID)
    for b in [
        "RevGlue account and Publisher ID — REMOVED",
        "RevGlue iframeapi endpoints — REMOVED",
        "Deals/Offers tab configuration — REMOVED",
    ]:
        bullet(doc, b)

    doc.save(path)
    print(f"✓ Deployment Guide saved: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# 8. STORE PUBLISHING GUIDE
# ══════════════════════════════════════════════════════════════════════════════

def create_store_guide(path):
    doc = new_doc()
    cover_block(doc, "Huddl Connect",
        "Store Publishing & Monetisation Guide",
        VERSION, DATE, "Classification: Confidential — Internal Product & Commercial")

    heading(doc, "1. App Identifiers", 1, ORANGE)
    for b in [
        "App name: Huddl Connect",
        "Android package: com.huddlconnect.app (verify in android/app/build.gradle.kts)",
        "iOS bundle ID: com.huddlconnect.app (verify in Xcode → Runner target)",
        "Version: 1.0.0 (initial launch)",
        "Category: Social Networking / Parenting",
    ]:
        bullet(doc, b)

    heading(doc, "2. Revenue Streams", 1, ORANGE)
    heading(doc, "2.1 Subscription Revenue (Primary)", 2)
    for b in [
        "Explorer: Free — conversion funnel entry point",
        "Neighbourhood: £5.99/month or £49.99/year (~30% annual saving)",
        "Inner Circle: £11.99/month or £99.99/year (~30% annual saving)",
        "7-day free trial on first paid subscription",
        "Managed via Apple IAP (iOS) and Google Play Billing (Android)",
    ]:
        bullet(doc, b)

    heading(doc, "2.2 B2B Revenue (Secondary)", 2)
    for b in [
        "Featured event listings for local businesses",
        "Sponsored group creation",
        "Borough partnership packages (councils, NHS)",
        "No affiliate/RevGlue revenue — removed from platform",
    ]:
        bullet(doc, b)

    heading(doc, "3. Google Play Store Listing", 1, ORANGE)
    heading(doc, "3.1 Store Description", 2)
    body(doc,
        "Huddl Connect — Find Your Local Tribe.\n\n"
        "The hyper-local community platform built for parents. "
        "Huddl connects you with parents in your borough — your actual neighbours — "
        "so you can share, support, and thrive together.\n\n"
        "Unlike national parenting apps, Huddl is borough-first. Everything you see "
        "— groups, events, marketplace listings — is local to you.")

    heading(doc, "3.2 Key Features for Store Listing", 2)
    for b in [
        "Join local parent groups matched to your postcode and parenting stage",
        "Discover AI-curated events and parent meetups near you — daily",
        "Buy, sell and give away preloved baby and children's items locally",
        "Chat with your AI-powered parenting Copilot — borough-aware and personal",
        "Create and RSVP to local parent meetups",
        "Get AI summaries of group conversations you've missed",
        "Your Digital Village — the 'next-door mum and dad' feeling, in your pocket",
    ]:
        bullet(doc, b)

    heading(doc, "3.3 Screenshots Required", 2)
    for b in [
        "Home screen (AI smart feed with borough badge)",
        "Connect tab (group chat with AI summary button)",
        "Discover tab (local events with AI discovery banner)",
        "Market tab (preloved listings with AI listing generator)",
        "Subscription plans screen (tier comparison)",
        "AI Copilot screen (conversation example)",
        "Onboarding — welcome/postcode screen",
    ]:
        bullet(doc, b)

    heading(doc, "4. Subscription IAP Configuration", 1, ORANGE)
    heading(doc, "4.1 Apple App Store (IAP)", 2)
    for b in [
        "Product IDs: neighbourhood_monthly, neighbourhood_annual, innercircle_monthly, innercircle_annual",
        "Free trial: 7 days on first subscription",
        "Configure in App Store Connect → In-App Purchases → Subscriptions",
        "Subscription group: 'Huddl Connect Premium'",
    ]:
        bullet(doc, b)

    heading(doc, "4.2 Google Play Billing", 2)
    for b in [
        "Product IDs: neighbourhood_monthly, neighbourhood_annual, innercircle_monthly, innercircle_annual",
        "Free trial: 7 days on first subscription",
        "Configure in Google Play Console → Monetise → Subscriptions",
        "Base plan + offer configuration for free trial",
    ]:
        bullet(doc, b)

    heading(doc, "5. Content Rating & Compliance", 1, ORANGE)
    for b in [
        "PEGI / ESRB: 4+ (family-friendly content)",
        "GDPR compliant: UK ICO registered, privacy policy in-app",
        "Children's Code aware: no data collection from children under 13",
        "Phone verification required: prevents anonymous accounts",
        "Content moderation: block/report tools built in",
        "No RevGlue or affiliate tracking — no third-party advertising cookies",
    ]:
        bullet(doc, b)

    doc.save(path)
    print(f"✓ Store Guide saved: {path}")


# ══════════════════════════════════════════════════════════════════════════════
# MAIN — Generate all documents
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    files = {
        "Huddl_Connect_Go_To_Market_Strategy.docx":   create_gtm,
        "Huddl_Connect_Investor_Pitch_Deck.docx":     create_pitch_deck,
        "Huddl_Connect_AI_System_Design_Document.docx": create_ai_doc,
        "Huddl_Connect_UX_Design_Documentation.docx": create_ux_doc,
        "Huddl_Connect_Functional_Design_Document.docx": create_fdd,
        "Huddl_Connect_Software_Design_Document.docx": create_sdd,
        "Huddl_Connect_Deployment_Guide.docx":        create_deployment,
        "Huddl_Connect_Store_Publishing_Guide.docx":  create_store_guide,
    }

    for filename, fn in files.items():
        full_path = os.path.join(OUTPUT_DIR, filename)
        try:
            fn(full_path)
        except Exception as e:
            print(f"✗ FAILED {filename}: {e}")
            import traceback; traceback.print_exc()

    # Also remove old trips supplement if present
    old_trips = os.path.join(OUTPUT_DIR, "Huddl_Trips_AI_Feature_Pitch_Deck_Supplement.docx")
    if os.path.exists(old_trips):
        os.remove(old_trips)
        print(f"✓ Removed old Trips supplement")

    print("\n✅ All documents generated.")
    print(f"Output directory: {OUTPUT_DIR}")
    import subprocess
    result = subprocess.run(["ls", "-lh", OUTPUT_DIR], capture_output=True, text=True)
    print(result.stdout)

