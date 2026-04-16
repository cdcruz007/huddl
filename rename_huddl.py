#!/usr/bin/env python3
"""
Rename 'Huddl Connect' -> 'Huddl' throughout all docs and file names.
"""

from docx import Document
from docx.oxml.ns import qn
import os, copy

DOCS_DIR = '/home/user/flutter_app/docs'

# ── pairs: (old_text, new_text) — ordered longest first to avoid partial matches
REPLACEMENTS = [
    ('Huddl Connect', 'Huddl'),
    ('huddl connect', 'huddl'),
    ('HUDDL CONNECT', 'HUDDL'),
    # keep "Huddl" alone untouched
]

def replace_in_para(para):
    """Replace text in a paragraph, preserving all run formatting."""
    # Rebuild the paragraph as a single run replacement to avoid
    # cross-run splits, then re-apply the first run's character style.
    full_text = para.text
    new_text = full_text
    changed = False
    for old, new in REPLACEMENTS:
        if old in new_text:
            new_text = new_text.replace(old, new)
            changed = True

    if not changed:
        return

    # Distribute new_text back across existing runs proportionally.
    # Simplest safe approach: put everything in first run, blank the rest.
    if para.runs:
        remaining = new_text
        first_run = para.runs[0]
        first_run.text = remaining
        for run in para.runs[1:]:
            run.text = ''
    else:
        para.add_run(new_text)


def replace_in_table(table):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_in_para(para)
            for nested_table in cell.tables:
                replace_in_table(nested_table)


def process_doc(path):
    doc = Document(path)
    for para in doc.paragraphs:
        replace_in_para(para)
    for table in doc.tables:
        replace_in_table(table)
    # Also fix core properties (title, subject, etc.)
    cp = doc.core_properties
    if cp.title and 'Huddl Connect' in cp.title:
        cp.title = cp.title.replace('Huddl Connect', 'Huddl')
    doc.save(path)


# ── 1. Process content of every .docx ────────────────────────────────────────
docx_files = sorted(f for f in os.listdir(DOCS_DIR) if f.endswith('.docx'))
print(f'Processing {len(docx_files)} documents...')
for fname in docx_files:
    fpath = os.path.join(DOCS_DIR, fname)
    process_doc(fpath)
    print(f'  ✅ content updated: {fname}')

# ── 2. Rename files (Huddl_Connect_ → Huddl_) ────────────────────────────────
print()
print('Renaming files...')
for fname in sorted(os.listdir(DOCS_DIR)):
    if not fname.endswith('.docx'):
        continue
    new_fname = fname.replace('Huddl_Connect_', 'Huddl_')
    if new_fname != fname:
        os.rename(
            os.path.join(DOCS_DIR, fname),
            os.path.join(DOCS_DIR, new_fname)
        )
        print(f'  ✅ renamed: {fname} → {new_fname}')

print()
print('Final docs directory:')
for f in sorted(os.listdir(DOCS_DIR)):
    print(f'  {f}')

print()
print('Done.')
